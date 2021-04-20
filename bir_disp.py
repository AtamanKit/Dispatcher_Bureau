from PySide2.QtGui import QPixmap
from PySide2.QtWidgets import QSplashScreen, QApplication
import sys

app = QApplication(sys.argv)
#Making a splash screen
pixmap = QPixmap("ataman_splash.png")
splash = QSplashScreen(pixmap)
splash.show()
splash.showMessage("Loading...")
app.processEvents()

from PySide2.QtGui import QIcon, QBrush, QColor, QFont, QPainter, QRegExpValidator
from PySide2.QtCore import QAbstractTableModel, Qt, QRegExp, QSize, QThreadPool, Signal, QObject, QRunnable
from PySide2.QtWidgets import QMdiSubWindow, QLineEdit, QMainWindow, QMdiArea, \
    QMenu, QAction, QComboBox, QLabel, QFrame, QVBoxLayout, QHBoxLayout, QSplitter, \
    QTableView, QAbstractItemView, QDialog, QCompleter, QGridLayout, QPushButton, QFileDialog, \
    QMessageBox, QTextEdit, QCheckBox, QCalendarWidget
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from qtwidgets import PasswordEdit
from pymongo import MongoClient
from docx import Document
from docx2pdf import convert
from playsound import playsound
from twilio.rest import Client
from matplotlib.backends.backend_qt5agg import (
    FigureCanvas, NavigationToolbar2QT as NavigationToolbar)
from matplotlib.figure import Figure

import pandas as pd
import numpy as np
import pymongo
import re
import datetime
import random
import webbrowser
import os.path
import shutil
import psycopg2


#Working with google API
import pickle
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from apiclient.http import MediaFileUpload

class tabSubWindow(QMdiSubWindow):
    def __init__(self):
        QMdiSubWindow.__init__(self)
    # def closeEvent(self, event):
    #     mainWindow.nameRegClose(self)
    #     # print("test")

class tabDecSubWd(QMdiSubWindow):
    def __init__(self):
        QMdiSubWindow.__init__(self)
    def closeEvent(self, event):
        pass
        # mainWindow.nameRegClose(self)
        # print("test")

#Formez Model/View
class TableModel(QAbstractTableModel):

    def __init__(self, data, headerdata):
        super(TableModel, self).__init__()
        self._data = data
        self.headerdata = headerdata

    def data(self, index, role):
        if role == Qt.DisplayRole:
            value = self._data.iloc[index.row(), index.column()]
            # if isinstance(value, int):
            #     return str(value)
            # else:
            return str(value)

        if role == Qt.TextAlignmentRole:
            value = self._data.iloc[index.row(), index.column()]
            matchedLink = re.search(":/", str(value))
            # matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", str(value))
            # matchedDelta = re.match("\d:\d\d:\d\d", str(value))
            # matchedLa = re.search("La ", str(value))
            # matchedInr = re.search("\d\d.\d\d.\d\d", str(value))
            # matchedAdm = re.search("admis", str(value))
            # matchedFider = re.search("Fider", str(value))
            # matchedFormatia = re.search("Formatia:", str(value))

            # if len(str(value)) <= 9 or bool(matched) or bool(matchedShort) or bool(matchedLa) or \
            #         str(value) == "Nu s-a lucrat" or str(value) == "Nou inregistrata" \
            #         or str(value) == "Programat" or str(value) == "Fara deconectari" \
            #         or str(value) == "Neprogramat" or bool(matchedDelta) \
            #         or str(value) == "/ ... /" or str(value) == "Intrerupere"\
            #         or bool(matchedInr) or bool(matchedAdm) or str(value) == "-//-" \
            #         or bool(matchedFider) or bool(matchedFormatia):
            #     return Qt.AlignCenter
            if matchedLink:
                return Qt.AlignVCenter
            if str(value):
                return Qt.AlignCenter
        if role == Qt.BackgroundRole:
            value = self._data.iloc[index.row(), index.column()]
            for i in range(10):
                value_t = self._data.iloc[index.row(), index.column() - i]
                if value_t == "TOTAL":
                    return QColor(80, 80, 80)
            # value_8 = self._data.iloc[index.row(), index.column() - 8]
            # value_7 = self._data.iloc[index.row(), index.column() - 7]
            matched = re.match("\d\d.\d\d.\d\d\d\d \d\d:\d\d", str(value))
            matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", str(value))
            # matchedDelta = re.match("\d:\d\d:\d\d", str(value))
            matchedPT = re.search("PT\d", str(value))
            matchedMont = re.search("/mont", str(value))
            matchedAdm = re.search("Admitere:", str(value))
            matchedGr = re.search("gr. ", str(value))
            matchedFormatia = re.search("Formatia:", str(value))
            matchedExcl = re.search("!", str(value))
            matchedExec = re.search("Executat", str(value))
            matchedAcc = re.search("Acceptat", str(value))
            matchedRef = re.search("Refuz", str(value))
            matchedPreg = re.search("Pregatire:", str(value))
            matchedTerm = re.search("Terminat:", str(value))
            matchedLink = re.search(":/", str(value))
            matchedConf = re.search("Confirmat:", str(value))
            matchedCerere = re.search("Cerere", str(value))
            matched2lines = re.search("--", str(value))
            matchedDepasit = re.search("Depasit", str(value))
            matchedIncadrat = re.search("Incadrat", str(value))
            if str(value) == "Nou inregistrata" or str(value) == "Neexecutat" or matchedCerere or \
                    matchedDepasit:
                # print(str(valuePlus))
                return QColor(179, 48, 48)
            elif bool(matchedShort) or bool(matchedRef)\
                    or bool(matchedTerm):
                return QColor(80, 80, 80)
            elif bool(matched2lines):
                return QColor(80, 80, 80)
            elif bool(matched) or str(value) == "/ ... /" or bool(matchedAdm):
                return QColor(194, 199, 40)
            elif str(value) == "Nu s-a lucrat":
                return QColor(13, 158, 163)
            elif str(value) == "/ --- /":
                return QColor(13, 158, 163)
            elif str(value) == "Programat":
                return QColor(177, 191, 147)
            elif str(value) == "Neprogramat":
                return QColor(209, 169, 169)
            elif str(value) == "UN":
                return QColor(180, 60, 60)
            elif str(value) == "FL":
                return QColor(87, 107, 148)
            elif str(value) == "GL":
                return QColor(150, 150, 150)
            elif str(value) == "RS":
                return QColor(130, 100, 200)
            elif bool(matchedIncadrat):
                return QColor(70, 122, 78)
            elif bool(matchedMont):
                return QColor(79, 74, 128)
            elif bool(matchedLink):
                return QColor(121, 113, 145)
            elif bool(matchedPT):
                return QColor(51, 105, 143)
            elif bool(matchedGr) and not bool(matchedFormatia) and not bool(matchedConf):
                return QColor(158, 81, 81)
            elif str(value) == "Intrerupere":
                return QColor(46, 130, 68)
            elif bool(matchedExcl):
                return QColor(13, 110, 106)
            elif bool(matchedExec):
                return QColor(68, 112, 71)
            elif bool(matchedAcc):
                return QColor(92, 156, 107)
            elif bool(matchedPreg):
                return QColor(35, 99, 145)
            elif bool(matchedConf):
                return QColor(92, 156, 107)
        if role == Qt.ForegroundRole:
            value = self._data.iloc[index.row(), index.column()]
            for i in range(10):
                value_t = self._data.iloc[index.row(), index.column() - i]
                if value_t == "TOTAL":
                    return QColor(255, 255, 255)
            matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", str(value))
            matchedRef = re.search("Refuz", str(value))
            matchedPreg = re.search("Pregatire:", str(value))
            matchedTerm = re.search("Terminat:", str(value))
            matched2lines = re.search("--", str(value))
            matchedCerere = re.search("Cerere", str(value))
            matchedDepasit = re.search("Depasit", str(value))
            if str(value) == "Nou inregistrata" or bool(matchedRef)\
                    or bool(matchedShort) or str(value) == "Neexecutat"\
                    or bool(matchedPreg) or bool(matchedTerm) or matched2lines \
                    or matchedCerere or matchedDepasit:
                return QColor(255, 255, 255)
            # elif str(value) == "Programat":
            #     return QColor(255, 255, 255)
        if role == Qt.FontRole:
            value = self._data.iloc[index.row(), index.column()]
            # matched = re.match("\d\d.\d\d.\d\d\d\d \d\d:\d\d", str(value))
            matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", str(value))
            matchedDelta = re.match("\d:\d\d:\d\d", str(value))
            if str(value) == "Nou inregistrata" or str(value) == "Nu s-a lucrat" \
                    or bool(matchedShort) or bool(matchedDelta):
                return QFont("Calibri", 13, QFont.Bold)

            else:
                return QFont("Calibri", 13, QFont.Bold)


    def rowCount(self, index):
        return self._data.shape[0]

    def columnCount(self, index):
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        if orientation == Qt.Horizontal and role == Qt.DisplayRole:
            return self.headerdata[section]
        if role == Qt.BackgroundRole:
            return QColor(66, 135, 245)
        if role == Qt.ForegroundRole:
            return QColor(0, 0, 0)
        if role == Qt.FontRole:
            return QFont("Calibri", 10, QFont.Bold)
        if role == Qt.DecorationRole:
            return QColor(66, 135, 245)

        # if orientation == Qt.Vertical and role == Qt.DisplayRole:
        #     return self.headerdataH[section]
        # if role == Qt.BackgroundRole:
        #     return QColor(66, 135, 245)
        # if role == Qt.ForegroundRole:
        #     return QColor(0, 0, 0)
        # if role == Qt.FontRole:
        #     return QFont("Calibri", 10, QFont.Bold)
        # if role == Qt.DecorationRole:
        #     return QColor(66, 135, 245)

class TableModelDeranj(QAbstractTableModel):

    def __init__(self, data, headerdata):
        super(TableModelDeranj, self).__init__()
        self._data = data
        self.headerdata = headerdata

    def data(self, index, role):
        if role == Qt.DisplayRole:
            value = self._data.iloc[index.row(), index.column()]
            # if isinstance(value, int):
            #     return str(value)
            # else:
            return value

        if role == Qt.TextAlignmentRole:
            value = self._data.iloc[index.row(), index.column()]
            if str(value):
                return Qt.AlignCenter
        if role == Qt.BackgroundRole:
            value = self._data.iloc[index.row(), index.column()]
            # value_minus = self._data.iloc[index.row(), index.column() - 1]
            matched = re.match("\d\d.\d\d.\d\d\d\d \d\d:\d\d", str(value))
            matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", str(value))
            matchedDelta = re.match("\d:\d\d:\d\d", str(value))
            matchedPT = re.search("PT\d", str(value))
            # matchedLa = re.search("La ", str(value))
            matchedGr = re.search("gr. ", str(value))
            matchedFormatia = re.search("Formatia:", str(value))
            matchedExcl = re.search("!", str(value))
            matchedExec = re.search("Executat", str(value))
            if str(value) == "Nou inregistrata" or str(value) == "Neexecutat":
                # print(str(valuePlus))
                return QColor(179, 48, 48)
            elif bool(matchedShort):
                return QColor(80, 80, 80)
            elif str(value) == "UN":
                return QColor(180, 60, 60)
            elif str(value) == "FL":
                return QColor(87, 107, 148)
            elif str(value) == "GL":
                return QColor(150, 150, 150)
            elif str(value) == "RS":
                return QColor(130, 100, 200)
            elif bool(matchedDelta):
                return QColor(70, 122, 78)
            # elif bool(matchedLa):
            #     return QColor(79, 74, 128)
            elif bool(matchedPT):
                return QColor(51, 105, 143)
            elif bool(matchedGr) and not bool(matchedFormatia):
                return QColor(158, 81, 81)
            elif bool(matchedExcl):
                return QColor(13, 110, 106)
            elif bool(matchedExec):
                return QColor(68, 112, 71)
        if role == Qt.ForegroundRole:
            value = self._data.iloc[index.row(), index.column()]
            matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", str(value))

            if str(value) == "Nou inregistrata"\
                    or bool(matchedShort) or str(value) == "Neexecutat":
                return QColor(255, 255, 255)
        if role == Qt.FontRole:
            return QFont("Calibri", 14, QFont.Bold)


    def rowCount(self, index):
        return self._data.shape[0]

    def columnCount(self, index):
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        if orientation == Qt.Horizontal and role == Qt.DisplayRole:
            return self.headerdata[section]
        if role == Qt.BackgroundRole:
            return QColor(66, 135, 245)
        if role == Qt.ForegroundRole:
            return QColor(0, 0, 0)
        if role == Qt.FontRole:
            return QFont("Calibri", 12, QFont.Bold)
        if role == Qt.DecorationRole:
            return QColor(66, 135, 245)

class TableModelII(QAbstractTableModel):

    def __init__(self, data, headerdata):
        super(TableModelII, self).__init__()
        self._data = data
        self.headerdata = headerdata

    def data(self, index, role):
        if role == Qt.DisplayRole or role == Qt.EditRole:
            value = self._data.iloc[index.row(), index.column()]
            return value

        if role == Qt.TextAlignmentRole:
            # value = self._data.iloc[index.row(), index.column()]
            return Qt.AlignCenter
        if role == Qt.BackgroundRole:
            value = self._data.iloc[index.row(), index.column()]
            if str(value) == "UN":
                return QColor(180, 60, 60)
            elif str(value) == "FL":
                return QColor(87, 107, 148)
            elif str(value) == "GL":
                return QColor(150, 150, 150)
            elif str(value) == "RS":
                return QColor(130, 100, 200)
        if role == Qt.FontRole:
            return QFont("Calibri", 14, QFont.Bold)

    def flags(self, index):
        flag = super(TableModelII, self).flags(index)
        return flag | Qt.ItemIsEditable

    def setData(self, index, value, role):
        if role == Qt.EditRole:
            self._data.iloc[index.row(), index.column()] = value
            # self.dataChanged.emit(index, index)
            mainWindow.dtContrDecZl(self)
            self.wsNrSol.cell(row=index.row()+5, column=index.column()+1).value = value
            self.wsNrSol.cell(row=index.row()+5, column=index.column()+1).alignment = \
                Alignment(horizontal="center", vertical="center")
            self.wbDecZil.save(self.rapFile)
            return True

    def destLoad(self):
        self.myDest = os.path.abspath(".") + "/Bundle/Destination.xlsx"
        self.wbDest = load_workbook(self.myDest)
        self.wsDest = self.wbDest["Destination"]


    def rowCount(self, index):
        return self._data.shape[0]

    def columnCount(self, index):
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        if orientation == Qt.Horizontal and role == Qt.DisplayRole:
            return self.headerdata[section]
        if role == Qt.BackgroundRole:
            return QColor(66, 135, 245)
        if role == Qt.ForegroundRole:
            return QColor(0, 0, 0)
        if role == Qt.FontRole:
            return QFont("Calibri", 10, QFont.Bold)
        if role == Qt.DecorationRole:
            return QColor(66, 135, 245)

class TableModelList(QAbstractTableModel):

    def __init__(self, data, headerdata):
        super(TableModelList, self).__init__()
        self._data = data
        self.headerdata = headerdata

    def data(self, index, role):
        if role == Qt.DisplayRole or role == Qt.EditRole:
            value = self._data.iloc[index.row(), index.column()]
            return value

        if role == Qt.TextAlignmentRole:
            # value = self._data.iloc[index.row(), index.column()]
            return Qt.AlignCenter
        if role == Qt.BackgroundRole:
            value = self._data.iloc[index.row(), index.column()]
            # valueMinus = self._data.iloc[index.row(), index.column()-1]
            matchedTelefon = re.search("\d\d\d\d\d", str(value))
            matchedGrTs3 = re.fullmatch("\d", str(value))
            matchedDa = re.match("DA!", str(value))
            matchedNu = re.match("NU!", str(value))
            if str(value) == "UN":
                return QColor(180, 60, 60)
            elif str(value) == "FL":
                return QColor(87, 107, 148)
            elif str(value) == "GL":
                return QColor(150, 150, 150)
            elif str(value) == "RS":
                return QColor(130, 100, 200)
            elif bool(matchedTelefon):
                return QColor(70, 153, 119)
            elif bool(matchedGrTs3):
                return QColor(156, 149, 103)
            elif bool(matchedDa):
                return QColor(194, 199, 40)
            elif bool(matchedNu):
                return QColor(158, 81, 81)
        if role == Qt.FontRole:
            return QFont("Calibri", 14, QFont.Bold)

    # def flags(self, index):
    #     flag = super(TableModelList, self).flags(index)
    #     return flag | Qt.ItemIsEditable
    #
    # def setData(self, index, value, role):
    #     if role == Qt.EditRole:
    #         self._data.iloc[index.row(), index.column()] = value
    #         # self.dataChanged.emit(index, index)
    #         mainWindow.dtContrDecZl(self)
    #         self.wsNrSol.cell(row=index.row()+5, column=index.column()+1).value = value
    #         self.wsNrSol.cell(row=index.row()+5, column=index.column()+1).alignment = \
    #             Alignment(horizontal="center", vertical="center")
    #         self.wbDecZil.save(self.rapFile)
    #         return True
    #
    # def destLoad(self):
    #     self.myDest = os.path.abspath(".") + "/Bundle/Destination.xlsx"
    #     self.wbDest = load_workbook(self.myDest)
    #     self.wsDest = self.wbDest["Destination"]


    def rowCount(self, index):
        return self._data.shape[0]

    def columnCount(self, index):
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        if orientation == Qt.Horizontal and role == Qt.DisplayRole:
            return self.headerdata[section]
        if role == Qt.BackgroundRole:
            return QColor(66, 135, 245)
        if role == Qt.ForegroundRole:
            return QColor(0, 0, 0)
        if role == Qt.FontRole:
            return QFont("Calibri", 10, QFont.Bold)
        if role == Qt.DecorationRole:
            return QColor(66, 135, 245)

class TableModelIII(QAbstractTableModel):

    def __init__(self, data):
        super(TableModelIII, self).__init__()
        self._data = data
        # self.headerdata = headerdata

    def data(self, index, role):
        if role == Qt.DisplayRole or role == Qt.EditRole:
            value = self._data.iloc[index.row(), index.column()]
            return value

        if role == Qt.BackgroundRole:
            value = self._data.iloc[index.row(), index.column()]
            if str(value) == "UN":
                return QColor(180, 60, 60)
            elif str(value) == "FL":
                return QColor(87, 107, 148)
            elif str(value) == "GL":
                return QColor(150, 150, 150)
            elif str(value) == "RS":
                return QColor(130, 100, 200)
        if role == Qt.FontRole:
            return QFont("Calibri", 14, QFont.Bold)

        if role == Qt.TextAlignmentRole:
            value = self._data.iloc[index.row(), index.column()]
            matched = re.match("\d", str(value))
            if bool(matched) or \
                    str(value) == "UN" or str(value) == "FL" \
                    or str(value) == "GL" or str(value) == "RS":
                return Qt.AlignCenter

    def flags(self, index):
        flag = super(TableModelIII, self).flags(index)
        return flag | Qt.ItemIsEditable

    def setData(self, index, value, role):
        if role == Qt.EditRole:
            self._data.iloc[index.row(), index.column()] = value
            self.dataChanged.emit(index, index)
            mainWindow.dtContrDecZl(self)
            self.wsNrSol.cell(row=index.row()+13, column=index.column()+1).value = value
            self.wbDecZil.save(self.rapFile)
            return True

    def destLoad(self):
        self.myDest = os.path.abspath(".") + "/Bundle/Destination.xlsx"
        self.wbDest = load_workbook(self.myDest)
        self.wsDest = self.wbDest["Destination"]

    def rowCount(self, index):
        return self._data.shape[0]

    def columnCount(self, index):
        return self._data.shape[1]


#Folosesc acest clas pentru a ma putea folosi de focus out, controlez oara QLineEdit
class DateLineEdit(QLineEdit):

    def __init__(self):
        QLineEdit.__init__(self)

    # def focusInEvent(self, e):
    #     QLineEdit.focusInEvent(self, e)
    #     self.initText = self.text()
    #     self.setText("")

    def focusOutEvent(self, e):
        QLineEdit.focusOutEvent(self, e)
        matchedShort = re.match("\d\d.\d\d.\d\d \d\d:\d\d", self.text())
        if not bool(matchedShort):
            mainWindow.msSecCall(self,'Formatul necesar pentru data si ora este: \n zz.ll.aa hh:mm')
            self.setFocus()

#Folosesc acest clas pentru a ma putea folosi de focus in QLineEdit
class MyLineEdit(QLineEdit):

    def __init__(self):
        QLineEdit.__init__(self)

    def focusInEvent(self, e):
        QLineEdit.focusInEvent(self, e)
        self.initText = self.text()
        self.setText("")

    def focusOutEvent(self, e):
        QLineEdit.focusOutEvent(self, e)
        if self.text() == "":
            self.setText(self.initText)

# Multiprossesing
class WorkerSignal(QObject):
    finished = Signal()

class Worker(QRunnable):
    signal = WorkerSignal()

    def run(self):
        mainWindow.loadMongo(self)

        myTime = datetime.datetime.now()
        for i in self.client.djUN.list_collection_names():
            if i == "al_" + str(myTime.year) + "_" + str(myTime.month):
                reg_alw = self.client.djUN.i

                change_stream = reg_alw.watch([{
                    '$match': {
                            'operationType': {'$in': ['update']}
                    }}, {
                    '$project': {
                        'fullDocument_id': '$fullDocument._id',
                        'pregatire': '$fullDocument.pregatire',
                        'admitere': '$fullDocument.admitere',
                        'terminare': '$fullDocument.terminare',
                    }
                }], full_document='updateLookup')
                for change in change_stream:
                    if change['pregatire'] == 'Cerere la pregatire' \
                        or change['admitere'] == 'Cerere la admitere' \
                        or change['terminare'] == 'Cerere la terminare':
                        self.signal.finished.emit()

class WorkerSound(QRunnable):

    def run(self):
        global elSoundContr
        elSoundContr = False

        while elSound:
            playsound('Sources/Sounds/cerere.mp3')

#Clasul principal
class mainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Parent Window create
        self.myMidi = QMdiArea()
        self.setCentralWidget(self.myMidi)
        self.setWindowTitle('Biroul dispecerului')
        self.setWindowIcon(QIcon('ataman_logo'))
        self.setGeometry(100, 100, 1250, 720)
        # self.myWindow.setStyleSheet('background-color: #6e6e6e')
        self.myMidi.setBackground(QBrush(QColor(100, 100, 100)))

        self.showMaximized()

        # Menu Create

        bar = self.menuBar()

        fileMenu = bar.addMenu('File')
        newMen = QMenu('New', self)
        fileMenu.addMenu(newMen)

        autSub = QAction('Autorizatie', self)
        newMen.addAction(autSub)
        autSub.triggered.connect(self.alTrig)

        disSub = QAction('Dispozitie', self)
        newMen.addAction(disSub)
        disSub.triggered.connect(self.dsTrig)

        decSub = QAction('Deconectare neplanificata', self)
        newMen.addAction(decSub)
        decSub.triggered.connect(self.decTrig)

        deranjSub = QAction('Deranjament', self)
        newMen.addAction(deranjSub)
        deranjSub.triggered.connect(self.deranjTrig)

        setMenu = QAction("Setari", self)
        fileMenu.addAction(setMenu)
        setMenu.triggered.connect(self.setTrig)

        fileMenu.addSeparator()

        exitMenu = QAction("Exit", self)
        fileMenu.addAction(exitMenu)
        # fileMenu.addMenu(exitMenu)
        exitMenu.triggered.connect(self.close)

        regMen = bar.addMenu('Registre')
        regAl = QAction('Autorizatii si dispozitii', self)
        regMen.addAction(regAl)
        regAl.triggered.connect(self.centrAlPop)
        regDec = QAction('Raport deconectari', self)
        regMen.addAction(regDec)
        regDec.triggered.connect(self.decTabWindow)
        regDeranj = QAction('Deranjamente', self)
        regMen.addAction(regDeranj)
        regDeranj.triggered.connect(self.deranjPop)


        angMen = bar.addMenu('Angajati')
        listAng = QAction(QIcon(os.path.abspath(".") + "/Sources/Context/list_ang.png"), \
                                              'Lista angajati', self)
        angMen.addAction(listAng)
        listAng.triggered.connect(self.listAngPop)

        angMen.addSeparator()
        cautAng = QAction(QIcon(os.path.abspath(".") + "/Sources/Context/search.png"), \
                                              'Cauta angajat', self)
        angMen.addAction(cautAng)
        cautAng.triggered.connect(self.angCautTrig)
        # cautSec = QAction('Cauta dupa sector', self)
        # angMen.addAction(cautSec)
        # cautSec.triggered.connect(self.cautSectTrig)
        angMen.addSeparator()

        angNou = QAction("Angajat nou", self)
        angMen.addAction(angNou)
        angNou.triggered.connect(self.angNouTrig)
        angModif = QAction('Modificare stare', self)
        angMen.addAction(angModif)
        angModif.triggered.connect(self.angModifTrig)
        angDemis = QAction('Sterge angajat', self)
        angMen.addAction(angDemis)
        angDemis.triggered.connect(self.angDelTrig)

        rpBar = bar.addMenu('Raport')
        anualMen = QMenu("Analiza anuala", self)
        rpBar.addMenu(anualMen)

        anualProg = QAction("Programat", self)
        anualMen.addAction(anualProg)
        anualProg.triggered.connect(self.AnualProg)

        anualNepr = QAction("Neprogramat", self)
        anualMen.addAction(anualNepr)
        anualNepr.triggered.connect(self.AnualNepr)

        autMen = bar.addMenu('LogIn/Out')
        autIn = QAction("Log In", self)
        autMen.addAction(autIn)
        autMen.triggered.connect(self.intTrig)
        autOut = QAction("Log Out", self)
        autMen.addAction(autOut)
        autOut.triggered.connect(self.logOut)

        self.nameFMenu = ""

        self.angDelFunc_control = False
        self.angFunc_control = False
        self.angModifFunc_control = False

        # Variabile pentru cautarea angajat, sector
        self.cautAngControl = False
        self.cautSecAngControl = False

        self.listAngControlPop = False
        self.deranjControlPop = False
        self.tabWindowControl = False

        self.tabDecControl = False
        self.anControl = False

        self.passControl = False
        self.contextAng = False

        self.dublControl = False

        #Tabela angajati
        self.tableAng = QTableView()
        #Controlez la terminarea lucrarilor ca nu e deschis registru autorizatii
        self.erContrAl = False

        # Multiprossesing run
        worker = Worker()
        self.threadpool = QThreadPool()
        self.threadpool.start(worker)
        worker.signal.finished.connect(self.msCerere)

        #Working with dates 'registru autorizatii'
        self.alTime = datetime.datetime.now()
        self.alYear = self.alTime.strftime('%Y')
        self.alMonth = self.alTime.strftime('%m')

        self.mnCombo = QComboBox()
        self.mnList = [
            "Ianuarie",
            "Februarie",
            "Martie",
            "Aprilie",
            "Mai",
            "Iunie",
            "Iulie",
            "August",
            "Septembrie",
            "Octombrie",
            "Noiembrie",
            "Decembrie"
        ]
        self.mnCombo.addItems(self.mnList)
        # self.mnCombo.setEditable(True)
        self.mnCombo.setCurrentText(self.NumbToMonth(self.alMonth))
        self.mnCombo.setFixedHeight(25)
        self.mnCombo.setFixedWidth(100)
        self.mnCombo.setStyleSheet('padding-left: 10%')
        self.mnCombo.currentTextChanged.connect(self.chMonth)


        self.intTrig()

    def MonthToNumb(self, month):
        if month == "Ianuarie":
            return "01"
        if month == "Februarie":
            return "02"
        if month == "Martie":
            return "03"
        if month == "Aprilie":
            return "04"
        if month == "Mai":
            return "05"
        if month == "Iunie":
            return "06"
        if month == "Iulie":
            return "07"
        if month == "August":
            return "08"
        if month == "Septembrie":
            return "09"
        if month == "Octombrie":
            return "10"
        if month == "Noiembrie":
            return "11"
        if month == "Decembrie":
            return "12"

    def NumbToMonth(self, numb):
        if numb == "01":
            return "Ianuarie"
        if numb == "02":
            return "Februarie"
        if numb == "03":
            return "Martie"
        if numb == "04":
            return "Aprilie"
        if numb == "05":
            return "Mai"
        if numb == "06":
            return "Iunie"
        if numb == "07":
            return "Iulie"
        if numb == "08":
            return "August"
        if numb == "09":
            return "Septembrie"
        if numb == "10":
            return "Octombrie"
        if numb == "11":
            return "Noiembrie"
        if numb == "12":
            return "Decembrie"

    def chMonth(self):
        self.centrAlPop()

    #Incarc mongoDB
    def loadMongo(self):
        try:
            self.client = MongoClient("mongodb+srv://PdjtUn:123pdj34@red-nord.lhwnm.mongodb.net/test?"
                     "retryWrites=true&w=majority")
        except:
            self.msSecCall("Lipseste legatura cu internetul!")

    def loadMongoGen(self):
        self.loadMongo()
        self.dbGen = self.client.General

    def loadOficii(self):
        self.loadMongoGen()
        self.ofList = []
        self.ofListAbr = []
        # with open("Bundle/Oficii.json", "r") as f:
        #     oficiiDict = json.load(f)

        for i in self.dbGen.oficii.find():
            self.ofList.append(i["name"])
            self.ofListAbr.append(i["abr"])

    def loadMongoUN(self):
        self.loadMongo()
        self.db = self.client.djUN_test

    def loadAng(self):
        self.loadAngSec()

        self.uList = []
        for i in self.angajati.find():
            self.uList.append(i["name"])
        self.uList.sort()
        self.uCombo.clear()
        self.uCombo.addItems(self.uList)

    def loadAngModif(self):
        self.loadAngSec()

        uList = []
        for i in self.angajati.find():
            uList.append(i["name"])
        uList.sort()
        self.uModifCombo.clear()
        self.uModifCombo.addItems(uList)

    def loadAngModif_func_sec(self):
        self.loadSect()
        try:
            self.functiaCombo.setCurrentText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["position"]
            )
            self.tabelLine.setText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["nr_tabel"]
            )
            self.sectCombo.setCurrentText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["sector"]
            )
            self.grtsCombo.setCurrentText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["gr_ts"]
            )
            self.semnLine.setText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["semnatura_el"]
            )
            self.telefonLine_serv.setText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["telefon_serv"]
            )
            self.telefonLine_pers.setText(
                self.angajati.find_one({"name": self.uModifCombo.currentText()})["telefon_pers"]
            )
        except:
            TypeError

    def loadAngModif_func(self):
        self.loadAngModif_func_sec()

        try:
            if self.angajati.find_one({"name": self.uModifCombo.currentText()})["emitent"] == "DA!":
                self.emCheck.setChecked(True)
            else:
                self.emCheck.setChecked(False)

            if self.angajati.find_one({"name": self.uModifCombo.currentText()})["conducator"] == "DA!":
                self.condCheck.setChecked(True)
            else:
                self.condCheck.setChecked(False)

            if self.angajati.find_one({"name": self.uModifCombo.currentText()})["admitent"] == "DA!":
                self.admCheck.setChecked(True)
            else:
                self.admCheck.setChecked(False)

            if self.angajati.find_one({"name": self.uModifCombo.currentText()})["sef"] == "DA!":
                self.sefCheck.setChecked(True)
            else:
                self.sefCheck.setChecked(False)

            if self.angajati.find_one({"name": self.uModifCombo.currentText()})["supraveghetor"] == "DA!":
                self.suprCheck.setChecked(True)
            else:
                self.suprCheck.setChecked(False)

            if self.angajati.find_one({"name": self.uModifCombo.currentText()})["membru"] == "DA!":
                self.memCheck.setChecked(True)
            else:
                self.memCheck.setChecked(False)
        except:
            TypeError


    def loadAngSec(self):
        if self.ofCombo.currentText() == "Ungheni":
            self.angajati = self.db.angajati_un
        if self.ofCombo.currentText() == "Falesti":
            self.angajati = self.db.angajati_fl
        if self.ofCombo.currentText() == "Glodeni":
            self.angajati = self.db.angajati_gl
        if self.ofCombo.currentText() == "Riscani":
            self.angajati = self.db.angajati_rs

    #Abrevieri oficii
    def abrOficii(self):
        if self.ofCombo.currentText() == "Ungheni":
            self.ofVar = "UN"
        if self.ofCombo.currentText() == "Falesti":
            self.ofVar = "FL"
        if self.ofCombo.currentText() == "Glodeni":
            self.ofVar = "GL"
        if self.ofCombo.currentText() == "Riscani":
            self.ofVar = "RS"

    def abrOficiiSec(self, myOfice):
        if myOfice == "Ungheni":
            return "UN"
        if myOfice == "Falesti":
            return "FL"
        if myOfice == "Glodeni":
            return "GL"
        if myOfice == "Riscani":
            return "RS"

    # Abreviere inversa
    def abrOficii_invers(self, myOfice):
        if myOfice == "UN":
            return "Ungheni"
        if myOfice == "FL":
            return "Falesti"
        if myOfice == "GL":
            return "Glodeni"
        if myOfice == "RS":
            return "Riscani"

    def lucruLineCh(self):
        if self.lucrLine.currentText() == "Schimbare contor, consumator non-casnic" or \
                self.lucrLine.currentText() == "Deconectarea, reconectarea consumatorilor":
            self.masLine_1.setText("Deconectati sarcina. Lucrarile se efectuiaza folosind mijloace individuale de protectie;")
        else:
            self.masLine_1.setText("Fara deconectari, de respectat distanta minim admisa (Tabelul nr. 1 NEI)")

    #Incarc informatia PT excel autorizati, dispozitie, deconectare neplanificata
    def loadPt(self):
        if self.ofCombo.currentText() == "Ungheni" or self.ofCombo.currentText() == "UN":
            self.wbPt = load_workbook("Bundle/Ungheni/PT_Ungheni.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_un
            if self.deranjControl:
                self.sector = self.db.sectoare_un
                self.f10kv = open("Bundle/Ungheni/Lista_f_10kV.txt")
        if self.ofCombo.currentText() == "Falesti" or self.ofCombo.currentText() == "FL":
            self.wbPt = load_workbook("Bundle/Falesti/PT_Falesti.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_fl
            if self.deranjControl:
                self.sector = self.db.sectoare_fl
        if self.ofCombo.currentText() == "Glodeni" or self.ofCombo.currentText() == "GL":
            self.wbPt = load_workbook("Bundle/Glodeni/PT_Glodeni.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_gl
                # self.wbAng = load_workbook("Bundle/Glodeni/Angajati.xlsx")
            if self.deranjControl:
                # self.wbAng = load_workbook("Bundle/Glodeni/Angajati.xlsx")
                self.sector = self.db.sectoare_gl
        if self.ofCombo.currentText() == "Riscani" or self.ofCombo.currentText() == "RS":
            self.wbPt = load_workbook("Bundle/Riscani/PT_Riscani.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_rs
                # self.wbAng = load_workbook("Bundle/Riscani/Angajati.xlsx")
            if self.deranjControl:
                # self.wbAng = load_workbook("Bundle/Riscani/Angajati.xlsx")
                self.sector = self.db.sectoare_rs
        self.wsPt = self.wbPt.active
        # if not self.decControl:
        #     self.wsAng = self.wbAng.active

        self.ptList = []
        for i in range(2, self.wsPt.max_row + 1):
            self.ptList.append(self.wsPt.cell(row=i, column=1).value)
        myCompleter = QCompleter(self.ptList)
        myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
        self.ptLine.setCompleter(myCompleter)

        if not self.decControl:
            if not self.deranjControl:
                # Completez forma pentru emitent
                self.emList = []
                for i in self.angajati.find():
                    if i["gr_ts"] == "5":
                        self.emList.append(i["name"] + " " + "gr. " + i["gr_ts"])
                myCompleter = QCompleter(self.emList)
                myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
                self.emLine.setCompleter(myCompleter)
                for i in self.angajati.find({"name": self.nameFMenu}):
                    # print(i["position"])
                    if i["position"] != "Dispecer":
                        self.emLine.setText(self.nameFMenu + " " + "gr. " + i["gr_ts"])
                        self.emLine.setEnabled(False)
                        self.ofCombo.setEnabled(False)
                # Completez forma pentru sef
                self.sefList = []
                # print(self.angajati)
                for i in self.angajati.find():
                    self.sefList.append(i["name"] + " " + "gr. " + i["gr_ts"])
                myCompleter = QCompleter(self.sefList)
                myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
                self.sfLine.setCompleter(myCompleter)
                self.memEchLine.setCompleter(myCompleter)

            if self.deranjControl:
                # Completez forma pentru sef
                self.sefList = []
                for i in self.angajati.find():
                    self.sefList.append(i["name"])
                myCompleter = QCompleter(self.sefList)
                myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
                self.sfLine.setCompleter(myCompleter)
                self.sectorList = []
                #Incarc sectoarele de MongoDB
                for i in self.sector.find():
                    self.sectorList.append(i["name"])
                self.sectorList.sort()
                self.sectCombo.clear()
                self.sectCombo.addItems(self.sectorList)
                self.sectCombo.setCurrentText("Alege:")
                self.f10kvList = []
                for i in self.f10kv:
                    i = i.strip()
                    self.f10kvList.append(i)
                myCompleter = QCompleter(self.f10kvList)
                myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
                self.f10kvLine.setCompleter(myCompleter)

    def loadPtCorect(self):
        if self.ofCorectCombo.currentText() == "Ungheni" or self.ofCorectCombo.currentText() == "UN":
            self.wbPt = load_workbook("Bundle/Ungheni/PT_Ungheni.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_un
            if self.deranjControl:
                self.sector = self.db.sectoare_un
                self.f10kv = open("Bundle/Ungheni/Lista_f_10kV.txt")
        if self.ofCorectCombo.currentText() == "Falesti" or self.ofCorectCombo.currentText() == "FL":
            self.wbPt = load_workbook("Bundle/Falesti/PT_Falesti.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_fl
            if self.deranjControl:
                self.sector = self.db.sectoare_fl
        if self.ofCorectCombo.currentText() == "Glodeni" or self.ofCorectCombo.currentText() == "GL":
            self.wbPt = load_workbook("Bundle/Glodeni/PT_Glodeni.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_gl
                # self.wbAng = load_workbook("Bundle/Glodeni/Angajati.xlsx")
            if self.deranjControl:
                # self.wbAng = load_workbook("Bundle/Glodeni/Angajati.xlsx")
                self.sector = self.db.sectoare_gl
        if self.ofCorectCombo.currentText() == "Riscani" or self.ofCorectCombo.currentText() == "RS":
            self.wbPt = load_workbook("Bundle/Riscani/PT_Riscani.xlsx")
            if not self.decControl:
                self.angajati = self.db.angajati_rs
                # self.wbAng = load_workbook("Bundle/Riscani/Angajati.xlsx")
            if self.deranjControl:
                # self.wbAng = load_workbook("Bundle/Riscani/Angajati.xlsx")
                self.sector = self.db.sectoare_rs
        self.wsPt = self.wbPt.active
        # if not self.decControl:
        #     self.wsAng = self.wbAng.active

        self.ptList = []
        for i in range(2, self.wsPt.max_row + 1):
            self.ptList.append(self.wsPt.cell(row=i, column=1).value)
        myCompleter = QCompleter(self.ptList)
        myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
        self.ptLine.setCompleter(myCompleter)

        if not self.decControl:
            if not self.deranjControl:
                # Completez forma pentru emitent
                self.emList = []
                for i in self.angajati.find():
                    if i["gr_ts"] == "5":
                        self.emList.append(i["name"] + " " + "gr. " + i["gr_ts"])
                myCompleter = QCompleter(self.emList)
                myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
                self.emLine.setCompleter(myCompleter)
                for i in self.angajati.find({"name": self.nameFMenu}):
                    # print(i["position"])
                    if i["position"] != "Dispecer":
                        self.emLine.setText(self.nameFMenu + " " + "gr. " + i["gr_ts"])
                        self.emLine.setEnabled(False)
                        self.ofCombo.setEnabled(False)
                # Completez forma pentru sef
                self.sefList = []
                # print(self.angajati)
                for i in self.angajati.find():
                    self.sefList.append(i["name"] + " " + "gr. " + i["gr_ts"])
                myCompleter = QCompleter(self.sefList)
                myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
                self.sfLine.setCompleter(myCompleter)
                self.memEchLine.setCompleter(myCompleter)

            # if self.deranjControl:
            #     # Completez forma pentru sef
            #     self.sefList = []
            #     for i in self.angajati.find():
            #         self.sefList.append(i["name"])
            #     myCompleter = QCompleter(self.sefList)
            #     myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
            #     self.sfLine.setCompleter(myCompleter)
            #     self.sectorList = []
            #     #Incarc sectoarele de MongoDB
            #     for i in self.sector.find():
            #         self.sectorList.append(i["name"])
            #     self.sectorList.sort()
            #     self.sectCombo.clear()
            #     self.sectCombo.addItems(self.sectorList)
            #     self.sectCombo.setCurrentText("Alege:")
            #     self.f10kvList = []
            #     for i in self.f10kv:
            #         i = i.strip()
            #         self.f10kvList.append(i)
            #     myCompleter = QCompleter(self.f10kvList)
            #     myCompleter.setCaseSensitivity(Qt.CaseInsensitive)
            #     self.f10kvLine.setCompleter(myCompleter)

    def loadPtSec(self):
        if self.data.at[self.modRow, 0] == "UN":
            self.wbPt = load_workbook("Bundle/Ungheni/PT_Ungheni.xlsx")
        if self.data.at[self.modRow, 0] == "FL":
            self.wbPt = load_workbook("Bundle/Falesti/PT_Falesti.xlsx")
        if self.data.at[self.modRow, 0] == "GL":
            self.wbPt = load_workbook("Bundle/Glodeni/PT_Glodeni.xlsx")
        if self.data.at[self.modRow, 0] == "RS":
            self.wbPt = load_workbook("Bundle/Riscani/PT_Riscani.xlsx")
        self.wsPt = self.wbPt.active

    # Incarc informatia pentru sectoare
    def loadSectSec(self):
        if self.ofCombo.currentText() == "Ungheni":
            self.sector = self.db.sectoare_un
        if self.ofCombo.currentText() == "Falesti":
            self.sector = self.db.sectoare_fl
        if self.ofCombo.currentText() == "Glodeni":
            self.sector = self.db.sectoare_gl
        if self.ofCombo.currentText() == "Riscani":
            self.sector = self.db.sectoare_rs

    def loadSectSecAng(self):
        if self.ofAngCombo.currentText() == "Ungheni":
            self.sector = self.db.sectoare_un
        if self.ofAngCombo.currentText() == "Falesti":
            self.sector = self.db.sectoare_fl
        if self.ofAngCombo.currentText() == "Glodeni":
            self.sector = self.db.sectoare_gl
        if self.ofAngCombo.currentText() == "Riscani":
            self.sector = self.db.sectoare_rs

    def loadSect(self):
        self.loadSectSec()
        sectorList = []
        for i in self.sector.find():
            sectorList.append(i["name"])
        sectorList.sort()
        sectorList.insert(0, "Alege sector:")
        if self.ofCombo.currentText() != "Toate oficiile":
            self.sectCombo.clear()
            self.sectCombo.addItems(sectorList)
        else:
            self.sectCombo.clear()

    def loadSectAng(self):
        self.loadSectSecAng()
        if self.ofAngCombo.currentText() != "Toate oficiile":
            sectorList = []
            for i in self.sector.find():
                sectorList.append(i["name"])
            sectorList.sort()
            sectorList.insert(0, "Alege sector:")
            self.sectorCombo.clear()
            self.sectorCombo.addItems(sectorList)
        else:
            self.sectorCombo.clear()
            # self.sectorCombo.addItem("Toate")

    def angNouTrig(self):
        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Angajati - Nou primit')
        self.dialBox.setStyleSheet('background-color: #424242;')

        # Oficiul nume
        ofFrame = QFrame()
        ofFrame.setFrameShape(QFrame.StyledPanel)
        ofFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Oficiul:")
        # Incarc oficiile din MongoDB
        self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(100)
        self.ofCombo.addItems(self.ofList)
        # self.ofCombo.setStyleSheet("margin-right: 50%")
        self.ofCombo.currentTextChanged.connect(self.loadSect)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)


        # # Nume , la label pun ofLabel
        # numeFrame = QFrame()
        # numeFrame.setFrameShape(QFrame.StyledPanel)
        # numeFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Nume (ex. Carp A.):")
        ofLabel.setStyleSheet("margin-left: 20%")
        self.numeLine = QLineEdit()
        self.numeLine.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.numeLine.setFixedWidth(100)

        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.numeLine)
        ofFrame.setLayout(hbox)
        # numeFrame.setLayout(hbox)

        # Functia nr_tabel
        functiaFrame = QFrame()
        functiaFrame.setFrameShape(QFrame.StyledPanel)
        functiaFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Functia:")
        myList = []
        for i in self.dbGen.position.find():
            myList.append(i["name"])
        myList.sort()
        myList.insert(0, "Alege functia:")
        self.functiaCombo = QComboBox()
        self.functiaCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.functiaCombo.setFixedWidth(100)
        self.functiaCombo.addItems(myList)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.functiaCombo)
        # functiaFrame.setLayout(hbox)

        # # Nr. tabel
        # tabelFrame = QFrame()
        # tabelFrame.setFrameShape(QFrame.StyledPanel)
        # tabelFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Nr. tabel:")
        ofLabel.setStyleSheet("margin-left: 70%")
        self.tabelLine = QLineEdit()
        self.tabelLine.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.tabelLine.setFixedWidth(100)
        rxTabel = QRegExp("\d\d\d\d")
        myValidator = QRegExpValidator(rxTabel)
        self.tabelLine.setValidator(myValidator)

        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.tabelLine)
        functiaFrame.setLayout(hbox)

        # Sector, gr_ts
        sectorFrame = QFrame()
        sectorFrame.setFrameShape(QFrame.StyledPanel)
        sectorFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Sector:")
        myList = []
        for i in self.db.sectoare_un.find():
            myList.append(i["name"])
        myList.sort()
        myList.insert(0, "Alege sector:")
        self.sectCombo = QComboBox()
        self.sectCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.sectCombo.setFixedWidth(100)
        self.sectCombo.addItems(myList)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.sectCombo)
        # sectorFrame.setLayout(hbox)


        # Grupa TS
        # grtsFrame = QFrame()
        # grtsFrame.setFrameShape(QFrame.StyledPanel)
        # grtsFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Grupa TS:")
        myList = ["1", "2", "3", "4", "5"]
        ofLabel.setStyleSheet("margin-left: 70%")
        self.grtsCombo = QComboBox()
        self.grtsCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.grtsCombo.setFixedWidth(100)
        self.grtsCombo.addItems(myList)

        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.grtsCombo)
        sectorFrame.setLayout(hbox)

        # Telefon
        telefonFrame = QFrame()
        telefonFrame.setFrameShape(QFrame.StyledPanel)
        telefonFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Tel./serv.:")
        self.telefonLine_serv = QLineEdit()
        self.telefonLine_serv.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.telefonLine_serv.setFixedWidth(100)
        # rxTabel = QRegExp("\d\d\d\d\d\d\d\d\d\d\d\d")
        # myValidator = QRegExpValidator(rxTabel)
        # self.telefonLine.setValidator(myValidator)

        ofLabel_pers = QLabel("Tel./pers.:")
        ofLabel_pers.setStyleSheet("margin-left: 65%")
        self.telefonLine_pers = QLineEdit()
        self.telefonLine_pers.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.telefonLine_pers.setFixedWidth(100)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.telefonLine_serv)

        # hbox_pers = QHBoxLayout()
        hbox.addWidget(ofLabel_pers)
        hbox.addWidget(self.telefonLine_pers)

        # vbox = QVBoxLayout()
        # vbox.addLayout(hbox)
        # vbox.addLayout(hbox_pers)
        telefonFrame.setLayout(hbox)

        #Semnatura
        semnFrame = QFrame()
        semnFrame.setFrameShape(QFrame.StyledPanel)
        semnFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Semnatura:")
        self.semnLine = QLineEdit()
        self.semnLine.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.semnLine.setValidator(myDtValidator)
        setRegButton = QPushButton()
        setRegButton.setIcon(QIcon("folder.ico"))
        setRegButton.setIconSize(QSize(20, 20))
        setRegButton.clicked.connect(self.semnLoad)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.semnLine)
        hbox.addWidget(setRegButton)
        semnFrame.setLayout(hbox)


        # Drepturi
        dreptFrame = QFrame()
        dreptFrame.setFrameShape(QFrame.StyledPanel)
        dreptFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        # emLabel = QLabel("Emitent:")
        self.emCheck = QCheckBox("Emitent AL, DS")
        self.emCheck.setStyleSheet("padding-right: 37%")
        self.condCheck = QCheckBox("Conducator de lucrari")
        self.admCheck = QCheckBox("Admitent")
        self.sefCheck = QCheckBox("Sef de lucrari")
        self.suprCheck = QCheckBox("Supraveghetor")
        self.memCheck = QCheckBox("Membru echipei")

        hbox_1 = QHBoxLayout()
        hbox_1.addWidget(self.emCheck)
        hbox_1.addWidget(self.condCheck)

        hbox_2 = QHBoxLayout()
        hbox_2.addWidget(self.admCheck)
        hbox_2.addWidget(self.sefCheck)

        hbox_3 = QHBoxLayout()
        hbox_3.addWidget(self.suprCheck)
        hbox_3.addWidget(self.memCheck)


        vbox = QVBoxLayout()
        vbox.addLayout(hbox_1)
        vbox.addLayout(hbox_2)
        vbox.addLayout(hbox_3)

        dreptFrame.setLayout(vbox)

        # Buttons Section (butoanele "Ok", "Cancel"
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okAng)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        vbox = QVBoxLayout()
        vbox.addWidget(ofFrame)
        # vbox.addWidget(numeFrame)
        vbox.addWidget(functiaFrame)
        # vbox.addWidget(tabelFrame)
        vbox.addWidget(sectorFrame)
        # vbox.addWidget(grtsFrame)
        vbox.addWidget(telefonFrame)
        vbox.addWidget(semnFrame)
        vbox.addWidget(dreptFrame)
        vbox.addWidget(btFrame)
        self.dialBox.setLayout(vbox)

        self.dialBox.exec()

    # def dialModifClose(self):
    #     if self.cautAngControl:
    #         self.dialBox.close()
    #         self.cautAngControl = False
    #     else:
    #         self.dialBox.close()

    def angCautTrig(self):
        self.cautAngControl = True
        # self.angModifTrig()

        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))

        self.dialBox.setWindowTitle('Angajati - Cauta angajat')
        self.dialBox.setStyleSheet('background-color: rgb(65, 84, 71);')
            # self.cautAngControl = False

        # Oficiul
        ofFrame = QFrame()
        ofFrame.setFrameShape(QFrame.StyledPanel)
        ofFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Oficiul:")
        # Incarc oficiile din MongoDB
        self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(100)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.setEditable(True)
        self.ofCombo.currentTextChanged.connect(self.loadAngModif)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        # ofFrame.setLayout(hbox)

        # Nume , la label pun ofLabel
        # numeFrame = QFrame()
        # numeFrame.setFrameShape(QFrame.StyledPanel)
        # numeFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Nume:")
        ofLabel.setStyleSheet("margin-left: 22%")
        self.uModifCombo = QComboBox()
        self.uModifCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.uModifCombo.setFixedWidth(100)
        self.uModifCombo.setEditable(True)
        self.loadAngModif()
        # self.uModifCombo.currentTextChanged.connect(self.loadAngModif_func)

        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.uModifCombo)
        ofFrame.setLayout(hbox)

        # Buttons Section (butoanele "Ok", "Cancel")
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okAngCaut)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        # # self.loadSectSec()
        # self.loadAngModif_func()
        vbox = QVBoxLayout()
        vbox.addWidget(ofFrame)
        # vbox.addWidget(numeFrame)
        # vbox.addWidget(functiaFrame)
        # vbox.addWidget(sectorFrame)
        # vbox.addWidget(grtsFrame)
        # vbox.addWidget(telefonFrame)
        # vbox.addWidget(semnFrame)
        # vbox.addWidget(dreptFrame)
        vbox.addWidget(btFrame)
        self.dialBox.setLayout(vbox)

        if self.contextAng:
            self.ofCombo.setCurrentText(self.abrOficii_invers(self.data.at[self.modRow, "oficiul"]))
            self.uModifCombo.setCurrentText(self.data.at[self.modRow, "name"])
            self.contextAng = False

        self.dialBox.exec()

    def angModifTrig(self):
        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Angajati - Modificare stare')
        self.dialBox.setStyleSheet('background-color: #424242;')

        # Oficiul
        ofFrame = QFrame()
        ofFrame.setFrameShape(QFrame.StyledPanel)
        ofFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Oficiul:")
        # Incarc oficiile din MongoDB
        self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(100)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.setEditable(True)
        self.ofCombo.currentTextChanged.connect(self.loadAngModif)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        # ofFrame.setLayout(hbox)

        # Nume , la label pun ofLabel
        # numeFrame = QFrame()
        # numeFrame.setFrameShape(QFrame.StyledPanel)
        # numeFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")


        ofLabel = QLabel("Nume:")
        ofLabel.setStyleSheet("margin-left: 22%")
        self.uModifCombo = QComboBox()
        self.uModifCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.uModifCombo.setFixedWidth(100)
        self.uModifCombo.setEditable(True)
        self.loadAngModif()
        self.uModifCombo.currentTextChanged.connect(self.loadAngModif_func)


        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.uModifCombo)
        ofFrame.setLayout(hbox)

        # Functia nr_tabel
        functiaFrame = QFrame()
        functiaFrame.setFrameShape(QFrame.StyledPanel)
        functiaFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Functia:")
        myList = []
        for i in self.dbGen.position.find():
            myList.append(i["name"])
        myList.sort()
        self.functiaCombo = QComboBox()
        self.functiaCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.functiaCombo.setFixedWidth(100)
        self.functiaCombo.addItems(myList)
        self.functiaCombo.setEditable(True)
        # self.functiaCombo.currentTextChanged.connect(self.loadAngModif_sector)


        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.functiaCombo)
        # functiaFrame.setLayout(hbox)

        ofLabel = QLabel("Nr. tabel:")
        ofLabel.setStyleSheet("margin-left: 8%")
        self.tabelLine = QLineEdit()
        self.tabelLine.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.tabelLine.setFixedWidth(100)
        rxTabel = QRegExp("\d\d\d\d")
        myValidator = QRegExpValidator(rxTabel)
        self.tabelLine.setValidator(myValidator)

        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.tabelLine)
        functiaFrame.setLayout(hbox)

        # Sector, gr_ts
        sectorFrame = QFrame()
        sectorFrame.setFrameShape(QFrame.StyledPanel)
        sectorFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Sector:")
        myList = []
        for i in self.db.sectoare_un.find():
            myList.append(i["name"])
        myList.sort()
        self.sectCombo = QComboBox()
        self.sectCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.sectCombo.setFixedWidth(100)
        self.sectCombo.addItems(myList)
        self.sectCombo.setEditable(True)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.sectCombo)
        # sectorFrame.setLayout(hbox)

        # Grupa TS
        # grtsFrame = QFrame()
        # grtsFrame.setFrameShape(QFrame.StyledPanel)
        # grtsFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Grupa TS:")
        myList = ["1", "2", "3", "4", "5"]
        ofLabel.setStyleSheet("margin-left: 7%")
        self.grtsCombo = QComboBox()
        self.grtsCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.grtsCombo.setFixedWidth(100)
        self.grtsCombo.addItems(myList)
        self.grtsCombo.setEditable(True)

        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.grtsCombo)
        sectorFrame.setLayout(hbox)

        # Telefon
        telefonFrame = QFrame()
        telefonFrame.setFrameShape(QFrame.StyledPanel)
        telefonFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Tel./serv.:")
        self.telefonLine_serv = QLineEdit()
        self.telefonLine_serv.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.telefonLine_serv.setFixedWidth(100)
        # rxTabel = QRegExp("\d\d\d\d\d\d\d\d\d\d\d\d")
        # myValidator = QRegExpValidator(rxTabel)
        # self.telefonLine.setValidator(myValidator)

        ofLabel_pers = QLabel("Tel./pers.:")
        ofLabel_pers.setStyleSheet("margin-left: 3%")
        self.telefonLine_pers = QLineEdit()
        self.telefonLine_pers.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        self.telefonLine_pers.setFixedWidth(100)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.telefonLine_serv)

        hbox.addWidget(ofLabel_pers)
        hbox.addWidget(self.telefonLine_pers)

        # vbox = QVBoxLayout()
        # vbox.addLayout(hbox)
        # vbox.addLayout(hbox_pers)
        telefonFrame.setLayout(hbox)

        # Semnatura
        semnFrame = QFrame()
        semnFrame.setFrameShape(QFrame.StyledPanel)
        semnFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Semnatura:")
        self.semnLine = QLineEdit()
        self.semnLine.setStyleSheet('background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.semnLine.setValidator(myDtValidator)
        setRegButton = QPushButton()
        setRegButton.setIcon(QIcon("folder.ico"))
        setRegButton.setIconSize(QSize(20, 20))
        setRegButton.clicked.connect(self.semnLoad)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.semnLine)
        hbox.addWidget(setRegButton)
        semnFrame.setLayout(hbox)

        # Drepturi
        dreptFrame = QFrame()
        dreptFrame.setFrameShape(QFrame.StyledPanel)
        dreptFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        # emLabel = QLabel("Emitent:")
        self.emCheck = QCheckBox("Emitent AL, DS")
        self.emCheck.setStyleSheet("padding-right: 37%")
        self.condCheck = QCheckBox("Conducator de lucrari")
        self.admCheck = QCheckBox("Admitent")
        self.sefCheck = QCheckBox("Sef de lucrari")
        self.suprCheck = QCheckBox("Supraveghetor")
        self.memCheck = QCheckBox("Membru echipei")

        hbox_1 = QHBoxLayout()
        hbox_1.addWidget(self.emCheck)
        hbox_1.addWidget(self.condCheck)

        hbox_2 = QHBoxLayout()
        hbox_2.addWidget(self.admCheck)
        hbox_2.addWidget(self.sefCheck)

        hbox_3 = QHBoxLayout()
        hbox_3.addWidget(self.suprCheck)
        hbox_3.addWidget(self.memCheck)

        vbox = QVBoxLayout()
        vbox.addLayout(hbox_1)
        vbox.addLayout(hbox_2)
        vbox.addLayout(hbox_3)

        dreptFrame.setLayout(vbox)

        # Buttons Section (butoanele "Ok", "Cancel")
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okAngModif)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        # self.loadSectSec()
        self.loadAngModif_func()
        vbox = QVBoxLayout()
        vbox.addWidget(ofFrame)
        # vbox.addWidget(numeFrame)
        vbox.addWidget(functiaFrame)
        vbox.addWidget(sectorFrame)
        # vbox.addWidget(grtsFrame)
        vbox.addWidget(telefonFrame)
        vbox.addWidget(semnFrame)
        vbox.addWidget(dreptFrame)
        vbox.addWidget(btFrame)
        self.dialBox.setLayout(vbox)

        if self.contextAng:
            self.ofCombo.setCurrentText(self.abrOficii_invers(self.data.at[self.modRow, "oficiul"]))
            self.uModifCombo.setCurrentText(self.data.at[self.modRow, "name"])
            self.contextAng = False

        self.dialBox.exec()

    # def cautSectTrig(self):
    #     self.cautSecAngControl = True
    #     # Dialog Window Create
    #     self.dialBox = QDialog()
    #     self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
    #     self.dialBox.setWindowIcon(QIcon('ataman_logo'))
    #     self.dialBox.setWindowTitle('Angajati - Cauta dupa sector')
    #     self.dialBox.setStyleSheet('background-color: rgb(65, 84, 71);')
    #
    #     # Oficiul
    #     ofFrame = QFrame()
    #     ofFrame.setFrameShape(QFrame.StyledPanel)
    #     ofFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")
    #
    #     ofLabel = QLabel("Oficiul:")
    #     # Incarc oficiile din MongoDB
    #     self.loadOficii()
    #     self.ofCombo = QComboBox()
    #     self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
    #     self.ofCombo.setFixedWidth(100)
    #     self.ofCombo.addItems(self.ofList)
    #     self.ofCombo.setEditable(True)
    #     self.ofCombo.currentTextChanged.connect(self.loadSect)
    #
    #     scLabel = QLabel("Sector:")
    #     scLabel.setStyleSheet("margin-left: 22%")
    #     myList = []
    #     for i in self.db.sectoare_un.find():
    #         myList.append(i["name"])
    #     myList.sort()
    #     self.sectCombo = QComboBox()
    #     self.sectCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
    #     self.sectCombo.setFixedWidth(100)
    #     self.sectCombo.addItems(myList)
    #     self.sectCombo.setEditable(True)
    #
    #     hbox = QHBoxLayout()
    #     hbox.addWidget(ofLabel)
    #     hbox.addWidget(self.ofCombo)
    #     hbox.addWidget(scLabel)
    #     hbox.addWidget(self.sectCombo)
    #     ofFrame.setLayout(hbox)
    #
    #     # Buttons Section (butoanele "Ok", "Cancel")
    #     btFrame = QFrame()
    #     btFrame.setFrameShape(QFrame.StyledPanel)
    #
    #     btOk = QPushButton('Ok')
    #     btOk.setStyleSheet('color: #e3e3e3')
    #     btOk.clicked.connect(self.okAngCaut)
    #
    #     btCancel = QPushButton('Cancel')
    #     btCancel.setStyleSheet('color: #e3e3e3')
    #     btCancel.clicked.connect(self.dialBox.close)
    #
    #     hbox = QHBoxLayout()
    #     hbox.addWidget(btOk)
    #     hbox.addWidget(btCancel)
    #     btFrame.setLayout(hbox)
    #
    #     # self.loadSectSec()
    #     self.loadAngModif_func()
    #     vbox = QVBoxLayout()
    #     vbox.addWidget(ofFrame)
    #     vbox.addWidget(btFrame)
    #     self.dialBox.setLayout(vbox)
    #
    #
    #     if self.contextAng:
    #         self.ofCombo.setCurrentText(self.abrOficii_invers(self.data.at[self.modRow, "oficiul"]))
    #         # self.uModifCombo.setCurrentText(self.data.at[self.modRow, "name"])
    #         self.contextAng = False
    #
    #     self.dialBox.exec()

    def angDelTrig(self):
        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Angajati - Sterge angajat')
        self.dialBox.setStyleSheet('background-color: #424242;')

        # Oficiul
        ofFrame = QFrame()
        ofFrame.setFrameShape(QFrame.StyledPanel)
        ofFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

        ofLabel = QLabel("Oficiul:")
        # Incarc oficiile din MongoDB
        self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(100)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.currentTextChanged.connect(self.loadAngModif)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        # ofFrame.setLayout(hbox)

        # Nume , la label pun ofLabel
        # numeFrame = QFrame()
        # numeFrame.setFrameShape(QFrame.StyledPanel)
        # numeFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")


        ofLabel = QLabel("Nume:")
        ofLabel.setStyleSheet("margin-left: 22%")
        self.uModifCombo = QComboBox()
        self.uModifCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.uModifCombo.setFixedWidth(100)
        self.uModifCombo.setEditable(True)
        self.loadAngModif()
        # self.uModifCombo.currentTextChanged.connect(self.loadAngModif)


        # hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.uModifCombo)
        ofFrame.setLayout(hbox)

        # Buttons Section (butoanele "Ok", "Cancel"
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okAngDel)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        # self.loadAngModif_func_sec()
        vbox = QVBoxLayout()
        vbox.addWidget(ofFrame)
        # vbox.addWidget(numeFrame)
        vbox.addWidget(btFrame)
        self.dialBox.setLayout(vbox)

        if self.contextAng:
            self.ofCombo.setCurrentText(self.abrOficii_invers(self.data.at[self.modRow, "oficiul"]))
            self.uModifCombo.setCurrentText(self.data.at[self.modRow, "name"])
            self.contextAng = False

        self.dialBox.exec()


    def intTrig(self):
        if not self.passControl:
        #Formez dereastra dialogului
            self.dialInt = QDialog()
            # self.dialInt.setWindowFlags(Qt.WindowCloseButtonHint)
            self.dialInt.setWindowFlags(Qt.FramelessWindowHint)
            self.dialInt.setWindowIcon(QIcon('ataman_logo'))
            self.dialInt.setWindowTitle('Log In:')
            self.dialInt.setStyleSheet('background-color: #424242;')

        #Oficiul
            ofFrame = QFrame()
            ofFrame.setFrameShape(QFrame.StyledPanel)
            ofFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

            ofLabel = QLabel("Oficiul:")
        #Incarc oficiile din MongoDB
            self.loadOficii()
            self.ofCombo = QComboBox()
            self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
            self.ofCombo.setFixedWidth(100)
            self.ofCombo.addItems(self.ofList)
            self.ofCombo.currentTextChanged.connect(self.loadAng)

            self.loadMongoUN()


            uLabel = QLabel("User:")


            self.uCombo = QComboBox()
            self.uCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
            self.uCombo.setFixedWidth(100)
            self.uCombo.setEditable(True)
            # self.uCombo.setFont(QFont("Calibri", QFont.Bold))
            self.loadAng()
            self.uCombo.setCurrentText("Alege:")
            self.uCombo.currentTextChanged.connect(self.menuName)

            hbox = QHBoxLayout()
            hbox.addWidget(ofLabel)
            hbox.addWidget(self.ofCombo)
            hbox.addWidget(uLabel)
            hbox.addWidget(self.uCombo)
            ofFrame.setLayout(hbox)

        #Auth Frame button
            self.authBtFrame = QFrame()
            self.authBtFrame.setFrameShape(QFrame.StyledPanel)

            self.btAuth = QPushButton("Generati cod de autentificare")
            self.btAuth.setStyleSheet('color: #e3e3e3')
            self.btAuth.clicked.connect(self.authFunc)

            hbox = QHBoxLayout()
            hbox.addWidget(self.btAuth)
            self.authBtFrame.setLayout(hbox)

        # Auth Frame
            self.authFrame = QFrame()
            self.authFrame.setFrameShape(QFrame.StyledPanel)
            self.authFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")

            authLabel = QLabel("Introduceti codul de autentificare:")
            self.authLine = QLineEdit()
            self.authLine.setStyleSheet('background-color: #ffffff; color: #050505')
            self.authLine.setFixedWidth(100)

            hbox = QHBoxLayout()
            hbox.addWidget(authLabel)
            hbox.addWidget(self.authLine)
            self.authFrame.setLayout(hbox)

        # User Frame
            self.psFrame = QFrame()
            self.psFrame.setFrameShape(QFrame.StyledPanel)
            self.psFrame.setStyleSheet("background-color: #535453; color: #e3e3e3")

            self.tableCheck = QCheckBox("Alegeti obtiunea numarului de tabel")
            self.tableCheck.stateChanged.connect(self.tbCheckFunc)

            psLabel = QLabel("Introduceti numarul de tabel:")
            self.psText = PasswordEdit()
            self.psText.setStyleSheet('background-color: #ffffff; color: #050505')
            self.psText.setFixedWidth(100)
            self.psText.setEnabled(False)
        # self.psText.setEchoMode(QLineEdit.Password)

            hbox = QHBoxLayout()
            vbox = QVBoxLayout()
            vbox.addWidget(self.tableCheck)
            hbox.addWidget(psLabel)
            hbox.addWidget(self.psText)
            vbox.addLayout(hbox)
            self.psFrame.setLayout(vbox)

        # Buttons Section (butoanele "Ok", "Cancel"
            btFrame = QFrame()
            btFrame.setFrameShape(QFrame.StyledPanel)

            btOk = QPushButton('Ok')
            btOk.setStyleSheet('color: #e3e3e3')
            btOk.clicked.connect(self.okPass)

            btCancel = QPushButton('Cancel')
            btCancel.setStyleSheet('color: #e3e3e3')
            btCancel.clicked.connect(self.cancelPass)

            hbox = QHBoxLayout()
            hbox.addWidget(btOk)
            hbox.addWidget(btCancel)
            btFrame.setLayout(hbox)

            vbox = QVBoxLayout()
            vbox.addWidget(ofFrame)
            vbox.addWidget(self.authBtFrame)
            vbox.addWidget(self.authFrame)
            vbox.addWidget(self.psFrame)
            vbox.addWidget(btFrame)
            self.dialInt.setLayout(vbox)

            self.dialInt.exec()
            #Pun aceasta instructiune ca sa functioneze contextMenu pentru toate ferestrele, dar nu
            # doar pentru prima care are self.tabel
            self.table = QTableView()
            self.tableDeranj = QTableView()

    def tbCheckFunc(self):
        if self.tableCheck.isChecked() == True:
            self.psFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")
            self.psText.setEnabled(True)
            self.authFrame.setStyleSheet("background-color: #535453; color: #e3e3e3")
            self.authLine.setEnabled(False)
            self.authBtFrame.setStyleSheet('background-color: #535453;')
            self.btAuth.setEnabled(False)
            self.btAuth.setStyleSheet('background-color: #535453')
        else:
            self.psFrame.setStyleSheet("background-color: #535453; color: #e3e3e3")
            self.psText.setEnabled(False)
            self.authFrame.setStyleSheet("background-color: #315240; color: #e3e3e3")
            self.authLine.setEnabled(True)
            self.authBtFrame.setStyleSheet('background-color: #424242;')
            self.btAuth.setEnabled(True)
            self.btAuth.setStyleSheet('background-color: #424242; color: #e3e3e3')

    def menuName(self):
        try:
            self.nameFMenu = self.uCombo.currentText()
            self.nameOf = self.ofCombo.currentText()
            self.namePosition = self.angajati.find_one({"name": self.nameFMenu})["position"]
            self.nameAngajati = self.angajati
        except:
            TypeError

    def setTrig(self):
        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Setari')
        self.dialBox.setStyleSheet('background-color: #424242;')

        setGenFrame = QFrame()
        setGenFrame.setFrameShape(QFrame.StyledPanel)

        #Destinatia registru autorizatii
        setRegFrame = QFrame()
        setRegFrame.setFrameShape(QFrame.StyledPanel)
        setRegFrame.setStyleSheet('background-color: #5c3838;')

        setRegLabel = QLabel("Introduceti destinatia registru AUTORIZATII (excel):")
        setRegLabel.setStyleSheet('color: #e3e3e3;')
        self.setRegLine = QLineEdit()
        self.setRegLine.setFixedWidth(250)
        self.setRegLine.setStyleSheet('background-color: #e3e3e3;')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.setRegLine.setValidator(myDtValidator)

        setRegButton = QPushButton()
        setRegButton.setIcon(QIcon("folder.ico"))
        setRegButton.setIconSize(QSize(20, 20))
        setRegButton.clicked.connect(self.destRegAl)

        hbox = QHBoxLayout()
        hbox.addWidget(setRegLabel)
        hbox.addWidget(self.setRegLine)
        hbox.addWidget(setRegButton)
        setRegFrame.setLayout(hbox)

        # Destinatia SAIDI
        setSaidiFrame = QFrame()
        setSaidiFrame.setFrameShape(QFrame.StyledPanel)
        setSaidiFrame.setStyleSheet('background-color: #5c3838;')

        setSaidiLabel = QLabel("Introduceti destinatia SAIDI (excel):")
        setSaidiLabel.setStyleSheet('color: #e3e3e3; margin-right: 70%')
        self.setSaidiLine = QLineEdit()
        self.setSaidiLine.setStyleSheet('background-color: #e3e3e3;')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.setSaidiLine.setValidator(myDtValidator)

        setSaidiButton = QPushButton()
        setSaidiButton.setIcon(QIcon("folder.ico"))
        setSaidiButton.setIconSize(QSize(20, 20))
        setSaidiButton.clicked.connect(self.destSaidi)

        hbox = QHBoxLayout()
        hbox.addWidget(setSaidiLabel)
        hbox.addWidget(self.setSaidiLine)
        hbox.addWidget(setSaidiButton)
        setSaidiFrame.setLayout(hbox)

        # Destinatia RAPORT PDJT
        setRapFrame = QFrame()
        setRapFrame.setFrameShape(QFrame.StyledPanel)
        setRapFrame.setStyleSheet('background-color: #5c3838;')

        setRapLabel = QLabel("Introduceti destinatia RAPORT PDJT (excel):")
        setRapLabel.setStyleSheet('color: #e3e3e3; margin-right: 30%')
        self.setRapLine = QLineEdit()
        self.setRapLine.setStyleSheet('background-color: #e3e3e3;')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.setRapLine.setValidator(myDtValidator)

        setRapButton = QPushButton()
        setRapButton.setIcon(QIcon("folder.ico"))
        setRapButton.setIconSize(QSize(20, 20))
        setRapButton.clicked.connect(self.destRap)

        hbox = QHBoxLayout()
        hbox.addWidget(setRapLabel)
        hbox.addWidget(self.setRapLine)
        hbox.addWidget(setRapButton)
        setRapFrame.setLayout(hbox)

        # Destinati Dispozitiilor
        setDsFrame = QFrame()
        setDsFrame.setFrameShape(QFrame.StyledPanel)
        setDsFrame.setStyleSheet('background-color: #5c3838;')

        setDsLabel = QLabel("Introduceti destinatia DISPOZITIILOR:")
        setDsLabel.setStyleSheet('color: #e3e3e3; margin-right: 57%')
        self.setDsLine = QLineEdit()
        self.setDsLine.setStyleSheet('background-color: #e3e3e3;')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.setDsLine.setValidator(myDtValidator)

        setDsButton = QPushButton()
        setDsButton.setIcon(QIcon("folder.ico"))
        setDsButton.setIconSize(QSize(20, 20))
        setDsButton.clicked.connect(self.destDs)

        hbox = QHBoxLayout()
        hbox.addWidget(setDsLabel)
        hbox.addWidget(self.setDsLine)
        hbox.addWidget(setDsButton)
        setDsFrame.setLayout(hbox)

        # Destinatia Analiza
        setAnFrame = QFrame()
        setAnFrame.setFrameShape(QFrame.StyledPanel)
        setAnFrame.setStyleSheet('background-color: #5c3838;')

        setAnLabel = QLabel("Introduceti destinatia ANALIZA ANUALA (excel):")
        setAnLabel.setStyleSheet('color: #e3e3e3; margin-right: 10%')
        self.setAnLine = QLineEdit()
        self.setAnLine.setStyleSheet('background-color: #e3e3e3;')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.setAnLine.setValidator(myDtValidator)

        setAnButton = QPushButton()
        setAnButton.setIcon(QIcon("folder.ico"))
        setAnButton.setIconSize(QSize(20, 20))
        setAnButton.clicked.connect(self.destAn)

        hbox = QHBoxLayout()
        hbox.addWidget(setAnLabel)
        hbox.addWidget(self.setAnLine)
        hbox.addWidget(setAnButton)
        setAnFrame.setLayout(hbox)

        #Butoanele
        setBtFrame = QFrame()
        # setBtFrame.setFrameShape(QFrame.StyledPanel)

        setBtOk = QPushButton('Ok')
        setBtOk.setStyleSheet('color: #e3e3e3')
        # setBtOk.setFixedWidth(100)
        setBtOk.clicked.connect(self.destOk)
        btReset = QPushButton("Reset")
        btReset.setStyleSheet('color: #e3e3e3')
        # btReset.setFixedWidth(100)
        btReset.clicked.connect(self.resetTrig)
        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        # btCancel.setFixedWidth(100)
        btCancel.clicked.connect(self.dialBox.close)

        # setEmptyLb = QLabel()

        hbox = QHBoxLayout()
        # hbox.addWidget(setEmptyLb)
        hbox.addWidget(setBtOk)
        hbox.addWidget(btReset)
        hbox.addWidget(btCancel)
        setBtFrame.setLayout(hbox)

        vbox = QVBoxLayout()
        vbox.addWidget(setRegFrame)
        vbox.addWidget(setSaidiFrame)
        vbox.addWidget(setRapFrame)
        vbox.addWidget(setDsFrame)
        vbox.addWidget(setAnFrame)
        vbox.addWidget(setBtFrame)
        setGenFrame.setLayout(vbox)

        hbox = QHBoxLayout()
        hbox.addWidget(setGenFrame)
        self.dialBox.setLayout(hbox)

        self.destLoad()
        self.setRegLine.setText(self.wsDest.cell(row=1, column=2).value)
        self.setSaidiLine.setText(self.wsDest.cell(row=2, column=2).value)
        self.setRapLine.setText(self.wsDest.cell(row=3, column=2).value)
        self.setDsLine.setText(self.wsDest.cell(row=4, column=2).value)
        self.setAnLine.setText(self.wsDest.cell(row=5, column=2).value)

        self.dialBox.exec()

    def dtButtonFunc(self):
        if self.formCheck.isChecked() == True:
            self.dialBoxDt = QDialog()
            # self.dialBoxDt.setWindowFlags(Qt.WindowCloseButtonHint)
            self.dialBoxDt.setWindowIcon(QIcon('ataman_logo'))
            self.dialBoxDt.setWindowTitle('Calendar')
            # self.dialBoxDt.setStyleSheet('background-color: #424242;')

            self.myDate = QCalendarWidget()
            # self.myDate.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)")
            self.myDate.clicked.connect(self.dateChange)

            hbox = QHBoxLayout()
            hbox.addWidget(self.myDate)
            self.dialBoxDt.setLayout(hbox)

            self.dialBoxDt.exec()

    def dateChange(self):
        myDate = self.myDate.selectedDate().toString(Qt.SystemLocaleShortDate)
        self.dtLine.setText(myDate)
        self.dialBoxDt.close()

    def newLink(self):
        dlg = QFileDialog()
        fileName = dlg.getOpenFileName()
        self.linkLine.setText(fileName[0])

    def okLink(self):
        self.data.at[self.modRow, 18] = self.linkLine.text()
        if self.data.at[self.modRow, 1] != "":
            self.reg_al.update_one({
                "nr_ds": self.data.at[self.modRow, 1]
            }, {
                "$set": {
                    "link": self.linkLine.text()
                }
            })
        elif self.data.at[self.modRow, 2] != "":
            self.reg_al.update_one({
                "nr_al": self.data.at[self.modRow, 2]
            }, {
                "$set": {
                    "link": self.linkLine.text()
                }
            })
        # self.centrAlPop()
        self.dialBox.close()

    def destRegAl(self):
        dlg = QFileDialog()
        myDirectory = dlg.getExistingDirectory()
        self.setRegLine.setText(myDirectory)
    def semnLoad(self):
        dlg = QFileDialog()
        smnaturaLink = dlg.getOpenFileName()
        self.semnLine.setText(smnaturaLink[0])
    def destSaidi(self):
        dlg = QFileDialog()
        myDirectory = dlg.getExistingDirectory()
        self.setSaidiLine.setText(myDirectory)
    def destRap(self):
        dlg = QFileDialog()
        myDirectory = dlg.getExistingDirectory()
        self.setRapLine.setText(myDirectory)
    def destDs(self):
        dlg = QFileDialog()
        myDirectory = dlg.getExistingDirectory()
        self.setDsLine.setText(myDirectory)
    def destAn(self):
        dlg = QFileDialog()
        myDirectory = dlg.getExistingDirectory()
        self.setAnLine.setText(myDirectory)

    def destOk(self):
        self.destLoad()
        self.wsDest.cell(row=1, column=2).value = self.setRegLine.text()
        self.wsDest.cell(row=2, column=2).value = self.setSaidiLine.text()
        self.wsDest.cell(row=3, column=2).value = self.setRapLine.text()
        self.wsDest.cell(row=4, column=2).value = self.setDsLine.text()
        self.wsDest.cell(row=5, column=2).value = self.setAnLine.text()
        try:
            self.wbDest.save(self.myDest)
        except PermissionError:
            self.msSecCall("Datele nu s-au introdus in fisierul Destination.xlsx!")
        self.dialBox.close()

    def resetFunc(self):
        self.wsDest.cell(row=1, column=2).value = ""
        self.wsDest.cell(row=2, column=2).value = ""
        self.wsDest.cell(row=3, column=2).value = ""
        self.wsDest.cell(row=4, column=2).value = ""
        self.wsDest.cell(row=5, column=2).value = ""
        try:
            self.wbDest.save(self.myDest)
        except PermissionError:
            self.msSecCall("Datele nu s-au introdus in fisierul Destination.xlsx!")

        self.setRegLine.setText("")
        self.setSaidiLine.setText("")
        self.setRapLine.setText("")
        self.setDsLine.setText("")
        self.setAnLine.setText("")

    def destLoad(self):
        self.myDest = os.path.abspath(".") + "/Bundle/Destination.xlsx"
        self.wbDest = load_workbook(self.myDest)
        self.wsDest = self.wbDest["Destination"]

    def deranjTrig(self):
        #Verific ferestrele din meniul New, le folosesc la loadPT
        self.deranjControl = True
        self.alControl = False
        self.decControl = False

        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Deranjament')
        self.dialBox.setStyleSheet('background-color: #424242;')

        # Oficiul
        ofFrame = QFrame()
        ofFrame.setFrameShape(QFrame.StyledPanel)
        ofFrame.setStyleSheet('background-color: #544637;')


        ofLabel = QLabel()
        ofLabel.setText('Oficiul:')
        ofLabel.setStyleSheet('color: #e3e3e3')

        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #544637; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(150)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.currentTextChanged.connect(self.loadPt)

        hbox = QHBoxLayout()
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        ofFrame.setLayout(hbox)

        # Transmis
        transmisFrame= QFrame()
        transmisFrame.setFrameShape(QFrame.StyledPanel)
        transmisFrame.setStyleSheet('background-color: #544637;')

        transmisLabel = QLabel()
        transmisLabel.setText('Transmis:')
        transmisLabel.setStyleSheet('color: #e3e3e3')

        self.sfLine = QLineEdit()
        self.sfLine.setStyleSheet('background-color: #ffffff')
        self.sfLine.setFixedWidth(150)

        hbox = QHBoxLayout()
        hbox.addWidget(transmisLabel)
        hbox.addWidget(self.sfLine)
        transmisFrame.setLayout(hbox)

        # Instalatia
        instalatiaFrame = QFrame()
        instalatiaFrame.setFrameShape(QFrame.StyledPanel)
        instalatiaFrame.setStyleSheet('background-color: #544637;')

        instalatiaLabel = QLabel()
        instalatiaLabel.setText('Instalatia:')
        instalatiaLabel.setStyleSheet('color: #e3e3e3')

        # instalatiaList = []
        # instalatia = open("Bundle/Instalatia_deranj.txt", "r")
        # for i in instalatia:
        #     i = i.strip()
        #     instalatiaList.append(i)
        for i in self.dbGen.nomenclator.find({"name": "instalatia"}):
            instList = i["lista"]

        self.instalatiaCombo = QComboBox()
        self.instalatiaCombo.setStyleSheet('background-color: #544637; color: #e3e3e3;  height:20')
        self.instalatiaCombo.setFixedWidth(150)
        self.instalatiaCombo.addItems(instList)
        self.instalatiaCombo.setEditable(True)
        self.instalatiaCombo.setCurrentText("Alege:")

        hbox = QHBoxLayout()
        hbox.addWidget(instalatiaLabel)
        hbox.addWidget(self.instalatiaCombo)
        instalatiaFrame.setLayout(hbox)

        # Sectorul
        sectFrame = QFrame()
        sectFrame.setFrameShape(QFrame.StyledPanel)
        sectFrame.setStyleSheet('background-color: #544637;')

        sectLabel = QLabel()
        sectLabel.setText('Sectorul:')
        sectLabel.setStyleSheet('color: #e3e3e3')

        self.sectCombo = QComboBox()
        self.sectCombo.setStyleSheet('background-color: #544637; color: #e3e3e3;  height:20')
        self.sectCombo.setFixedWidth(150)
        self.sectCombo.setEditable(True)

        hbox = QHBoxLayout()
        hbox.addWidget(sectLabel)
        hbox.addWidget(self.sectCombo)
        sectFrame.setLayout(hbox)

        # F 10kV
        f10kvFrame = QFrame()
        f10kvFrame.setFrameShape(QFrame.StyledPanel)
        f10kvFrame.setStyleSheet('background-color: #544637;')

        f10kvLabel = QLabel()
        f10kvLabel.setText('Fid.10kV nr:')
        f10kvLabel.setStyleSheet('color: #e3e3e3')

        self.f10kvLine = QLineEdit()
        self.f10kvLine.setStyleSheet('background-color: #ffffff')
        self.f10kvLine.setFixedWidth(150)

        hbox = QHBoxLayout()
        hbox.addWidget(f10kvLabel)
        hbox.addWidget(self.f10kvLine)
        f10kvFrame.setLayout(hbox)

        #PT frame
        ptFrame = QFrame()
        ptFrame.setFrameShape(QFrame.StyledPanel)
        ptFrame.setStyleSheet('background-color: #544637;')

        ptLabel = QLabel()
        ptLabel.setText('PT, Fider:')
        ptLabel.setStyleSheet('color: #e3e3e3')

        # PT Section
        self.ptLine = QLineEdit()
        self.ptLine.setStyleSheet('background-color: #ffffff')
        self.ptLine.setFixedWidth(150)
        self.loadPt()
        self.ptLine.setText("PT")
        rxPt = QRegExp("(PT|PD)\d.(dot)?.(dot)")
        rxPt.setCaseSensitivity(Qt.CaseInsensitive)
        myValidator = QRegExpValidator(rxPt)
        self.ptLine.setValidator(myValidator)

        # Pentru QlineEdit folosesc clasul MyLineEdit ca sa pot folosi FocusInEvent
        self.ptFidLine = QLineEdit()
        self.ptFidLine.setStyleSheet('background-color: #ffffff')
        self.ptFidLine.setFixedWidth(150)
        self.ptFidLine.setText("Fider nr.")
        rxGr = QRegExp("Fider nr.\d\d?\d?")
        myGrValidator = QRegExpValidator(rxGr)
        self.ptFidLine.setValidator(myGrValidator)

        gridPtFrame = QGridLayout()
        ptFrame.setLayout(gridPtFrame)

        gridPtFrame.addWidget(ptLabel, 0, 0)
        gridPtFrame.addWidget(self.ptLine, 0, 1)
        gridPtFrame.addWidget(self.ptFidLine, 1, 1)

        # Continutul
        continFrame = QFrame()
        continFrame.setFrameShape(QFrame.StyledPanel)
        continFrame.setStyleSheet('background-color: #544637;')

        continLabel = QLabel()
        continLabel.setText('Continutul:')
        continLabel.setStyleSheet('color: #e3e3e3')

        self.continText = QTextEdit()
        self.continText.setStyleSheet("Background-color: rgb(255, 255, 255)")
        self.continText.setFixedHeight(45)
        self.continText.setFixedWidth(150)

        hbox = QHBoxLayout()
        hbox.addWidget(continLabel)
        hbox.addWidget(self.continText)
        continFrame.setLayout(hbox)

        # Buttons Section (butoanele "Ok", "Cancel"
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okDeranjTrigger)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        # Gneral grid
        grid = QGridLayout()
        self.dialBox.setLayout(grid)
        grid.addWidget(ofFrame, 1, 0)
        grid.addWidget(transmisFrame, 2, 0)
        grid.addWidget(sectFrame, 3, 0)
        grid.addWidget(instalatiaFrame, 4, 0)
        grid.addWidget(f10kvFrame, 5, 0)
        grid.addWidget(ptFrame, 6, 0)
        grid.addWidget(continFrame, 7, 0)
        grid.addWidget(btFrame, 8, 0)

        self.dialBox.exec()

    def decTrig(self):
        self.decControl = True
        self.deranjControl = False

        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Deconectare neplanificata')
        self.dialBox.setStyleSheet('background-color: #424242;')

        # Nr. Deconectarii
        nrDecFrame = QFrame()
        nrDecFrame.setFrameShape(QFrame.StyledPanel)
        nrDecFrame.setStyleSheet('background-color: #5c3838;')


        ofLabel = QLabel()
        ofLabel.setText('Oficiul:')
        ofLabel.setStyleSheet('color: #e3e3e3')

    #Incarc oficiile din MongoDB
        # self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #5c3838; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(140)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.currentTextChanged.connect(self.loadPt)

        hbox = QHBoxLayout()
        # hbox.addWidget(nrLabel)
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        nrDecFrame.setLayout(hbox)

        #PT frame
        ptFrame = QFrame()
        ptFrame.setFrameShape(QFrame.StyledPanel)
        ptFrame.setStyleSheet('background-color: #5c3838;')

        ptLabel = QLabel()
        ptLabel.setText('Numarul PT, Fider:')
        ptLabel.setStyleSheet('color: #e3e3e3; margin-right: 25%')

        # PT Section
        self.ptLine = QLineEdit()
        self.ptLine.setStyleSheet('background-color: #ffffff')
        self.ptLine.setFixedWidth(140)
        self.loadPt()
        # self.ptList = []
        # for i in range(2, self.wsPt.max_row + 1):
        #     self.ptList.append(self.wsPt.cell(row=i, column=1).value)
        # ptCompleter = QCompleter(self.ptList)
        # ptCompleter.setCaseSensitivity(Qt.CaseInsensitive)
        # self.ptLine.setCompleter(ptCompleter)
        self.ptLine.setText("PT")
        rxPt = QRegExp("(PT|PD)\d.(dot)?.(dot)")
        rxPt.setCaseSensitivity(Qt.CaseInsensitive)
        myValidator = QRegExpValidator(rxPt)
        self.ptLine.setValidator(myValidator)

        # Pentru QlineEdit folosesc clasul MyLineEdit ca sa pot folosi FocusInEvent
        self.ptFidLine = QLineEdit()
        self.ptFidLine.setStyleSheet('background-color: #ffffff')
        self.ptFidLine.setFixedWidth(140)
        self.ptFidLine.setText("Fider nr.")
        rxGr = QRegExp("Fider nr.\d\d?\d?")
        myGrValidator = QRegExpValidator(rxGr)
        self.ptFidLine.setValidator(myGrValidator)

        gridPtFrame = QGridLayout()
        ptFrame.setLayout(gridPtFrame)

        gridPtFrame.addWidget(ptLabel, 0, 0)
        gridPtFrame.addWidget(self.ptLine, 0, 1)
        gridPtFrame.addWidget(self.ptFidLine, 1, 1)

        # Data frame
        dtFrame = QFrame()
        dtFrame.setFrameShape(QFrame.StyledPanel)
        dtFrame.setStyleSheet('background-color: #5c3838;')

        dtLabel = QLabel()
        dtLabel.setText('Data si ora deconectarii:')
        dtLabel.setStyleSheet('color: #e3e3e3')

        self.dtLine = DateLineEdit()
        self.dtLine.setStyleSheet("Background-color: rgb(255, 255, 255)")
        self.dtLine.setFixedWidth(140)

        rxDt = QRegExp("\d\d.\d\d.\d\d \d:\d")
        myDtValidator = QRegExpValidator(rxDt)
        self.dtLine.setValidator(myDtValidator)
        pyDateTime = datetime.datetime.now().strftime("%d.%m.%y %H:%M")
        self.dtLine.setText(pyDateTime)

        hbox = QHBoxLayout()
        hbox.addWidget(dtLabel)
        hbox.addWidget(self.dtLine)
        dtFrame.setLayout(hbox)

        cauzaFrame = QFrame()
        cauzaFrame.setFrameShape(QFrame.StyledPanel)
        cauzaFrame.setStyleSheet('background-color: #314652;')

        cauzaLabel = QLabel()
        cauzaLabel.setText('Cauza deconectarii:')
        cauzaLabel.setStyleSheet('color: #e3e3e3; margin-right: 15%')

        # wbCauza = load_workbook(os.path.abspath(".") + "/Bundle/Cauza_deconectare.xlsx")
        # wsCauza = wbCauza.active
        # cauzaList = []
        # for i in range(1, wsCauza.max_row + 1):
        #     cauzaList.append(wsCauza.cell(row=i, column=1).value)
        for i in self.dbGen.nomenclator.find({"name": "cauza_dec"}):
            cauzaList = i["lista"]
        self.cauzaCombo = QComboBox()
        self.cauzaCombo.setStyleSheet('background-color: #314652; color: #e3e3e3;  height:20')
        self.cauzaCombo.setFixedWidth(140)
        self.cauzaCombo.setEditable(True)
        self.cauzaCombo.addItems(cauzaList)
        # self.cauzaCombo.setFixedWidth(120)

        hbox = QHBoxLayout()
        hbox.addWidget(cauzaLabel)
        hbox.addWidget(self.cauzaCombo)
        cauzaFrame.setLayout(hbox)

        # termenFrame = QFrame()
        # termenFrame.setFrameShape(QFrame.StyledPanel)
        # termenFrame.setStyleSheet('background-color: #314652;')
        #
        # termenLabel = QLabel()
        # termenLabel.setText('Termen reglementat:')
        # termenLabel.setStyleSheet('color: #e3e3e3')
        #
        # termenList = ['Incadrat', 'Neincadrat']
        # self.termenCombo = QComboBox()
        # self.termenCombo.setStyleSheet('background-color: #314652; color: #e3e3e3;  height:20')
        # self.termenCombo.setFixedWidth(140)
        # self.termenCombo.addItems(termenList)
        #
        # hbox = QHBoxLayout()
        # hbox.addWidget(termenLabel)
        # hbox.addWidget(self.termenCombo)
        # termenFrame.setLayout(hbox)

        # Buttons Section (butoanele "Ok", "Cancel"
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okDecTrigger)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        # Gneral grid
        grid = QGridLayout()
        self.dialBox.setLayout(grid)
        grid.addWidget(nrDecFrame, 1, 0)
        grid.addWidget(ptFrame, 2, 0)
        grid.addWidget(dtFrame, 3, 0)
        grid.addWidget(cauzaFrame, 4, 0)
        # grid.addWidget(termenFrame, 5, 0)
        grid.addWidget(btFrame, 6, 0)

        self.dialBox.exec()


    def dsTrig(self):
        self.alControl = False
        self.decControl = False
        self.deranjControl = False


        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Dispozitie')
        self.dialBox.setStyleSheet('background-color: #424242;')

        #Nr. Dispozitiei frame
        nrDsFrame = QFrame()
        nrDsFrame.setFrameShape(QFrame.StyledPanel)
        nrDsFrame.setStyleSheet('background-color: #314652;')

        # for i in range(self.wsRegAl.max_row, 3, -1):
        #     if self.wsRegAl.cell(row=i, column=3).value != None:
        #         self.nrAlDs = int(self.wsRegAl.cell(row=i, column=3).value) + 1
        #         break
        # for i in self.db.reg_al_un.find({"nr_ds":{"$ne":""}}).sort({"_id":-1}).limit(1):
        #     print(i["nr_ds"])

        # for i in self.db.reg_al_un.find({"nr_ds":{"$ne":""}}).sort("_id",-1).limit(1):
        #     self.nrAlDs = int(i["nr_ds"]) + 1
        #
        # nrLabel = QLabel()
        # nrLabel.setStyleSheet('color: #e3e3e3;')
        # nrLabel.setText('Numarul dispozitiei: ' + str(self.nrAlDs))

        ofLabel = QLabel()
        ofLabel.setText('Oficiul:')
        ofLabel.setStyleSheet('color: #e3e3e3;')

        # Incarc oficiile din MongoDB
        # self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #314652; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(180)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.setEnabled(True)
        self.ofCombo.setEditable(True)
        self.ofCombo.setCurrentText(self.nameOf)
        self.ofCombo.currentTextChanged.connect(self.loadPt)

        hbox = QHBoxLayout()
        # hbox.addWidget(nrLabel)
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        nrDsFrame.setLayout(hbox)

        # Sef de lucrari Frame (Sef de lucrari, grupa TS), Load Excel files
        sefFrame = QFrame()
        sefFrame.setFrameShape(QFrame.StyledPanel)
        sefFrame.setStyleSheet('background-color: #314652;')

        sefLabel = QLabel()
        sefLabel.setStyleSheet('color: #e3e3e3;')
        sefLabel.setText('Sef de lucrari:')

        self.sfLine = QLineEdit()
        self.sfLine.setStyleSheet('background-color: #ffffff')
        self.sfLine.setFixedWidth(180)

        hbox = QHBoxLayout()
        hbox.addWidget(sefLabel)
        hbox.addWidget(self.sfLine)
        sefFrame.setLayout(hbox)

        # Membrii echipei frame
        memEchFrame = QFrame()
        memEchFrame.setFrameShape(QFrame.StyledPanel)
        memEchFrame.setStyleSheet('background-color: #314652;')

        memEchLabel = QLabel()
        memEchLabel.setStyleSheet('color: #e3e3e3;')
        memEchLabel.setText('Membrii echipei:')

        self.memEchLine = QLineEdit()
        self.memEchLine.setStyleSheet('background-color: #ffffff')
        self.memEchLine.setFixedWidth(180)

        hbox = QHBoxLayout()
        hbox.addWidget(memEchLabel)
        hbox.addWidget(self.memEchLine)
        memEchFrame.setLayout(hbox)

        # Emitent Frame
        emFrame = QFrame()
        emFrame.setFrameShape(QFrame.StyledPanel)
        emFrame.setStyleSheet('background-color: #314652;')

        emLabel = QLabel()
        emLabel.setStyleSheet('color: #e3e3e3;')
        emLabel.setText('Emitent :')

        self.emLine = QLineEdit()
        self.emLine.setStyleSheet('background-color: #ffffff')
        self.emLine.setFixedWidth(180)

        hbox = QHBoxLayout()
        hbox.addWidget(emLabel)
        hbox.addWidget(self.emLine)
        emFrame.setLayout(hbox)

        # Se executa frame
        exFrame = QFrame()
        exFrame.setFrameShape(QFrame.StyledPanel)
        exFrame.setStyleSheet('background-color: #6b6843;')

        exLabel = QLabel()
        exLabel.setStyleSheet('color: #e3e3e3;')
        exLabel.setText('Se executa:')

        # Instalatia section

        # wb = load_workbook('Bundle/Instalatia.xlsx')
        # ws = wb.active
        # instList = []
        # for i in range(1, ws.max_row + 1):
        #     instList.append(ws.cell(row=i, column=1).value)
        for i in self.dbGen.nomenclator.find({"name": "instalatia"}):
            instList = i["lista"]

        # Widgetul QlineEdit este inlocuit cu clasul MyLineEdit creat
        # mai sus pentru a putea folosi FocusIN si OUT Event
        self.instLine = QComboBox()
        self.instLine.addItems(instList)
        self.instLine.setEditable(True)
        self.instLine.setStyleSheet('color: #e3e3e3;  height:20')
        self.instLine.setFixedWidth(180)

        # PT Sectiom
        self.ptLine = QLineEdit()
        self.ptLine.setStyleSheet('background-color: #ffffff')
        self.loadPt()
        self.ptLine.setText("PT")
        self.ptLine.setFixedWidth(180)

        # Pentru QlineEdit folosesc clasul MyLineEdit ca sa pot folosi FocusInEvent
        self.ptFidLine = QLineEdit()
        self.ptFidLine.setStyleSheet('background-color: #ffffff')
        # self.ptFidLine.setFixedWidth(20)
        self.ptFidLine.setText("Fider nr.")
        rxGr = QRegExp("Fider nr.\d(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)"
                       "(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)")
        myGrValidator = QRegExpValidator(rxGr)
        self.ptFidLine.setValidator(myGrValidator)
        self.ptFidLine.setFixedWidth(180)

        # Lucrari Section
        # wb = load_workbook('Bundle/Lucrarile.xlsx')
        # ws = wb.active
        # lucrList = []
        for i in self.dbGen.nomenclator.find({"name": "lucrari"}):
            lucrList = i["lista"]
        lucrList.sort()
        self.lucrLine = QComboBox()
        self.lucrLine.setStyleSheet('color: #e3e3e3;  height:20')
        self.lucrLine.addItems(lucrList)
        self.lucrLine.setEditable(True)
        self.lucrLine.setCurrentText("Lucrarile efectuate:")
        self.lucrLine.setFixedWidth(180)
        self.lucrLine.currentTextChanged.connect(self.lucruLineCh)

        gridExFrame = QGridLayout()
        exFrame.setLayout(gridExFrame)

        gridExFrame.addWidget(exLabel, 0, 0)
        gridExFrame.addWidget(self.instLine, 0, 1)
        gridExFrame.addWidget(self.ptLine, 1, 1)
        gridExFrame.addWidget(self.ptFidLine, 2, 1)
        gridExFrame.addWidget(self.lucrLine, 3, 1)
        # gridExFrame.addWidget(self.decCombo, 4, 1)

        # Masurile de protectie si prevenire FRAME
        masFrame = QFrame()
        masFrame.setFrameShape(QFrame.StyledPanel)
        masFrame.setStyleSheet('background-color: #6b6843;')

        self.formCheck = QCheckBox("Completeaza formular. Masurile de protectie si prevenire TS:")
        self.formCheck.setStyleSheet("color: rgb(255, 255, 255)")
        self.formCheck.stateChanged.connect(self.changeForm)

        self.dtLabel = QLabel("Alege data:")
        self.dtLabel.setStyleSheet("color: rgb(255, 255, 255)")

        self.dtButton = QPushButton()
        self.dtButton.setIconSize(QSize(25, 25))
        self.dtButton.setIcon(QIcon("calendar-blue.ico"))
        self.dtButton.setFlat(True)
        # self.dtButton.setAutoFillBackground(QPalette.Disabled)
        self.dtButton.setFixedSize(30, 30)
        self.dtButton.setStyleSheet("border-style: outset")
        self.dtButton.clicked.connect(self.dtButtonFunc)

        self.dtLine = QLineEdit()
        self.dtLine.setStyleSheet("background-color: rgb(255, 255, 255)")
        self.dtLine.setReadOnly(True)

        hbox = QHBoxLayout()
        hbox.addWidget(self.dtLabel)
        hbox.addWidget(self.dtLine)
        hbox.addWidget(self.dtButton)

        # self.dsCalendar = QLabel("Alege data lucrarilor: ")
        # self.dsCalendar.setStyleSheet("color: rgb(255, 255, 255)")

        self.masLine_1 = QLineEdit()
        self.masLine_1.setStyleSheet('background-color: #ffffff;  height:20')
        self.masLine_1.setText("Fara deconectari, de respectat distanta minim admisa (Tabelul nr. 1 NEI);")
        self.masLine_1.setCursorPosition(0)
        self.masLine_1.setEnabled(False)
        self.masLine_2 = QLineEdit()
        self.masLine_2.setStyleSheet('background-color: #ffffff;  height:20')
        self.masLine_2.setText("De respectat cerintele instructiunei SSM ME1 si ME2;")
        self.masLine_2.setCursorPosition(0)
        self.masLine_2.setEnabled(False)
        self.masLine_3 = QLineEdit()
        self.masLine_3.setStyleSheet('background-color: #ffffff;  height:20')
        self.masLine_3.setText("De respectat cerintele comisiei nationale extraordinare de sanatate"
                               "publica de prevenire a infectiei COVID 19.")
        self.masLine_3.setCursorPosition(0)
        self.masLine_3.setEnabled(False)

        # self.progressDs = QProgressBar()

        self.vboxPr = QVBoxLayout()
        self.vboxPr.addWidget(self.formCheck)
        self.vboxPr.addLayout(hbox)
        self.vboxPr.addWidget(self.masLine_1)
        self.vboxPr.addWidget(self.masLine_2)
        self.vboxPr.addWidget(self.masLine_3)
        # self.vboxPr.addWidget(self.progressDs)
        masFrame.setLayout(self.vboxPr)


        # Buttons Section (butoanele "Ok", "Cancel"
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okDsTrigger)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        hbox = QHBoxLayout()
        hbox.addWidget(btOk)
        hbox.addWidget(btCancel)
        btFrame.setLayout(hbox)

        # Gneral grid
        grid = QGridLayout()
        self.dialBox.setLayout(grid)
        grid.addWidget(nrDsFrame, 1, 0)
        grid.addWidget(sefFrame, 2, 0)
        grid.addWidget(memEchFrame, 3, 0)
        grid.addWidget(emFrame, 4, 0)
        grid.addWidget(exFrame, 5, 0)
        grid.addWidget(masFrame, 6, 0)
        grid.addWidget(btFrame, 7, 0)

        # Fac control pentru dublare
        if self.dublControl:
            self.ofCombo.setCurrentText(self.abrOficii_invers(self.data.at[self.modRow, 0]))
            self.sfLine.setText(self.data.at[self.modRow, 8])
            if self.data.at[self.modRow, 9] != "":
                myList = self.data.at[self.modRow, 9].split(":")
                self.memEchLine.setText(myList[1])
            if not re.search("Confirmat:", self.data.at[self.modRow, 10]):
                self.emLine.setText(self.data.at[self.modRow, 10])
            else:
                myList = self.data.at[self.modRow, 10].split("\n")
                self.emLine.setText(myList[1])
            self.instLine.setCurrentText(self.data.at[self.modRow, 3])
            self.ptLine.setText(self.data.at[self.modRow, 4])
            self.ptFidLine.setText(self.data.at[self.modRow, 6])
            self.lucrLine.setCurrentText(self.data.at[self.modRow, 7])

        self.dialBox.exec()

    def alTrig(self):
        self.alControl = True
        self.decControl = False
        self.deranjControl = False

        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Autorizatie')
        self.dialBox.setStyleSheet('background-color: #424242;')

        # Dialog First Frame (Nr. autorizatie, Oficiul)
        nrautFrame = QFrame()
        nrautFrame.setFrameShape(QFrame.StyledPanel)
        nrautFrame.setStyleSheet('background-color: #315240;')

        # #Incarc registrul din excel
        # for i in range(self.wsRegAl.max_row, 3, -1):
        #     if self.wsRegAl.cell(row=i, column=2).value != None:
        #         self.nrAlDs = int(self.wsRegAl.cell(row=i, column=2).value) + 1
        #         break

        # for i in self.db.reg_al_un.find({"nr_al":{"$ne":""}}).sort("_id",-1).limit(1):
        #     self.nrAlDs = int(i["nr_al"]) + 1
        #
        # nrLabel = QLabel()
        # nrLabel.setStyleSheet('color: #e3e3e3;')
        # nrLabel.setText('Numarul autorizatiei: ' + str(self.nrAlDs))

        ofLabel = QLabel()
        ofLabel.setText('Oficiul:')
        ofLabel.setStyleSheet('color: #e3e3e3;')

        # Incarc oficiile din MongoDB
        # self.loadOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.setStyleSheet('background-color: #315240; color: #e3e3e3;  height:20')
        self.ofCombo.setFixedWidth(180)
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.setEnabled(True)
        self.ofCombo.setCurrentText(self.nameOf)
        self.ofCombo.currentTextChanged.connect(self.loadPt)

        hbox = QHBoxLayout()
        # hbox.addWidget(nrLabel)
        hbox.addWidget(ofLabel)
        hbox.addWidget(self.ofCombo)
        nrautFrame.setLayout(hbox)

        # Sef de lucrari Frame (Sef de lucrari, grupa TS), Load Excel files
        sefFrame = QFrame()
        sefFrame.setFrameShape(QFrame.StyledPanel)
        sefFrame.setStyleSheet('background-color: #315240;')

        sefLabel = QLabel()
        sefLabel.setStyleSheet('color: #e3e3e3;')
        sefLabel.setText('Sef de lucrari:')

        self.sfLine = QLineEdit()
        self.sfLine.setStyleSheet('background-color: #ffffff')
        self.sfLine.setFixedWidth(180)

        gridSefFrame = QGridLayout()
        sefFrame.setLayout(gridSefFrame)

        gridSefFrame.addWidget(sefLabel, 0, 0)
        gridSefFrame.addWidget(self.sfLine, 0, 1)

        # Membrii echipei frame
        memEchFrame = QFrame()
        memEchFrame.setFrameShape(QFrame.StyledPanel)
        memEchFrame.setStyleSheet('background-color: #315240;')

        memEchLabel = QLabel()
        memEchLabel.setStyleSheet('color: #e3e3e3;')
        memEchLabel.setText('Membrii echipei:')

        self.memEchLine = QLineEdit()
        self.memEchLine.setStyleSheet('background-color: #ffffff')
        self.memEchLine.setFixedWidth(180)

        hbox = QHBoxLayout()
        hbox.addWidget(memEchLabel)
        hbox.addWidget(self.memEchLine)
        memEchFrame.setLayout(hbox)

        # Emitent Frame
        emFrame = QFrame()
        emFrame.setFrameShape(QFrame.StyledPanel)
        emFrame.setStyleSheet('background-color: #315240;')

        emLabel = QLabel()
        emLabel.setStyleSheet('color: #e3e3e3;')
        emLabel.setText('Emitent :        ')

        self.emLine = QLineEdit()
        self.emLine.setStyleSheet('background-color: #ffffff')
        self.emLine.setFixedWidth(180)

        gridEmFrame = QGridLayout()
        emFrame.setLayout(gridEmFrame)

        gridEmFrame.addWidget(emLabel, 0, 0)
        gridEmFrame.addWidget(self.emLine, 0, 1)

        # Se executa frame
        exFrame = QFrame()
        exFrame.setFrameShape(QFrame.StyledPanel)
        exFrame.setStyleSheet('background-color: #6b6843;')

        exLabel = QLabel()
        exLabel.setStyleSheet('color: #e3e3e3; margin-right: 2')
        exLabel.setText('Se executa:')

        # Instalatia section

        # wb = load_workbook('Bundle/Instalatia.xlsx')
        # ws = wb.active
        # instList = []
        # for i in range(1, ws.max_row + 1):
        #     instList.append(ws.cell(row=i, column=1).value)
        for i in self.dbGen.nomenclator.find({"name": "instalatia"}):
            instList = i["lista"]

        # Widgetul QlineEdit este inlocuit cu clasul MyLineEdit creat
        # mai sus pentru a putea folosi FocusIN si OUT Event
        self.instLine = QComboBox()
        self.instLine.addItems(instList)
        self.instLine.setEditable(True)
        self.instLine.setStyleSheet('color: #e3e3e3;  height:20')
        self.instLine.setFixedWidth(180)

        # PT Section
        self.ptLine = QLineEdit()
        self.ptLine.setStyleSheet('background-color: #ffffff')
        self.loadPt()

        self.ptLine.setText("PT")
        rxPt = QRegExp("(PT|PD)\d.(dot)?.(dot)")
        rxPt.setCaseSensitivity(Qt.CaseInsensitive)
        myValidator = QRegExpValidator(rxPt)
        self.ptLine.setValidator(myValidator)
        self.ptLine.setFixedWidth(180)

        #Pentru QlineEdit folosesc clasul MyLineEdit ca sa pot folosi FocusInEvent
        self.ptFidLine = QLineEdit()
        self.ptFidLine.setStyleSheet('background-color: #ffffff')
        self.ptFidLine.setFixedWidth(180)
        self.ptFidLine.setText("Fider nr.")
        # rxGr = QRegExp("Fider nr.\d\d?\d?")
        # myGrValidator = QRegExpValidator(rxGr)
        # self.ptFidLine.setValidator(myGrValidator)

        #Lucrari Section
        # wb = load_workbook('Bundle/Lucrarile.xlsx')
        # ws = wb.active
        # lucrList = []
        # for i in range(1, ws.max_row + 1):
        #     lucrList.append(ws.cell(row=i, column=1).value)
        for i in self.dbGen.nomenclator.find({"name": "lucrari"}):
            lucrList = i["lista"]
        lucrList.sort()
        self.lucrLine = QComboBox()
        self.lucrLine.setStyleSheet('color: #e3e3e3;  height:20')
        self.lucrLine.addItems(lucrList)
        self.lucrLine.setEditable(True)
        self.lucrLine.setCurrentText("Lucrarile efectuate:")
        self.lucrLine.setFixedWidth(180)

        self.smNr = QLineEdit()
        self.smNr.setStyleSheet('background-color: #ffffff')
        self.smNr.setText("SM nr.")
        rxGr = QRegExp("SM nr.\d(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)"
                       "(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)(\d?|,)")
        myGrValidator = QRegExpValidator(rxGr)
        self.smNr.setValidator(myGrValidator)
        self.smNr.setFixedWidth(180)

        for i in self.dbGen.nomenclator.find({"name": "cu_dec"}):
            cu_decList = i["lista"]
        self.decCombo = QComboBox()
        self.decCombo.setStyleSheet('background-color: #6b6843; color: #e3e3e3; height:20')
        self.decCombo.addItems(cu_decList)
        self.decCombo.setEditable(True)

        gridExFrame = QGridLayout()
        exFrame.setLayout(gridExFrame)

        gridExFrame.addWidget(exLabel, 0, 0)
        gridExFrame.addWidget(self.instLine, 0, 1)
        gridExFrame.addWidget(self.ptLine, 1, 1)
        gridExFrame.addWidget(self.ptFidLine, 2, 1)
        gridExFrame.addWidget(self.lucrLine, 3, 1)
        gridExFrame.addWidget(self.smNr, 4, 1)
        gridExFrame.addWidget(self.decCombo, 5, 1)

        # Buttons Section (butoanele "Ok", "Cancel"
        btFrame = QFrame()
        btFrame.setFrameShape(QFrame.StyledPanel)

        btOk = QPushButton('Ok')
        btOk.setStyleSheet('color: #e3e3e3')
        btOk.clicked.connect(self.okTrigger)

        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        btCancel.clicked.connect(self.dialBox.close)

        gridBt = QGridLayout()
        btFrame.setLayout(gridBt)
        gridBt.addWidget(btOk, 0, 0)
        gridBt.addWidget(btCancel, 0, 1)

        # General grid
        grid = QGridLayout()
        self.dialBox.setLayout(grid)
        grid.addWidget(nrautFrame,1,0)
        grid.addWidget(sefFrame, 2, 0)
        grid.addWidget(memEchFrame, 3, 0)
        grid.addWidget(emFrame, 4, 0)
        grid.addWidget(exFrame, 5,0)
        grid.addWidget(btFrame, 6, 0)

        # Fac control pentru dublare
        if self.dublControl:
            self.ofCombo.setCurrentText(self.abrOficii_invers(self.data.at[self.modRow, 0]))
            self.sfLine.setText(self.data.at[self.modRow, 8])
            if self.data.at[self.modRow, 9] != "":
                myList = self.data.at[self.modRow, 9].split(":")
                self.memEchLine.setText(myList[1])
            if not re.search("Confirmat:", self.data.at[self.modRow, 10]):
                self.emLine.setText(self.data.at[self.modRow, 10])
            else:
                myList = self.data.at[self.modRow, 10].split("\n")
                self.emLine.setText(myList[1])
            self.instLine.setCurrentText(self.data.at[self.modRow, 3])
            self.ptLine.setText(self.data.at[self.modRow, 4])
            self.ptFidLine.setText(self.data.at[self.modRow, 6])
            self.lucrLine.setCurrentText(self.data.at[self.modRow, 7])
            if self.data.at[self.modRow, 12] != "":
                myList = self.data.at[self.modRow, 12].split("SM nr.")
                self.smNr.setText("SM nr." + myList[1])
            self.decCombo.setCurrentText(self.data.at[self.modRow, 11])

        self.dialBox.exec()

    #  Functie prelucrare erori cimpul autorizatii
    def msCall(self, intrCimp):
        self.intrCimp = intrCimp
        self.intrMB = QMessageBox()
        self.intrMB.setIcon(QMessageBox.Warning)
        self.intrMB.setWindowTitle('Atentie!')
        self.intrMB.setText('Introduceti datele in cimpul:\n' + self.intrCimp)
        self.intrMB.setWindowIcon(QIcon('ataman_logo.ico'))
        self.intrMB.exec()

    def msSecCall(self, myMess):
        self.myMess = myMess
        self.secMB = QMessageBox()
        self.secMB.setIcon(QMessageBox.Information)
        self.secMB.setWindowTitle('Pentru informare:')
        self.secMB.setText(self.myMess)
        self.secMB.setWindowIcon(QIcon('ataman_logo.ico'))
        self.secMB.exec()

    def msThrdCall(self, myMess):
        self.myMess = myMess
        self.secMB = QMessageBox()
        self.secMB.setIcon(QMessageBox.Information)
        self.secMB.setWindowTitle('Pentru informare:')
        self.secMB.setText(self.myMess)
        self.secMB.setWindowIcon(QIcon('ataman_logo.ico'))

    def resetCall(self, myMess):
        self.myMess = myMess
        self.secMB = QMessageBox()
        self.secMB.setIcon(QMessageBox.Warning)
        self.secMB.setWindowTitle('Atentie')
        self.secMB.setText(self.myMess)
        self.secMB.setWindowIcon(QIcon('ataman_logo.ico'))
        self.secMB.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        self.secMB.setDefaultButton(QMessageBox.Cancel)

    global elSoundContr
    elSoundContr = True

    def msCerere(self):
        if elSoundContr:
            global elSound
            elSound = True

            workerSound = WorkerSound()
            self.threadpoolSound = QThreadPool()
            self.threadpoolSound.start(workerSound)
            # self.signalSound.finished.connect(self.myStop)

            if self.angDoc["position"] == "Dispecer":
                self.resetCall("Aveti o cerere in registru de DS/AL.\n"
                               "Apasati 'Ok', daca doriti sa o vedeti!")
                if self.secMB.exec() == QMessageBox.Ok:
                    self.centrAlPop()
                    elSound = False
                else:
                    self.secMB.close()
                    elSound = False

    # def myStop(self):
    #     pass

    def resetTrig(self):
        self.resetCall("Toate datele din cimpuri vor fi sterse!")
        if  self.secMB.exec() == QMessageBox.Ok:
            self.resetFunc()
        else:
            self.secMB.close()

    # def msCerereThread(self):
    #     Thread(target=self.myPlay).start()
    #     Thread(target=self.msCerere).start()


    #Functie populez Biroul dispecerului cu Autorizatie
    def centrAlPop(self):
        self.alFunc()
        global elSoundContr
        elSoundContr = True

        if self.tabWindowControl == True:
            self.tabWindow.close()

        myTime = datetime.datetime.now()
        myYear = myTime.strftime('%Y')
        myMonth = self.MonthToNumb(self.mnCombo.currentText())
        myVar = "al_" + myYear + "_" + myMonth

        self.reg_al = self.db[myVar]

        try:
            self.data = pd.DataFrame(self.reg_al.find())
            myColumn = self.data.pop("_id")
            self.data.insert(20, "_id", myColumn)
            myColumn = self.data.pop("id")
            self.data.insert(20, "id", myColumn)
            self.data.columns = range(21)
            # print(self.data.at[2, 19])
            self.data.sort_index(ascending=False, inplace=True, ignore_index=True)
            header = ["Oficiul", "Nr. \nDS", "Nr. \nAL", "Instalatia", \
                      "PT", "Localitatea", "Nr. \nFider", "Lucrarile\nefectuate", \
                      "Sef de lucrari\nsau supravegh.\n(nume, prenum.)\ngrupa TS)", \
                      "Membrii form.\nce lucr.pe AL,DS\n(nume, prenum.,\n grupa TS)",
                      "Lucratorul care\n a emis AL (DS)\n(nume, prenum.,\ngrupa TS)",\
                      "Cu deconectare",\
                      "Masurile tehnice\n de asigurare a TS\ncu indic. deconect.\nlocurilor de\nmont legaturilor\n la pamint",\
                      "Semnatura\nlucratorilor\ncare au executat \ninstr. periodica \nsi care au\nfost instruiti",\
                      "Starea", "Pregat. locului\nde munca (data,\nora)", "Admiterea\nechipei (data,\nora)",
                      "Terminarea\nlucrarilor (data,\nora)", "Linkul", "_id", "id"]
            self.table = QTableView()
            self.model = TableModel(self.data, header)

            self.table.setModel(self.model)
            self.table.setWordWrap(True)
            self.table.setTextElideMode(Qt.ElideMiddle)
            self.table.resizeRowsToContents()
            self.table.resizeColumnsToContents()
            self.table.hideColumn(19)
            self.table.hideColumn(20)
            for i in range(20):
                self.table.setRowHeight(i, 100)
            # self.table.setColumnWidth(1, 50)
            # self.table.setColumnWidth(2, 100)
            for i in range(3, 18):
                self.table.setColumnWidth(i, 110)
            self.table.setColumnWidth(6, 50)
            # self.table.setColumnWidth(20, 150)
            # self.table.setColumnWidth(18, 170)
            self.table.setStyleSheet("Background-color: rgb(200, 200, 200)")
            self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
            self.table.verticalHeader().hide()

            regFrame = QFrame()
            regAlTitle = QLabel()
            regAlTitle_month = QLabel()

            regAlTitle.setText("Registru de autorizatii, oficiul:")
            regAlTitle.setStyleSheet('padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)')

            regAlTitle_month.setText("pentru luna:")
            regAlTitle_month.setStyleSheet('padding-left: 20%; font-size: 24px; color: rgb(191, 60, 60)')

            ofList = ['Toate oficiile']
            ofList = ofList + self.ofList
            self.ofCombo = QComboBox()
            self.ofCombo.addItems(ofList)
            # self.ofComboReg.setStyleSheet("margin-right:1400%")
            self.ofCombo.setStyleSheet("padding-left:10%; font-size:12px")
            self.ofCombo.setFixedHeight(25)
            self.ofCombo.setFixedWidth(100)
            self.ofCombo.currentTextChanged.connect(self.ofChangeReg)

            # refreshBt = QPushButton("Refresh")
            # refreshBt.setFixedSize(110, 29)
            # refreshBt.clicked.connect(self.refreshAl)

            emptyLb = QLabel("")

            hbox = QHBoxLayout()
            hbox.addWidget(regAlTitle)
            hbox.addWidget(self.ofCombo)
            hbox.addWidget(regAlTitle_month)
            hbox.addWidget(self.mnCombo)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)
            # hbox.setStretch(1, 1)

            vbox = QVBoxLayout()
            vbox.addLayout(hbox)
            vbox.addWidget(self.table)
            regFrame.setLayout(vbox)

            self.tabWindow = QMdiSubWindow()
            self.myMidi.addSubWindow(self.tabWindow)
            self.tabWindow.setWindowIcon(QIcon(QPixmap(1, 1)))
            self.tabWindow.setWidget(regFrame)
            self.tabWindow.setGeometry(100, 100, 1000, 600)
            self.tabWindow.showMaximized()
            self.tabWindowControl = True
            self.tabWindow.show()

            self.table.doubleClicked.connect(self.showAl)
        # except AttributeError:
        #     self.msSecCall("Nu exista date pentru aceasta perioada in registru DS/AL!")
        except KeyError:
            self.msThrdCall("Nu exista date pentru aceasta perioada in registru DS/AL!")
            if self.secMB.exec() == QMessageBox.Ok:
                self.mnCombo.setCurrentText(self.NumbToMonth(myTime.strftime('%m')))


    def changeForm(self):
        if self.formCheck.isChecked() == True:
            self.masLine_1.setEnabled(True)
            self.masLine_2.setEnabled(True)
            self.masLine_3.setEnabled(True)
        else:
            self.masLine_1.setEnabled(False)
            self.masLine_2.setEnabled(False)
            self.masLine_3.setEnabled(False)

    def changeExecDeranj(self):
        for i in range(len(self.data.count(axis=1))):
            if self.execCheck.isChecked() == True:
                if self.tableDeranj.isRowHidden(i) == False:
                    if self.data.at[i, 11] != "Neexecutat":
                        self.tableDeranj.hideRow(i)
            elif self.execCheck.isChecked() == False:
                self.tableDeranj.showRow(i)
                if self.ofCombo.currentText() == "Toate oficiile":
                    pass
                elif self.ofCombo.currentText() != "Toate oficiile":
                    if self.data.at[i, 0] != self.ofVar:
                        self.tableDeranj.hideRow(i)
                    if self.sectCombo.currentText() != "Alege sector:":
                        if self.data.at[i, 3] != self.sectCombo.currentText():
                            self.tableDeranj.hideRow(i)
        # print(self.execCheck.isChecked())

    def changeOfAng(self):
        # self.abrOficii()
        # print(self.data)
        for i in range(len(self.data.count(axis=1))):
            #Mai intii controlez daca a fost sau nu sters angajatul (altfel apare cind schimb oficiile)
            if self.data.at[i, "name"] != "":
                self.tableAng.showRow(i)
            if self.ofAngCombo.currentText() != "Toate oficiile":
                if self.data.at[i, "oficiul"] != self.abrOficiiSec(self.ofAngCombo.currentText()):
                    self.tableAng.hideRow(i)
        self.loadSectAng()

    def changeSectDeranj(self):
        self.execCheck.setChecked(False)
        for i in range(len(self.data.count(axis=1))):
            self.tableDeranj.showRow(i)
            if self.ofCombo.currentText() != "Toate oficiile":
                if self.data.at[i, 0] != self.ofVar:
                    self.tableDeranj.hideRow(i)
            if self.sectCombo.currentText() != "Alege sector:":
                if self.data.at[i, 3] != self.sectCombo.currentText():
                    self.tableDeranj.hideRow(i)

    def changeSectAng(self):
        # print(self.data.at[0, "position"])
        for i in range(len(self.data.count(axis=1))):
            self.tableAng.showRow(i)
            if self.ofAngCombo.currentText() != "Toate oficiile":
                if self.data.at[i, "oficiul"] != self.abrOficiiSec(self.ofAngCombo.currentText()):
                    self.tableAng.hideRow(i)
            if self.sectorCombo.currentText() != "Alege sector:":
                if self.data.at[i, "sector"] != self.sectorCombo.currentText():
                    self.tableAng.hideRow(i)
            #Selectez angajatul cautat
            if self.cautAngControl:
                if self.data.at[i, "name"] == self.uModifCombo.currentText():
                    self.tableAng.selectRow(i)

    def ofChangeReg(self):
        self.abrOficii()
        for i in range(len(self.data.count(axis=1))):
            self.table.showRow(i)
            if self.ofCombo.currentText() != "Toate oficiile":
                if self.data.at[i, 0] != self.ofVar:
                    self.table.hideRow(i)
        if self.deranjControlPop == True:
            self.execCheck.setChecked(False)
            self.loadSect()

    def ofChangeDeranj(self):
        self.abrOficii()
        for i in range(len(self.data.count(axis=1))):
            self.tableDeranj.showRow(i)
            if self.ofCombo.currentText() != "Toate oficiile":
                if self.data.at[i, 0] != self.ofVar:
                    self.tableDeranj.hideRow(i)
        self.execCheck.setChecked(False)
        self.loadSect()

    def showAl(self):
        self.modRow = self.table.currentIndex().row()
        if self.data.at[self.modRow, 1] != "":
            for i in self.reg_al.find({"nr_ds": self.data.at[self.modRow, 1]}):
                if i["link"] != "":
                    webbrowser.open(i["link"])
                else:
                    self.msSecCall("Nu exista link!")
        elif self.data.at[self.modRow, 2] != "":
            for i in self.reg_al.find({"nr_al": self.data.at[self.modRow, 2]}):
                if i["link"] != "":
                    webbrowser.open(i["link"])
                else:
                    self.msSecCall("Nu exista link!")


    def decTabWindow(self):
        if self.tabDecControl == True:
            self.tabWindDec.close()
            # self.tabDecControl = False

        self.dtContrDecZl()

        self.centralReg = []
        for i in range(7, self.wsDecPT.max_row + 1):
            myColumn = []
            for j in range(1, self.wsDecPT.max_column - 1):
                # data = np.array(centralReg.append(self.wsRegAl.cell(row=i, column=j)))
                myColumn.append(self.wsDecPT.cell(row=i, column=j).value)
            self.centralReg.append(myColumn)
        self.dataDec = pd.DataFrame(self.centralReg)
        # print(self.data)
        # pd.set_option('display.max_columns', None)
        header = ["Nr.", "Oficiul", "PT", "Fider", "Data si ora deconectarii", "Data si ora conectarii", "Durata intreruperii",
                  "Consumatori casnici", "Consumatori non-casnici",
                  "Total",
                  "Localitate", "Cauza deconectarii",
                  "Termenul\nreglementat\n6/12"]
        self.tableDecPt = QTableView()
        self.model = TableModel(self.dataDec, header)

        self.tableDecPt.setModel(self.model)
            # self.table.setWordWrap(True)
        self.tableDecPt.setTextElideMode(Qt.ElideMiddle)
        self.tableDecPt.resizeColumnsToContents()
        # self.table.columnsWidth(9, 300)
        self.tableDecPt.resizeRowsToContents()
        self.tableDecPt.setStyleSheet("Background-color: rgb(200, 200, 200)")
        self.tableDecPt.setSelectionBehavior(QAbstractItemView.SelectRows)
        for i in range(0, 13):
            self.tableDecPt.setColumnWidth(i, 135)
        self.tableDecPt.setColumnWidth(4, 160)
        self.tableDecPt.setColumnWidth(5, 160)
        self.tableDecPt.setColumnWidth(12, 200)
        self.tableDecPt.verticalHeader().hide()

        #Creez a II-lea tabel
        self.centralReg = []
        for i in range(5, 9):
            myColumn = []
            for j in range(1, 12):
                # data = np.array(centralReg.append(self.wsRegAl.cell(row=i, column=j)))
                if isinstance(self.wsNrSol.cell(row=i, column=j).value, int):
                    myColumn.append(str(self.wsNrSol.cell(row=i, column=j).value))
                else:
                    myColumn.append(self.wsNrSol.cell(row=i, column=j).value)
            self.centralReg.append(myColumn)
        self.dataDec = pd.DataFrame(self.centralReg)

        header = ["Nr.", "Oficiul", "Total", "Remediat", "In Executare",
                  "Termen reglementat\n6/12 nerespectat",
                  "Depasire, ore", "Cauza SA RED-Nord",
                  "Cauza IS ME", "Cauza consumator",
                  "Altele"]
        # verHeader = ["UN", "FL", "GL", "RS"]

        self.tableNrSol = QTableView()
        self.model = TableModelII(self.dataDec, header)

        self.tableNrSol.setModel(self.model)
        # self.table.setWordWrap(True)
        self.tableNrSol.setTextElideMode(Qt.ElideMiddle)
        self.tableNrSol.resizeColumnsToContents()
        # self.tableNrSol.setColumnWidth(9, 300)
        self.tableNrSol.resizeRowsToContents()
        self.tableNrSol.setStyleSheet("Background-color: rgb(200, 200, 200)")
        for i in range(0, 11):
            self.tableNrSol.setColumnWidth(i, 140)
        self.tableNrSol.verticalHeader().hide()
        # self.tableNrSol.clicked.connect(self.nrSolChange)

        # Creez a III-lea tabel
        self.centralReg = []
        for i in range(13, 17):
            myColumn = []
            for j in range(1, 12):
                # data = np.array(centralReg.append(self.wsRegAl.cell(row=i, column=j)))
                if isinstance(self.wsNrSol.cell(row=i, column=j).value, int):
                    myColumn.append(str(self.wsNrSol.cell(row=i, column=j).value))
                else:
                    myColumn.append(self.wsNrSol.cell(row=i, column=j).value)
            self.centralReg.append(myColumn)
        self.dataDec = pd.DataFrame(self.centralReg)


        # header = []
        self.tableInf = QTableView()
        self.model = TableModelIII(self.dataDec)

        self.tableInf.setModel(self.model)
        # self.table.setWordWrap(True)
        self.tableInf.setTextElideMode(Qt.ElideMiddle)
        self.tableInf.resizeColumnsToContents()
        # self.tableInf.setColumnWidth(9, 300)
        self.tableInf.resizeRowsToContents()
        self.tableInf.setStyleSheet("Background-color: rgb(200, 200, 200)")
        # self.tableInf.horizontalHeader().hide()
        self.tableInf.verticalHeader().hide()
        for i in range(3, 11):
            self.tableInf.setColumnHidden(i, True)
        self.tableInf.setColumnWidth(2, 1460)


        #I Frame Deconectari PT

        rapDecFrame = QFrame()

        rapDecTitle = QLabel("Raport deconectari zilnice")
        rapDecTitle.setStyleSheet("padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)")

        rapData = QLabel(" " + datetime.datetime.now().strftime("%d.%m.%y"))
        rapData.setStyleSheet("font-size:24px")

        mySpace = QLabel(" ")

        hbox = QHBoxLayout()
        hbox.addWidget(rapDecTitle)
        hbox.addWidget(rapData)
        hbox.addWidget(mySpace)
        hbox.addWidget(mySpace)
        hbox.addWidget(mySpace)
        hbox.addWidget(mySpace)

        decPT = QLabel("Deconectari PT")
        decPT.setStyleSheet("padding-left: 100%; font-size:14px; color:rgb(0, 0, 0)")

        vbox = QVBoxLayout()
        vbox.addLayout(hbox)
        vbox.addWidget(decPT)
        vbox.addWidget(self.tableDecPt)
        # vbox.addWidget(spaceFrame)
        rapDecFrame.setLayout(vbox)

        #Numar de solicitari, informatii suplimentare

        nrSolFrame = QFrame()

        nrSol = QLabel("Numar solicitari")
        nrSol.setStyleSheet("padding-left: 100%; font-size:14px; color:rgb(0, 0, 0)")

        infSupl = QLabel("Informatii suplimentare")
        infSupl.setStyleSheet("padding-left: 100%; font-size:14px; color:rgb(0, 0, 0)")

        vbox = QVBoxLayout()
        vbox.addWidget(nrSol)
        vbox.addWidget(self.tableNrSol)
        vbox.addWidget(infSupl)
        vbox.addWidget(self.tableInf)
        nrSolFrame.setLayout(vbox)

        myScrRes = app.primaryScreen()
        myScrAvailable = myScrRes.availableGeometry()
        myHeight = round(myScrAvailable.height()/4)

        mySplitter = QSplitter(Qt.Vertical)
        mySplitter.addWidget(rapDecFrame)
        mySplitter.addWidget(nrSolFrame)
        mySplitter.setSizes([myHeight, 1])

        mySplitter.handle(1).setStyleSheet("Background-color: rgb(191, 60, 60)")
        # mySplitter.setHandleWidth(10)
        # print(mySplitter.handleWidth())

        self.tabWindDec = QMdiSubWindow()
        self.myMidi.addSubWindow(self.tabWindDec)
        self.tabWindDec.setWindowIcon(QIcon(QPixmap(1, 1)))
        self.tabWindDec.setWidget(mySplitter)
        self.tabWindDec.setGeometry(100, 100, 1000, 600)
        self.tabWindDec.showMaximized()
        self.tabDecControl = True
        self.tabWindDec.show()

    def executat(self):
        if self.data.at[self.modRow, 11] == "Neexecutat":
            pyDateTime = datetime.datetime.now()
            self.data.at[self.modRow, 11] = \
                "Executat: " + self.uCombo.currentText() + "\n" + pyDateTime.strftime("%d.%m.%y")
            self.db.der_app_deranj.update_one({
                "nr_ordine": self.data.at[self.modRow, 1]
            }, {
                "$set": {
                    "starea": self.data.at[self.modRow, 11]
                }
            })
        else:
            self.msSecCall("Deranjamentul a fost remediat!")


    # def intrerupere(self):
    #     # self.tabWindow.close()
    #     # self.centrAlPop()
    #     self.data.at[self.modRow, self.modColumn] = \
    #         "Intrerupere"
    #
    #     # Fac update la Excel dupa introducerea datelor la "Nou inregistrata"
    #     newRegEx = self.wsRegAl.max_row - self.modRow
    #     self.wsRegAl.cell(row=newRegEx, column=16).value = \
    #         self.data.at[self.modRow, self.modColumn]
    #     try:
    #         self.wbRegAl.save(self.regFile)
    #     except PermissionError:
    #         self.msSecCall("Registru AUTORIZATIILOR excel este deschis undeva,\n"
    #                        "(nu este permisa introducerea datelor). Incercati mai tirziu!")
    #         self.centrAlPop()


    # def viberTrig(self):
    #     command = "Viber.exe"
    #     # os.system(command)
    #     subprocess.Popen(command)

    def intLinkTrig(self):
        try:
            try:
                if self.data.at[self.modRow, 1] != "":
                    myList = self.reg_al.find_one({"nr_ds": self.data.at[self.modRow, 1]})["link"].split("/")
                    myOriginal = self.reg_al.find_one({"nr_ds": self.data.at[self.modRow, 1]})["link"]
                if self.data.at[self.modRow, 2] != "":
                    myList = self.reg_al.find_one({"nr_al": self.data.at[self.modRow, 2]})["link"].split("/")
                    myOriginal = self.reg_al.find_one({"nr_al": self.data.at[self.modRow, 2]})["link"]

                #Lucrez cu google Drive API v3
                SCOPES = ['https://www.googleapis.com/auth/drive']
                creds = None
                # The file token.pickle stores the user's access and refresh tokens, and is
                # created automatically when the authorization flow completes for the first
                # time.
                if os.path.exists('google_my/token.pickle'):
                    with open('google_my/token.pickle', 'rb') as token:
                        creds = pickle.load(token)
                # If there are no (valid) credentials available, let the user log in.
                if not creds or not creds.valid:
                    if creds and creds.expired and creds.refresh_token:
                        creds.refresh(Request())
                    else:
                        flow = InstalledAppFlow.from_client_secrets_file(
                            'google_my/credentials.json', SCOPES)
                        creds = flow.run_local_server(port=0)
                    # Save the credentials for the next run
                    with open('google_my/token.pickle', 'wb') as token:
                        pickle.dump(creds, token)

                service = build('drive', 'v3', credentials=creds)

                #Partea care raspunde de upload
                file_metadata = {'name': myList[len(myList)-1], 'parents': ['1TSSHbkN8EVfwKuyJ1RrSucmwrUSoEg46']}
                media = MediaFileUpload(myOriginal)
                file = service.files().create(body=file_metadata,
                                              media_body=media,
                                              fields='id').execute()
                self.data.at[self.modRow, 18] = "https://drive.google.com/file/d/" + \
                                                file.get('id') + "/view?usp=sharing"
                if self.data.at[self.modRow, 1] != "":
                    self.reg_al.update_one({
                        "nr_ds": self.data.at[self.modRow, 1]
                    }, {
                        "$set": {
                            "link": "https://drive.google.com/file/d/" + \
                                                file.get('id') + "/view?usp=sharing"
                        }
                    })
                if self.data.at[self.modRow, 2] != "":
                    self.reg_al.update_one({
                        "nr_al": self.data.at[self.modRow, 2]
                    }, {
                        "$set": {
                            "link": "https://drive.google.com/file/d/" + \
                                    file.get('id') + "/view?usp=sharing"
                        }
                    })
            except FileNotFoundError:
                self.msSecCall("Nu se poate de facut upload pentru ca\n"\
                               "nu exista link valabil pe calculatorul dvs.")
        except (OSError, AttributeError):
            pass

    def linkTrig(self):
        # Dialog Window Create
        self.dialBox = QDialog()
        self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
        self.dialBox.setWindowIcon(QIcon('ataman_logo'))
        self.dialBox.setWindowTitle('Schimba linkul')
        self.dialBox.setStyleSheet('background-color: #424242;')

        setGenFrame = QFrame()
        setGenFrame.setFrameShape(QFrame.StyledPanel)

        # Destinatia registru autorizatii
        setRegFrame = QFrame()
        setRegFrame.setFrameShape(QFrame.StyledPanel)
        setRegFrame.setStyleSheet('background-color: #5c3838')

        setRegLabel = QLabel("Introdu linkul:")
        setRegLabel.setStyleSheet('color: rgb(255, 255, 255)')
        self.linkLine = QLineEdit()
        self.linkLine.setFixedWidth(250)
        self.linkLine.setStyleSheet('background-color: #e3e3e3;')
        rxDt = QRegExp("")
        myDtValidator = QRegExpValidator(rxDt)
        self.linkLine.setValidator(myDtValidator)
        if self.data.at[self.modRow, 1] != "":
            self.linkLine.setText(self.reg_al.find_one({"nr_ds": self.data.at[self.modRow, 1]})["link"])
        if self.data.at[self.modRow, 2] != "":
            self.linkLine.setText(self.reg_al.find_one({"nr_al": self.data.at[self.modRow, 2]})["link"])

        setRegButton = QPushButton()
        setRegButton.setIcon(QIcon("folder.ico"))
        setRegButton.setIconSize(QSize(20, 20))
        setRegButton.clicked.connect(self.newLink)

        hbox = QHBoxLayout()
        hbox.addWidget(setRegLabel)
        hbox.addWidget(self.linkLine)
        hbox.addWidget(setRegButton)
        setRegFrame.setLayout(hbox)

        #Butoane
        setBtFrame = QFrame()
        # setBtFrame.setFrameShape(QFrame.StyledPanel)

        setBtOk = QPushButton('Ok')
        setBtOk.setStyleSheet('color: #e3e3e3')
        # setBtOk.setFixedWidth(100)
        setBtOk.clicked.connect(self.okLink)
        btCancel = QPushButton('Cancel')
        btCancel.setStyleSheet('color: #e3e3e3')
        # btCancel.setFixedWidth(100)
        btCancel.clicked.connect(self.dialBox.close)

        # setEmptyLb = QLabel()

        hbox = QHBoxLayout()
        # hbox.addWidget(setEmptyLb)
        hbox.addWidget(setBtOk)
        hbox.addWidget(btCancel)
        setBtFrame.setLayout(hbox)

        vbox = QVBoxLayout()
        vbox.addWidget(setRegFrame)
        # vbox.addWidget(setSaidiFrame)
        # vbox.addWidget(setRapFrame)
        # vbox.addWidget(setDsFrame)
        vbox.addWidget(setBtFrame)
        setGenFrame.setLayout(vbox)

        hbox = QHBoxLayout()
        hbox.addWidget(setGenFrame)
        self.dialBox.setLayout(hbox)

        self.dialBox.exec()


    def semneaza(self):
        searchInstr = re.search("Instruire:", self.data.at[self.modRow, 13])
        if self.data.at[self.modRow, 13] == "Semnatura":
            pyDateTime = datetime.datetime.now()
            self.data.at[self.modRow, 13] = "Instruire:\n" + self.nameFMenu +\
                "\n" + pyDateTime.strftime("%d.%m.%y")
        elif searchInstr:
            self.data.at[self.modRow, 13] = self.data.at[self.modRow, 13] + "\n" + self.nameFMenu

        if self.data.at[self.modRow, 1] != "":
            self.reg_al.update_one({
                "nr_ds": self.data.at[self.modRow, 1]
            }, {
                "$set": {
                    "semnatura": self.data.at[self.modRow, 13]
                }
            })
        elif self.data.at[self.modRow, 2] != "":
            self.reg_al.update_one({
                "nr_al": self.data.at[self.modRow, 2]
            }, {
                "$set": {
                    "semnatura": self.data.at[self.modRow, 13]
                }
            })


    def semnVaz(self):
        if self.data.at[self.modRow, 10] == "Semnatura":
            pyDateTime = datetime.datetime.now()
            self.data.at[self.modRow, 10] = "Vazut: \n" + self.nameFMenu + \
                                                    "\n" + pyDateTime.strftime("%d.%m.%y")
        else:
            self.data.at[self.modRow, 10] = self.data.at[self.modRow, 10] + "\n" + self.nameFMenu
        self.db.der_app_deranj.update_one({
                "nr_ordine": self.data.at[self.modRow, 1]
            }, {
                "$set": {
                    "responsabil": self.data.at[self.modRow, 10]
                }
            })


    def accFunc(self):
        if self.namePosition == "Dispecer":
            if self.data.at[self.modRow, 14] == "Nou inregistrata":
                myDateTime = datetime.datetime.now()
                self.data.at[self.modRow, 14] = "Acceptat:" + \
                                                "\n" + self.uCombo.currentText() + "\n" + \
                                                myDateTime.strftime("%d.%m.%y")

                if self.data.at[self.modRow, 1] != "":
                    self.reg_al.update_one({
                        "nr_ds": self.data.at[self.modRow, 1]
                    }, {
                        "$set": {
                            "starea": self.data.at[self.modRow, 14]
                        }
                    })
                elif self.data.at[self.modRow, 2] != "":
                    self.reg_al.update_one({
                        "nr_al": self.data.at[self.modRow, 2]
                    }, {
                        "$set": {
                            "starea": self.data.at[self.modRow, 14]
                        }
                    })
            else:
                self.msSecCall("Inregistrarea a fost deja efectuata!")
        else:
            self.msSecCall("Nu aveti suficiente drepturi\npentru "
                           "a efectua inregistrarea!")

    def refFunc(self):
        if self.namePosition == "Dispecer":
            if self.data.at[self.modRow, 14] == "Nou inregistrata":
                self.data.at[self.modRow, 14] = "Refuz:" + \
                                                "\n" + self.uCombo.currentText()

                if self.data.at[self.modRow, 1] != "":
                    self.reg_al.update_one({
                        "nr_ds": self.data.at[self.modRow, 1]
                    }, {
                        "$set": {
                            "starea": self.data.at[self.modRow, 14]
                        }
                    })
                elif self.data.at[self.modRow, 2] != "":
                    self.reg_al.update_one({
                        "nr_al": self.data.at[self.modRow, 2]
                    }, {
                        "$set": {
                            "starea": self.data.at[self.modRow, 14]
                        }
                    })
            else:
                self.msSecCall("Inregistrarea a fost deja efectuata!")
        else:
            self.msSecCall("Nu aveti suficiente drepturi\npentru "
                           "a efectua inregistrarea!")

    def corectFunc(self):
        if self.data.at[self.modRow, 15] == "Pregatire":
            self.decControl = False
            self.deranjControl = False
            # Dialog Window Create
            self.dialBox = QDialog()
            self.dialBox.setWindowFlags(Qt.WindowCloseButtonHint)
            self.dialBox.setWindowIcon(QIcon('ataman_logo'))
            self.dialBox.setWindowTitle('Registru autorizatii - Corectare')
            self.dialBox.setStyleSheet('background-color: #424242;')

            # Oficiul
            ofFrame = QFrame()
            ofFrame.setFrameShape(QFrame.StyledPanel)
            ofFrame.setStyleSheet("background-color: rgb(55, 66, 89); color: #e3e3e3")

            ofLabel = QLabel("Oficiul:")
            ofLabel.setStyleSheet("margin-right: 55%;")
            # Incarc oficiile din MongoDB
            self.loadOficii()
            self.ofCorectCombo = QComboBox()
            self.ofCorectCombo.setStyleSheet('color: #e3e3e3; height:20;')
            self.ofCorectCombo.setFixedWidth(100)
            self.ofCorectCombo.addItems(self.ofListAbr)
            self.ofCorectCombo.setCurrentText(self.data.at[self.modRow, 0])
            self.ofCorectCombo.currentTextChanged.connect(self.loadPtCorect)

            if self.data.at[self.modRow, 1] != "":
                nrDsAl = self.data.at[self.modRow, 1]
                nrLabel = QLabel("Nr. DS: ")
            elif self.data.at[self.modRow, 2] != "":
                nrDsAl = self.data.at[self.modRow, 2]
                nrLabel = QLabel("Nr. AL: ")

            nrLabel.setStyleSheet("margin-left: 38%;")

            self.nrLine = QLineEdit()
            self.nrLine.setText(str(nrDsAl))
            self.nrLine.setStyleSheet('background-color: #ffffff; color: rgb(130, 130, 130);')
            self.nrLine.setFixedWidth(20)
            self.nrLine.setEnabled(False)

            empty = QLabel()
            empty.setStyleSheet("margin-right: 260%;")

            hbox = QHBoxLayout()
            hbox.addWidget(ofLabel)
            hbox.addWidget(self.ofCorectCombo)
            hbox.addWidget(nrLabel)
            hbox.addWidget(self.nrLine)
            hbox.addWidget(empty)
            # hbox.addWidget(empty)
            ofFrame.setLayout(hbox)

            #Instalatia, PT, Fider
            instPtFrame = QFrame()
            instPtFrame.setFrameShape(QFrame.StyledPanel)
            instPtFrame.setStyleSheet("background-color: rgb(55, 66, 89); color: #e3e3e3")

            instLabel = QLabel("Instalatia: ")
            instLabel.setStyleSheet("margin-right:37%")
            for i in self.dbGen.nomenclator.find({"name": "instalatia"}):
                instList = i["lista"]
            self.instLine = QComboBox()
            self.instLine.addItems(instList)
            self.instLine.setEditable(True)
            self.instLine.setStyleSheet('color: #e3e3e3;  height:20')
            self.instLine.setFixedWidth(100)
            self.instLine.setCurrentText(self.data.at[self.modRow, 3])

            ptLabel = QLabel("PT: ")
            ptLabel.setStyleSheet("margin-left:53%")
            self.ptLine = QLineEdit()
            self.ptLine.setStyleSheet('background-color: #ffffff; color: rgb(0, 0, 0)')
            emptyLabel = QLabel("")

            rxPt = QRegExp("(PT|PD)\d.(dot)?.(dot)")
            rxPt.setCaseSensitivity(Qt.CaseInsensitive)
            myValidator = QRegExpValidator(rxPt)
            self.ptLine.setValidator(myValidator)
            self.ptLine.setFixedWidth(100)
            self.ptLine.setText(self.data.at[self.modRow, 4])

            self.ptFidLine = QLineEdit()
            self.ptFidLine.setStyleSheet('background-color: #ffffff; color: rgb(0, 0, 0)')
            self.ptFidLine.setFixedWidth(100)
            self.ptFidLine.setText(self.data.at[self.modRow, 6])

            hbox = QHBoxLayout()
            hbox.addWidget(instLabel)
            hbox.addWidget(self.instLine)
            hbox.addWidget(ptLabel)
            hbox.addWidget(self.ptLine)
            hbox.addWidget(emptyLabel)
            hbox.addWidget(self.ptFidLine)
            instPtFrame.setLayout(hbox)

            # Sef, Membrii, Emitent
            sefEmFrame = QFrame()
            sefEmFrame.setFrameShape(QFrame.StyledPanel)
            sefEmFrame.setStyleSheet("background-color: rgb(55, 66, 89); color: #e3e3e3")

            sefLabel = QLabel("Sef de lucrari: ")
            sefLabel.setStyleSheet("margin-right: 19%")
            memLabel = QLabel("Membrii echipei: ")
            # memLabel.setStyleSheet("margin-right:0,9%")
            emLabel = QLabel("Emitent: ")

            self.sfLine = QLineEdit()
            self.sfLine.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)")
            self.sfLine.setFixedWidth(100)
            self.sfLine.setText(self.data.at[self.modRow, 8])
            self.memEchLine = QLineEdit()
            self.memEchLine.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)")
            self.memEchLine.setFixedWidth(100)
            if self.data.at[self.modRow, 9] != "":
                myList = self.data.at[self.modRow, 9].split(":")
                self.memEchLine.setText("Formatia: " + myList[1])
            self.emLine = QLineEdit()
            self.emLine.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)")
            self.emLine.setFixedWidth(100)

            hbox = QHBoxLayout()
            hbox.addWidget(sefLabel)
            hbox.addWidget(self.sfLine)
            hbox.addWidget(memLabel)
            hbox.addWidget(self.memEchLine)
            hbox.addWidget(emLabel)
            hbox.addWidget(self.emLine)
            sefEmFrame.setLayout(hbox)

            # Ce deconectare, Masurile
            decMasFrame = QFrame()
            decMasFrame.setFrameShape(QFrame.StyledPanel)
            decMasFrame.setStyleSheet("background-color: rgb(55, 66, 89); color: #e3e3e3")

            lucrLabel = QLabel("Lucrarile efectuate:")
            for i in self.dbGen.nomenclator.find({"name": "lucrari"}):
                lucrList = i["lista"]
            lucrList.sort()
            self.lucrLine = QComboBox()
            self.lucrLine.setStyleSheet('color: #e3e3e3;  height:20')
            self.lucrLine.addItems(lucrList)
            self.lucrLine.setEditable(True)
            self.lucrLine.setFixedWidth(100)
            self.lucrLine.setCurrentText(self.data.at[self.modRow, 7])

            cuDeconectare = QLabel("Cu deconectare: ")
            for i in self.dbGen.nomenclator.find({"name": "cu_dec"}):
                cu_decList = i["lista"]
            self.decCombo = QComboBox()
            self.decCombo.setStyleSheet('color: #e3e3e3; height:20')
            self.decCombo.setFixedWidth(100)
            self.decCombo.addItems(cu_decList)
            self.decCombo.setEditable(True)
            self.decCombo.setCurrentText(self.data.at[self.modRow, 11])

            masTeh = QLabel("Masurile tehnice:")
            self.masLine_corect = QLineEdit()
            self.masLine_corect.setStyleSheet("background-color: rgb(255, 255, 255); color: rgb(0, 0, 0)")
            self.masLine_corect.setText(self.data.at[self.modRow, 12])
            self.masLine_corect.setFixedWidth(100)

            hbox = QHBoxLayout()
            hbox.addWidget(lucrLabel)
            hbox.addWidget(self.lucrLine)
            hbox.addWidget(cuDeconectare)
            hbox.addWidget(self.decCombo)
            hbox.addWidget(masTeh)
            hbox.addWidget(self.masLine_corect)
            decMasFrame.setLayout(hbox)

            self.loadPtCorect()
            self.emLine.setText(self.data.at[self.modRow, 10])

            # Buttons Section (butoanele "Ok", "Cancel")
            btFrame = QFrame()
            btFrame.setFrameShape(QFrame.StyledPanel)

            btOk = QPushButton('Ok')
            btOk.setStyleSheet('color: #e3e3e3')
            btOk.clicked.connect(self.corectieFunc)

            btCancel = QPushButton('Cancel')
            btCancel.setStyleSheet('color: #e3e3e3')
            btCancel.clicked.connect(self.dialBox.close)

            hbox = QHBoxLayout()
            hbox.addWidget(btOk)
            hbox.addWidget(btCancel)
            btFrame.setLayout(hbox)


            vbox = QVBoxLayout()
            vbox.addWidget(ofFrame)
            vbox.addWidget(instPtFrame)
            vbox.addWidget(sefEmFrame)
            vbox.addWidget(decMasFrame)
            vbox.addWidget(btFrame)
            self.dialBox.setLayout(vbox)

            self.dialBox.exec()
        else:
            self.msSecCall("Modificari nu mai pot fi efectuate!")

        # Functie sterge rind DS/AL

    def dublFunc(self):
        self.dublControl = True
        if self.data.at[self.modRow, 1] != "":
            self.dsTrig()
        if self.data.at[self.modRow, 2] != "":
            self.alTrig()
        self.dublControl = False



    def stergeFunc(self):
        if self.data.at[self.modRow, 15] == "Pregatire":
            self.resetCall("Sinteti deacord sa stergeti inregistrarea?\n"
                        'Apasati "Cancel" daca nu!')
            if self.secMB.exec() == QMessageBox.Ok:
                self.stergeTrig()
            else:
                self.secMB.close()
        else:
            self.msSecCall("Modificari nu mai pot fi efectuate!")

    def stergeTrig(self):
        # Sterg la MongoDB
        if self.data.at[self.modRow, 1] != "":
            self.reg_al.delete_one({
                "nr_ds": self.data.at[self.modRow, 1]
            })
        elif self.data.at[self.modRow, 2] != "":
            self.reg_al.delete_one({
                "nr_al": self.data.at[self.modRow, 2]
            })
        self.table.hideRow(self.modRow)

    def pregFunc(self):
        if self.namePosition == "Dispecer":
            searchAcc = re.search("Acceptat:", self.data.at[self.modRow, 14])
            if searchAcc:
                if self.data.at[self.modRow, 15] == "Pregatire" \
                    or self.data.at[self.modRow, 15] == "Cerere la pregatire":
                    self.dtRegAl()
                    myMaxRow = self.wsRegAl.max_row + 1
                    myDateTime = datetime.datetime.now()
                    self.data.at[self.modRow, 15] = "Pregatire:\n" + \
                        myDateTime.strftime("%d.%m.%y %H:%M")+ "\n" + self.uCombo.currentText()

                    if self.data.at[self.modRow, 11] == "" or self.data.at[self.modRow, 11] == "Fara deconectari":
                    # if self.data.at[self.modRow, 1] != "":
                        self.data.at[self.modRow, 16] = "Admitere:\n" + \
                        myDateTime.strftime("%d.%m.%y %H:%M")+ "\n" + self.uCombo.currentText()

                    #Fac update la MongoDB, ulterior excel
                    if self.data.at[self.modRow, 1] != "":
                        self.reg_al.update_one({
                            "nr_ds": self.data.at[self.modRow, 1]
                        }, {
                            "$set": {
                                "pregatire": self.data.at[self.modRow, 15],
                                "admitere": self.data.at[self.modRow, 16]
                            }
                        })
                        # Fac update la excel registru
                        for i in self.reg_al.find({"nr_ds": self.data.at[self.modRow, 1]}):
                            self.wsRegAl.cell(row=myMaxRow, column=1).value = i["oficiul"]
                            self.wsRegAl.cell(row=myMaxRow, column=1).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=2).value = i["nr_ds"]
                            self.wsRegAl.cell(row=myMaxRow, column=2).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=3).value = i["nr_al"]
                            self.wsRegAl.cell(row=myMaxRow, column=3).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=4).value = i["instalatia"]
                            self.wsRegAl.cell(row=myMaxRow, column=4).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=5).value = i["pt"]
                            self.wsRegAl.cell(row=myMaxRow, column=5).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=6).value = i["localitatea"]
                            self.wsRegAl.cell(row=myMaxRow, column=6).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=7).value = i["fid_nr"]
                            self.wsRegAl.cell(row=myMaxRow, column=7).alignment = \
                                 Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=8).value = i["lucrarile"]
                            self.wsRegAl.cell(row=myMaxRow, column=8).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=9).value = i["sef"]
                            self.wsRegAl.cell(row=myMaxRow, column=9).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=10).value = i["mem_ech"]
                            self.wsRegAl.cell(row=myMaxRow, column=10).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=11).value = i["emitent"]
                            self.wsRegAl.cell(row=myMaxRow, column=11).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=12).value = i["cu_dec"]
                            self.wsRegAl.cell(row=myMaxRow, column=12).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=13).value = i["mas_teh"]
                            self.wsRegAl.cell(row=myMaxRow, column=13).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=14).value = i["semnatura"]
                            self.wsRegAl.cell(row=myMaxRow, column=14).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=15).value = i["starea"]
                            self.wsRegAl.cell(row=myMaxRow, column=15).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=16).value = i["pregatire"]
                            self.wsRegAl.cell(row=myMaxRow, column=16).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=17).value = i["admitere"]
                            self.wsRegAl.cell(row=myMaxRow, column=17).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=18).value = i["terminare"]
                            self.wsRegAl.cell(row=myMaxRow, column=18).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)

                    elif self.data.at[self.modRow, 2] != "":
                        self.reg_al.update_one({
                            "nr_al": self.data.at[self.modRow, 2]
                        }, {
                            "$set": {
                                "pregatire": self.data.at[self.modRow, 15],
                                "admitere": self.data.at[self.modRow, 16]
                            }
                        })
                        # Fac update la excel registru
                        for i in self.reg_al.find({"nr_al": self.data.at[self.modRow, 2]}):
                            self.wsRegAl.cell(row=myMaxRow, column=1).value = i["oficiul"]
                            self.wsRegAl.cell(row=myMaxRow, column=1).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=2).value = i["nr_ds"]
                            self.wsRegAl.cell(row=myMaxRow, column=2).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=3).value = i["nr_al"]
                            self.wsRegAl.cell(row=myMaxRow, column=3).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=4).value = i["instalatia"]
                            self.wsRegAl.cell(row=myMaxRow, column=4).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=5).value = i["pt"]
                            self.wsRegAl.cell(row=myMaxRow, column=5).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=6).value = i["localitatea"]
                            self.wsRegAl.cell(row=myMaxRow, column=6).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=7).value = i["fid_nr"]
                            self.wsRegAl.cell(row=myMaxRow, column=7).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=8).value = i["lucrarile"]
                            self.wsRegAl.cell(row=myMaxRow, column=8).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=9).value = i["sef"]
                            self.wsRegAl.cell(row=myMaxRow, column=9).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=10).value = i["mem_ech"]
                            self.wsRegAl.cell(row=myMaxRow, column=10).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=11).value = i["emitent"]
                            self.wsRegAl.cell(row=myMaxRow, column=11).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=12).value = i["cu_dec"]
                            self.wsRegAl.cell(row=myMaxRow, column=12).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=13).value = i["mas_teh"]
                            self.wsRegAl.cell(row=myMaxRow, column=13).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=14).value = i["semnatura"]
                            self.wsRegAl.cell(row=myMaxRow, column=14).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=15).value = i["starea"]
                            self.wsRegAl.cell(row=myMaxRow, column=15).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=16).value = i["pregatire"]
                            self.wsRegAl.cell(row=myMaxRow, column=16).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=17).value = i["admitere"]
                            self.wsRegAl.cell(row=myMaxRow, column=17).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)
                            self.wsRegAl.cell(row=myMaxRow, column=18).value = i["terminare"]
                            self.wsRegAl.cell(row=myMaxRow, column=18).alignment = \
                                Alignment(horizontal="center", vertical="center", wrap_text=True)

                    try:
                        self.wbRegAl.save(self.regFile)
                    except PermissionError:
                        self.msSecCall("Registru AUTORIZATIILOR excel este deschis undeva,\n"
                                        "(nu este permisa introducerea datelor). Incercati mai tirziu!")
                        self.centrAlPop()
                else:
                    self.msSecCall("Permiterea la pregatirea locului de munca\n"
                                   "a fost efectuata sau nu s-a lucrat!")
            else:
                self.msSecCall("Permiterea la pregatirea locului de munca\n"
                               "nu poate fi efectuata fara a accepta inregistrarea!")
        else:
            self.msSecCall("Nu aveti suficiente drepturi\n"
                       "pentru a permite pregatirea locului de munca!")

    def admFunc(self):
        if self.namePosition == "Dispecer":
            searchPreg = re.search("Pregatire:", self.data.at[self.modRow, 15])
            if self.data.at[self.modRow, 16] == "Admitere" \
                    or self.data.at[self.modRow, 16] == "Cerere la admitere":
                if self.data.at[self.modRow, 11] == "" or self.data.at[self.modRow, 11] == "Fara deconectari":
                    self.pregFunc()
                elif searchPreg:
                    self.dtRegAl()
                    # myMaxRow = self.wsRegAl.max_row + 1
                    myDateTime = datetime.datetime.now()
                    self.data.at[self.modRow, 16] = "Admitere:\n" + \
                        myDateTime.strftime("%d.%m.%y %H:%M")+ "\n" + self.uCombo.currentText()

                    # Fac update la MongoDB dupa introducerea datelor la "Nou inregistrata"
                    self.reg_al.update_one({
                        "nr_al": self.data.at[self.modRow, 2]
                    }, {
                        "$set": {
                            "admitere": self.data.at[self.modRow, 16]
                         }
                    })
                    #Fac update la registru excel
                    for i in range(5, self.wsRegAl.max_row + 1):
                        if self.wsRegAl.cell(row=i, column=3).value == self.data.at[self.modRow, 2]:
                            for j in self.reg_al.find({"nr_al": self.data.at[self.modRow, 2]}):
                                self.wsRegAl.cell(row=i, column=17).value = j["admitere"]
                                self.wsRegAl.cell(row=i, column=17).alignment = \
                                    Alignment(horizontal="center", vertical="center", wrap_text=True)

                    try:
                        self.wbRegAl.save(self.regFile)
                    except PermissionError:
                        self.msSecCall("Registru AUTORIZATIILOR excel este deschis undeva,\n"
                                        "(nu este permisa introducerea datelor). Incercati mai tirziu!")
                        self.centrAlPop()
                else:
                    self.msSecCall("Admiterea echipei nu poate fi efectuat fara\n"
                                   "permisiunea la pregatirea locului de munca!")
            else:
                self.msSecCall("Admiterea echipei deja a fost efectuata sau exista\n"
                               "alta cauza pentru care admiterea nu poate fi efectuata!")
        else:
            self.msSecCall("Nu aveti suficiente drepturi\n"
                            "pentru a efectua admiterea echipei!")

    # Functie chemata din context menu pentru neefectuarea lucrarilor
    def nuLucr(self):
        if self.namePosition == "Dispecer":
            searchAcc = re.search("Acceptat:", self.data.at[self.modRow, 14])
            searchPreg = re.search("Pregatire:", self.data.at[self.modRow, 15])
            if not searchPreg:
                if searchAcc:
                    self.data.at[self.modRow, 15] = "Nu s-a lucrat"
                    self.data.at[self.modRow, 16] = "Nu s-a lucrat"
                    self.data.at[self.modRow, 17] = "Nu s-a lucrat"

                    # Fac update la MongoDB dupa introducerea datelor la "Nou inregistrata"
                    if self.data.at[self.modRow, 1] != "":
                        self.reg_al.update_one({
                            "nr_ds": self.data.at[self.modRow, 1]
                        }, {
                            "$set": {
                                "pregatire": self.data.at[self.modRow, 15],
                                "admitere": self.data.at[self.modRow, 16],
                                "terminare": self.data.at[self.modRow, 17]
                            }
                        })
                    elif self.data.at[self.modRow, 2] != "":
                        self.reg_al.update_one({
                            "nr_al": self.data.at[self.modRow, 2]
                        }, {
                            "$set": {
                                "pregatire": self.data.at[self.modRow, 15],
                                "admitere": self.data.at[self.modRow, 16],
                                "terminare": self.data.at[self.modRow, 17]

                            }
                        })
                else:
                    self.msSecCall("Acceptati inregistrarea, dupa care,\n"
                                    "efectuati aceasta operatiune!")
            else:
                self.msSecCall("Aceasta operatiune nu poate fi\n"
                               "efectuata dupa permiterea la pregatire!")
        else:
            self.msSecCall("Nu aveti suficiente drepturi\n"
                            "pentru a efectua aceasta operatiune!")




    #Functie terminarea lucrarilor
    def termLucr(self):
        if self.namePosition == "Dispecer":
            searchAdm = re.search("Admitere:", self.data.at[self.modRow, 16])
            if searchAdm:
                if self.data.at[self.modRow, 17] == "Terminare" \
                        or self.data.at[self.modRow, 17] == "Cerere la terminare":
                    self.dtRegAl()
                    # myMaxRow = self.wsRegAl.max_row
                    # Controlez daca exista mapa cu anul
                    self.dtContrSaidi()
                    # self.dtAnAnual()
                    self.loadPtSec()
                    myDateTime = datetime.datetime.now()
                    self.data.at[self.modRow, 17] = "Terminat:\n" + \
                        myDateTime.strftime("%d.%m.%y %H:%M")+ "\n" + self.uCombo.currentText() + "\n--"
                    self.data.at[self.modRow, 16] = self.data.at[self.modRow, 16] +\
                        "\n--"
                    self.data.at[self.modRow, 15] = self.data.at[self.modRow, 15] + \
                                                                    "\n--"
                    valuePreg = self.data.at[self.modRow, 15]
                    valueTerm = self.data.at[self.modRow, 17]
                    # valueMinList = []
                    valuePregList = valuePreg.split(sep="\n")
                    valueTermList = valueTerm.split(sep="\n")
                    # print(valuePregList[1])
                    # strToDatePreg = datetime.datetime.strptime(valueMinList[1], "%d.%m.%y %H:%M")
                    # print(strToDate)
                    # self.data.at[self.modRow, self.modColumn - 1] = \
                    #     strToDate.strftime("%d.%m.%y %H:%M")

                    # Fac update la MongoDB, ulterior excel
                    if self.data.at[self.modRow, 1] != "":
                        self.reg_al.update_one({
                            "nr_ds": self.data.at[self.modRow, 1]
                        }, {
                            "$set": {
                                "terminare": self.data.at[self.modRow, 17],
                                "pregatire": self.data.at[self.modRow, 15],
                                "admitere": self.data.at[self.modRow, 16]
                            }
                        })
                        # Fac update la registru excel
                        for i in range(5, self.wsRegAl.max_row + 1):
                            if self.wsRegAl.cell(row=i, column=2).value == self.data.at[self.modRow, 1]:
                                for j in self.reg_al.find({"nr_ds": self.data.at[self.modRow, 1]}):
                                    self.wsRegAl.cell(row=i, column=18).value = j["terminare"]
                                    self.wsRegAl.cell(row=i, column=18).alignment = \
                                        Alignment(horizontal="center", vertical="center", wrap_text=True)
                    elif self.data.at[self.modRow, 2] != "":
                        self.reg_al.update_one({
                            "nr_al": self.data.at[self.modRow, 2]
                        }, {
                            "$set": {
                                "terminare": self.data.at[self.modRow, 17],
                                "pregatire": self.data.at[self.modRow, 15],
                                "admitere": self.data.at[self.modRow, 16]
                            }
                        })
                        # Fac update la registru excel
                        for i in range(5, self.wsRegAl.max_row + 1):
                            if self.wsRegAl.cell(row=i, column=3).value == self.data.at[self.modRow, 2]:
                                for j in self.reg_al.find({"nr_al": self.data.at[self.modRow, 2]}):
                                    self.wsRegAl.cell(row=i, column=18).value = j["terminare"]
                                    self.wsRegAl.cell(row=i, column=18).alignment = \
                                        Alignment(horizontal="center", vertical="center", wrap_text=True)
                    try:
                        self.wbRegAl.save(self.regFile)
                    except PermissionError:
                        self.msSecCall("Registru AUTORIZATIILOR excel este deschis undeva,\n"
                                       "(nu este permisa introducerea datelor). Incercati mai tirziu!")
                        self.centrAlPop()
                        self.erContrAl = True
                    # Introduc datele in Excel deconectari

                    if self.data.at[self.modRow, 11] == "Programat" and not self.erContrAl:
                        myMaxRow = self.wsDecProg.max_row + 1
                        self.wsDecProg.cell(row=myMaxRow, column=1).value = \
                            self.wsDecProg.cell(row=myMaxRow - 1, column=1).value + 1
                        self.wsDecProg.cell(row=myMaxRow, column=1).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=2).value = \
                            self.data.at[self.modRow, 0]
                        self.wsDecProg.cell(row=myMaxRow, column=2).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=3).value = \
                            datetime.date.today().strftime("%d.%m.%y")
                        self.wsDecProg.cell(row=myMaxRow, column=3).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=4).value = \
                            "JT"
                        self.wsDecProg.cell(row=myMaxRow, column=4).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=5).value = \
                            self.data.at[self.modRow, 4]
                        self.wsDecProg.cell(row=myMaxRow, column=5).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=6).value = \
                            self.data.at[self.modRow, 6]
                        self.wsDecProg.cell(row=myMaxRow, column=6).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=7).value = \
                            0
                        self.wsDecProg.cell(row=myMaxRow, column=7).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=8).value = \
                            1
                        self.wsDecProg.cell(row=myMaxRow, column=8).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=9).value = \
                            1
                        self.wsDecProg.cell(row=myMaxRow, column=9).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=10).value = \
                            valuePregList[1]
                        self.wsDecProg.cell(row=myMaxRow, column=10).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecProg.cell(row=myMaxRow, column=11).value = \
                            valueTermList[1]
                        self.wsDecProg.cell(row=myMaxRow, column=11).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        #Calculez si pun orele in coloana 12
                        valueMaxDate_11 = self.wsDecProg.cell(row=myMaxRow, column=11).value
                        valueMaxDate_10 = self.wsDecProg.cell(row=myMaxRow, column=10).value
                        strToDate_11 = datetime.datetime.strptime(valueMaxDate_11, "%d.%m.%y %H:%M")
                        strToDate_10 = datetime.datetime.strptime(valueMaxDate_10, "%d.%m.%y %H:%M")
                        delta_11_10 = strToDate_11 - strToDate_10
                        myList = str(delta_11_10).split()
                        try:
                            if int(myList[0]):
                                myDateSplit = myList[2].split(":")
                                myDayPlus = 24 * int(myList[0]) + int(myDateSplit[0])
                                myDateStr = str(myDayPlus) + ":" + myDateSplit[1] + ":" + myDateSplit[2]
                                myDeltaHour = myDateStr
                        except ValueError:
                            myDeltaHour = str(delta_11_10)
                        self.wsDecProg.cell(row=myMaxRow, column=12).value = myDeltaHour
                        self.wsDecProg.cell(row=myMaxRow, column=12).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        #Calculez si pun numarul de consumatori si localitatea
                        for i in range(2, self.wsPt.max_row + 1):
                            if self.wsPt.cell(row=i, column=1).value == self.data.at[self.modRow, 4]:
                                self.totNrCas = int(self.wsPt.cell(row=i, column=4).value)
                                self.fidNrCas = round(self.totNrCas/3)
                                if self.fidNrCas > 65:
                                    self.fidNrCas = random.randrange(60, 70)
                                myComma = re.search(",", self.data.at[self.modRow, 6])
                                if myComma:
                                    myNumber = self.data.at[self.modRow, 6].count(",")
                                    self.fidNrCas = (myNumber + 1) * self.fidNrCas
                                self.wsDecProg.cell(row=myMaxRow, column=13).value = self.fidNrCas
                                self.wsDecProg.cell(row=myMaxRow, column=13).alignment = \
                                    Alignment(horizontal="center", vertical="center")

                                self.totNrEc = self.wsPt.cell(row=i, column=5).value
                                if self.totNrEc <= 2 and self.totNrCas == 0:
                                    self.fidNrEc = 1
                                elif self.totNrEc > 12:
                                    self.fidNrEc = random.randrange(2, 5)
                                else:
                                    self.fidNrEc = round(self.totNrEc/3)
                                self.wsDecProg.cell(row=myMaxRow, column=14).value = self.fidNrEc
                                self.wsDecProg.cell(row=myMaxRow, column=14).alignment = \
                                    Alignment(horizontal="center", vertical="center")
                                self.wsDecProg.cell(row=myMaxRow, column=16).value = self.wsPt.cell(row=i, column=2).value

                        self.wsDecProg.cell(row=myMaxRow, column=15).value = \
                            str(self.wsDecProg.cell(row=myMaxRow, column=13).value + \
                            self.wsDecProg.cell(row=myMaxRow, column=14).value)
                        self.wsDecProg.cell(row=myMaxRow, column=15).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        self.wsDecProg.cell(row=myMaxRow, column=19).value = \
                            self.data.at[self.modRow, 7]

                        #Introduc datele in Excel analiza anuala
                        anAnualMaxRow = self.wsAnAnualP.max_row + 1
                        self.wsAnAnualP.cell(row=anAnualMaxRow, column=1).value = \
                            self.data.at[self.modRow, 0]
                        self.wsAnAnualP.cell(row=anAnualMaxRow, column=2).value = \
                            self.data.at[self.modRow, 4] + " " + self.data.at[self.modRow, 6]
                        self.wsAnAnualP.cell(row=anAnualMaxRow, column=3).value = \
                            self.data.at[self.modRow, 5]
                        self.wsAnAnualP.cell(row=anAnualMaxRow, column=4).value = \
                            self.fidNrCas + self.fidNrEc
                        self.wsAnAnualP.cell(row=anAnualMaxRow, column=5).value = \
                            myDeltaHour
                        self.wsAnAnualP.cell(row=anAnualMaxRow, column=6).value = \
                            1
                        try:
                            self.wbAnAnual.save(self.fileAnAnual)
                        except PermissionError:
                            self.msSecCall("Datele din autorizatie, sectiunea PROGRAMAT \n"
                                           "nu vor participa la analiza anuala (undeva este deschisa analiza anuala excel)!")
                        try:
                            self.wbDec.save(self.saidiFile)
                        except PermissionError:
                            self.msSecCall("Datele din autorizatie, sectiunea PROGRAMAT \n"
                                           "nu s-au introdus in fisierul SAIDI excel (cineva foloseste aplicatia)!")

                    if self.data.at[self.modRow, 11] == "Neprogramat" and not self.erContrAl:
                        myMaxRow = self.wsDecNeProg.max_row + 1
                        self.wsDecNeProg.cell(row=myMaxRow, column=1).value = \
                            self.wsDecNeProg.cell(row=myMaxRow - 1, column=1).value + 1
                        self.wsDecNeProg.cell(row=myMaxRow, column=1).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=2).value = \
                            self.data.at[self.modRow, 0]
                        self.wsDecNeProg.cell(row=myMaxRow, column=2).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=3).value = \
                            datetime.date.today().strftime("%d.%m.%y")
                        self.wsDecNeProg.cell(row=myMaxRow, column=3).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=4).value = \
                            self.data.at[self.modRow, 4]
                        self.wsDecNeProg.cell(row=myMaxRow, column=4).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=5).value = \
                            self.data.at[self.modRow, 6]
                        self.wsDecNeProg.cell(row=myMaxRow, column=5).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=6).value = \
                            0
                        self.wsDecNeProg.cell(row=myMaxRow, column=6).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=7).value = \
                            0
                        self.wsDecNeProg.cell(row=myMaxRow, column=7).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=8).value = \
                            1
                        self.wsDecNeProg.cell(row=myMaxRow, column=8).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=9).value = \
                            1
                        self.wsDecNeProg.cell(row=myMaxRow, column=9).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=10).value = \
                            valuePregList[1]
                        self.wsDecNeProg.cell(row=myMaxRow, column=10).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecNeProg.cell(row=myMaxRow, column=11).value = \
                            valueTermList[1]
                        self.wsDecNeProg.cell(row=myMaxRow, column=11).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        # Calculez si pun orele diferenta lor
                        valueMaxDate_11 = self.wsDecNeProg.cell(row=myMaxRow, column=11).value
                        valueMaxDate_10 = self.wsDecNeProg.cell(row=myMaxRow, column=10).value
                        strToDate_11 = datetime.datetime.strptime(valueMaxDate_11, "%d.%m.%y %H:%M")
                        strToDate_10 = datetime.datetime.strptime(valueMaxDate_10, "%d.%m.%y %H:%M")
                        delta_11_10 = strToDate_11 - strToDate_10
                        myList = str(delta_11_10).split()
                        try:
                            if int(myList[0]):
                                myDateSplit = myList[2].split(":")
                                myDayPlus = 24 * int(myList[0]) + int(myDateSplit[0])
                                myDateStr = str(myDayPlus) + ":" + myDateSplit[1] + ":" + myDateSplit[2]
                                myDeltaHour = myDateStr
                        except ValueError:
                            myDeltaHour = str(delta_11_10)
                        self.wsDecNeProg.cell(row=myMaxRow, column=19).value = myDeltaHour
                        self.wsDecNeProg.cell(row=myMaxRow, column=19).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        # Calculez si pun numarul de consumatori si localitatea
                        for i in range(2, self.wsPt.max_row + 1):
                            if self.wsPt.cell(row=i, column=1).value == self.data.at[self.modRow, 4]:
                                self.totNrCas = int(self.wsPt.cell(row=i, column=4).value)
                                self.fidNrCas = round(self.totNrCas / 3)
                                if self.fidNrCas > 65:
                                    self.fidNrCas = random.randrange(60, 70)
                                myComma = re.search(",", self.data.at[self.modRow, 6])
                                if myComma:
                                    myNumber = self.data.at[self.modRow, 6].count(",")
                                    self.fidNrCas = (myNumber + 1) * self.fidNrCas
                                self.wsDecNeProg.cell(row=myMaxRow, column=25).value = self.fidNrCas
                                self.wsDecNeProg.cell(row=myMaxRow, column=25).alignment = \
                                    Alignment(horizontal="center", vertical="center")

                                self.totNrEc = self.wsPt.cell(row=i, column=5).value
                                if self.totNrEc <= 2 and self.totNrCas == 0:
                                    self.fidNrEc = 1
                                elif self.totNrEc > 12:
                                    self.fidNrEc = random.randrange(2, 5)
                                else:
                                    self.fidNrEc = round(self.totNrEc/3)
                                self.wsDecNeProg.cell(row=myMaxRow, column=26).value = self.fidNrEc
                                self.wsDecNeProg.cell(row=myMaxRow, column=26).alignment = \
                                    Alignment(horizontal="center", vertical="center")
                                self.wsDecNeProg.cell(row=myMaxRow, column=27).value = \
                                    self.wsDecNeProg.cell(row=myMaxRow, column=25).value + \
                                    self.wsDecNeProg.cell(row=myMaxRow, column=26).value
                                self.wsDecNeProg.cell(row=myMaxRow, column=27).alignment = \
                                    Alignment(horizontal="center", vertical="center")
                                self.wsDecNeProg.cell(row=myMaxRow, column=28).value = self.wsPt.cell(row=i, column=2).value
                                myLocalitate = self.wsPt.cell(row=i, column=2).value

                        self.wsDecNeProg.cell(row=myMaxRow, column=32).value = \
                            self.data.at[self.modRow, 7]
                        try:
                            self.wbDec.save(self.saidiFile)
                        except PermissionError:
                            self.msSecCall("Datele din autorizatie, sectiunea NEPROGRAMAT \n"
                                           "nu s-au introdus in fisierul SAIDI excel (cineva foloseste aplicatia)!")
                        self.dtContrDecZl()
                        myMaxRow = self.wsDecPT.max_row + 1

                        # if self.wsDecPT.cell(row=2, column=8).value == None:
                        #     self.wsDecPT.cell(row=2, column=8).value = datetime.datetime.now().strftime("%d.%m.%Y")
                        self.wsDecPT.cell(row=myMaxRow, column=1).value = \
                            str(int(self.wsDecPT.cell(row=myMaxRow - 1, column=1).value) + 1)
                        self.wsDecPT.cell(row=myMaxRow, column=1).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecPT.cell(row=myMaxRow, column=2).value = \
                            self.data.at[self.modRow, 0]
                        self.wsDecPT.cell(row=myMaxRow, column=2).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecPT.cell(row=myMaxRow, column=3).value = \
                            self.data.at[self.modRow, 4]
                        self.wsDecPT.cell(row=myMaxRow, column=4).value = \
                            self.data.at[self.modRow, 6]
                        self.wsDecPT.cell(row=myMaxRow, column=4).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecPT.cell(row=myMaxRow, column=5).value = \
                            valuePregList[1]
                        self.wsDecPT.cell(row=myMaxRow, column=5).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecPT.cell(row=myMaxRow, column=6).value = \
                            valueTermList[1]
                        self.wsDecPT.cell(row=myMaxRow, column=6).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        # Calculez si pun orele diferenta lor
                        # valueDate_6 = self.wsDecPT.cell(row=myMaxRow, column=6).value
                        # valueDate_5 = self.wsDecPT.cell(row=myMaxRow, column=5).value
                        # strToDate_6 = datetime.datetime.strptime(valueDate_6, "%d.%m.%y %H:%M")
                        # strToDate_5 = datetime.datetime.strptime(valueDate_5, "%d.%m.%y %H:%M")
                        # delta_6_5 = valueTermList[1] - valuePregList[1]
                        self.wsDecPT.cell(row=myMaxRow, column=7).value = myDeltaHour
                        self.wsDecPT.cell(row=myMaxRow, column=7).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        # # Calculez si pun numarul de consumatori si localitatea
                        # for i in range(2, self.wsPt.max_row + 1):
                        #     if self.wsPt.cell(row=i, column=1).value == self.wsDecPT.cell(row=myMaxRow, column=3).value:
                        #         self.totNrCas = int(self.wsPt.cell(row=i, column=4).value)
                        #         self.fazaNrCas = round(self.totNrCas / 9)
                        #         if self.fazaNrCas > 25:
                        #             self.fazaNrCas = random.randrange(20, 30)
                        self.wsDecPT.cell(row=myMaxRow, column=8).value = str(self.fidNrCas)
                        self.wsDecPT.cell(row=myMaxRow, column=8).alignment = \
                                Alignment(horizontal="center", vertical="center")

                                # self.totNrEc = self.wsPt.cell(row=i, column=5).value
                                # if self.totNrEc <= 2 and self.totNrCas == 0:
                                #     self.fazaNrEc = 1
                                # elif self.totNrEc > 12:
                                #     self.fazaNrEc = random.randrange(1, 3)
                                # else:
                                #     self.fazaNrEc = round(self.totNrEc / 9)
                        self.wsDecPT.cell(row=myMaxRow, column=9).value = str(self.fidNrEc)
                        self.wsDecPT.cell(row=myMaxRow, column=9).alignment = \
                                Alignment(horizontal="center", vertical="center")
                        self.wsDecPT.cell(row=myMaxRow, column=11).value = \
                                myLocalitate

                        self.wsDecPT.cell(row=myMaxRow, column=10).value = \
                            str(int(self.wsDecPT.cell(row=myMaxRow, column=8).value) + \
                                int(self.wsDecPT.cell(row=myMaxRow, column=9).value))
                        self.wsDecPT.cell(row=myMaxRow, column=10).alignment = \
                            Alignment(horizontal="center", vertical="center")
                        self.wsDecPT.cell(row=myMaxRow, column=12).value = self.data.at[self.modRow, 7]
                        self.wsDecPT.cell(row=myMaxRow, column=13).value = "Incadrat"
                        self.wsDecPT.cell(row=myMaxRow, column=13).alignment = \
                            Alignment(horizontal="center", vertical="center")

                        try:
                            self.wbDecZil.save(self.rapFile)
                        except PermissionError:
                            self.msSecCall("Datele din autorizatie, sectiunea NEPROGRAMAT \n"
                                           "nu s-au introdus in fisierul RAPORT PDJT excel (cineva foloseste aplicatia)!")

                        # Introduc datele in Excel analiza anuala
                        self.postgresLoad()
                        self.cur.execute(f"""INSERT INTO anlzan21n (
                                                oficiul,
                                                pt_fider,
                                                localitate,
                                                nr_cons,
                                                ore,
                                                nr_dec
                                            ) 
                                            VALUES (
                                                '{self.data.at[self.modRow, 0]}',
                                                '{self.data.at[self.modRow, 4] + " " +self.data.at[self.modRow, 6]}',
                                                '{self.data.at[self.modRow, 5]}',
                                                '{self.fidNrCas + self.fidNrEc}',
                                                '{myDeltaHour}',
                                                '{1}'
                                            )"""
                        )
                        self.conn.commit()
                        self.cur.close()

                        # Determin incadrarea termenului urban, rural
                        myList = myDeltaHour.split(":")
                        myMinute = int(myList[0]) * 60 + int(myList[1])
                        termText = "Incadrat"
                        compens = 0
                        if myLocalitate == None or bool(re.search("or[.]", myLocalitate)):
                            if myMinute > 6 * 60:
                                difMinute = myMinute - 6 * 60
                                termText = "Depasit cu: " + str(round(difMinute/60)) + "H " + myList[1] + "min."
                                if difMinute / 60 <= 3:
                                    k = 1
                                elif (difMinute / 60 > 3) and (difMinute / 60 <= 6):
                                    k = 4
                                elif (difMinute / 60 > 6) and (difMinute / 60 <= 9):
                                    k = 7
                                elif (difMinute / 60 > 9):
                                    k = 10
                                compens = round(0.01 * 160 * 2.04 * k, 2)

                        else:
                            if myMinute > 12 * 60:
                                difMinute = myMinute - 12 * 60
                                termText = "Depasit cu: " + str(round(difMinute/60)) + "H " + myList[1] + "min."
                                if difMinute / 60 <= 3:
                                    k = 1
                                elif (difMinute / 60 > 3) and (difMinute / 60 <= 6):
                                    k = 4
                                elif (difMinute / 60 > 6) and (difMinute / 60 <= 9):
                                    k = 7
                                elif (difMinute / 60 > 9):
                                    k = 10
                                compens = round(0.01 * 160 * 2.04 * k, 2)


                        #Introduc datele in MongoDB analiza lunara
                        for i in self.db.deconect_app_deconect.find().sort("_id", -1).limit(1):
                            self.nrDec = int(i["id"]) + 1

                        self.db.deconect_app_deconect.insert_one({
                            "id": self.nrDec,
                            "oficiul": self.data.at[self.modRow, 0],
                            "nr_ordine": self.nrDec,
                            "pt": self.data.at[self.modRow, 4],
                            "fid_04kv": self.data.at[self.modRow, 6],
                            "data_dec": valuePregList[1],
                            "data_conect": datetime.datetime.now().strftime("%d.%m.%y %H:%M"),
                            "durata": myDeltaHour,
                            "cons_cas": self.fidNrCas,
                            "cons_ec": self.fidNrEc,
                            "total": self.fidNrCas + self.fidNrEc,
                            "localitate": self.data.at[self.modRow, 5],
                            "cauza": self.data.at[self.modRow, 7],
                            "termen": termText,
                            "compens": compens
                        })

                        # anAnualMaxRow = self.wsAnAnualN.max_row + 1
                        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=1).value = \
                        #     self.data.at[self.modRow, 0]
                        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=2).value = \
                        #     self.data.at[self.modRow, 4] + " " +self.data.at[self.modRow, 6]
                        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=3).value = \
                        #     self.data.at[self.modRow, 5]
                        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=4).value = \
                        #     self.fidNrCas + self.fidNrEc
                        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=5).value = \
                        #     myDeltaHour
                        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=6).value = \
                        #     1
                        # try:
                        #     self.wbAnAnual.save(self.fileAnAnual)
                        # except PermissionError:
                        #     self.msSecCall("Datele din autorizatie, sectiunea NEPROGRAMAT \n"
                        #                    "nu vor participa la analiza anuala (undeva este deschisa analiza anuala excel)!")
                    self.erContrAl = False

                else:
                    self.msSecCall("Pe aceasta DS/AL lucrarile deja s-au terminat!")
            else:
                self.msSecCall("Terminarea lucrarilor nu poate fi efectuata\n"
                               "deoarece nu s-a facut admiterea echipei!")
        else:
            self.msSecCall("Nu aveti suficiente drepturi\n"
                            "pentru a permite terminarea lucrarilor!")

    #Functie pentru context menu, admite echipa, folosesc try...except
    #ca sa depasesc eroarea cind nu-i tabelul
    def contextMenuEvent(self, event):
        try:
            if self.table.hasFocus():
                self.modRow = self.table.currentIndex().row()
                self.modColumn = self.table.currentIndex().column()
                ctxMenu = QMenu(self)
                refreshAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/refresh.png"), \
                                              "Refresh")
                ctxMenu.addSeparator()
                accAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/accept.png"),\
                                        "Accepta inregistrarea")
                refAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/refuz.png"),\
                                        "Refuza inregistrarea")
                ctxMenu.addSeparator()
                corectAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/corectare.png"), \
                                                 "Corecteaza DS/AL")
                dublAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/copy.png"), \
                                                 "Dubleaza DS/AL")
                erAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/sterge.png"),\
                                        "Sterge obtiune")
                ctxMenu.addSeparator()
                pregAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/pregatire.png"),\
                                        "Pregatirea locului de munca")
                admAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/admitere.png"),\
                                        "Admiterea echipei")
                termAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/terminat.png"),\
                                        "Lucrarile s-au terminat")
                nuLucrAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/nu_lucr.png"),\
                                        "Nu s-a lucrat")
                ctxMenu.addSeparator()
                semnAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/semneaza.png"),\
                                        "Semneaza instructor/instructat")
                ctxMenu.addSeparator()
                linkAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/link.png"), \
                                               "Linkul DS/AL")
                chLinkAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/schimba_link.png"), \
                                               "Schimba linkul DS/AL")
                intLinkAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/internet.png"), \
                                                 "Pune link DS/AL pe internet")
                # viberAction = ctxMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/viber.png"), \
                #                                   "Trimite Viber")
                ctxAction = ctxMenu.exec_(self.mapToGlobal(event.pos()))
                if ctxAction == refreshAction:
                    self.centrAlPop()
                if ctxAction == accAction:
                    self.accFunc()
                if ctxAction == refAction:
                    self.refFunc()
                if ctxAction == corectAction:
                    self.corectFunc()
                if ctxAction == dublAction:
                    self.dublFunc()
                if ctxAction == erAction:
                    self.stergeFunc()
                if ctxAction == pregAction:
                    self.pregFunc()
                if ctxAction == admAction:
                    self.admFunc()
                if ctxAction == termAction:
                    self.termLucr()
                if ctxAction == nuLucrAction:
                    self.nuLucr()
                if ctxAction == semnAction:
                    self.semneaza()
                if ctxAction == linkAction:
                    self.showAl()
                if ctxAction == chLinkAction:
                    self.linkTrig()
                if ctxAction == intLinkAction:
                    self.intLinkTrig()
                # if ctxAction == viberAction:
                #     self.viberTrig()

            if self.tableDeranj.hasFocus():
                self.modRow = self.tableDeranj.currentIndex().row()
                self.modColumn = self.tableDeranj.currentIndex().column()
                derMenu = QMenu(self)
                refreshAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/refresh.png"), \
                                                  "Refresh")
                derMenu.addSeparator()
                semnAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/semneaza.png"), \
                                               "Semneaza, deranjament avizat")
                remAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/reparat.png"), \
                                                "Deranjament remediat")
                derAction = derMenu.exec_(self.mapToGlobal(event.pos()))
                if derAction == refreshAction:
                    self.deranjPop()
                if derAction == semnAction:
                    self.semnVaz()
                if derAction == remAction:
                    self.executat()

            if self.tableAng.hasFocus():
                self.contextAng = True
                self.modRow = self.tableAng.currentIndex().row()
                self.modColumn = self.tableAng.currentIndex().column()
                derMenu = QMenu(self)
                refreshAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/refresh.png"), \
                                                  "Refresh")
                # derMenu.addSeparator()
                cautAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/search.png"), \
                                               "Cauta angajat")
                # cautSectAction = derMenu.addAction("Cauta dupa sector")
                # derMenu.addSeparator()
                nouAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/new_circle.png"), \
                                               "Angajat nou")
                modAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/modifica.png"), \
                                                "Modifica starea")
                remAction = derMenu.addAction(QIcon(os.path.abspath(".") + "/Sources/Context/sterge.png"), \
                                              "Sterge angajat")
                derAction = derMenu.exec_(self.mapToGlobal(event.pos()))
                if derAction == cautAction:
                    self.angCautTrig()
                # if derAction == cautSectAction:
                #     self.cautSectTrig()
                if derAction == refreshAction:
                    self.listAngPop()
                if derAction == nouAction:
                    self.angNouTrig()
                if derAction == modAction:
                    self.angModifTrig()
                if derAction == remAction:
                    self.angDelTrig()

        except AttributeError:
            pass

    def alFunc(self):
        myTime = datetime.datetime.now()
        myDelta = datetime.timedelta(days=10)
        myMonthMinus = myTime - myDelta

        myTime_year = myTime.strftime("%Y")
        myTime_month = myTime.strftime("%m")

        myMthMinus_year = myMonthMinus.strftime("%Y")
        myMthMinus_month = myMonthMinus.strftime("%m")

        self.reg_al = None

        for i in self.db.list_collection_names():
            if i == "al_" + myTime_year + "_" + myTime_month:
                self.reg_al = self.db[i]

        if self.reg_al == None:
            myVar = "al_" + myTime_year + "_" + myTime_month
            myVarMin = "al_" + myMthMinus_year + "_" + \
                       myMthMinus_month

            for j in self.db[myVarMin].find().sort("_id", -1).limit(1):
                myNumb = int(j["id"])
                self.reg_al = self.db[myVar]
                self.reg_al.insert_one({
                    "id": myNumb,
                    "oficiul": "",
                    "nr_ds": "",
                    "nr_al": "",
                    "instalatia": "",
                    "pt": "",
                    "localitatea": "",
                    "fid_nr": "",
                    "lucrarile": "",
                    "sef": "",
                    "mem_ech": "",
                    "emitent": "",
                    "cu_dec": "",
                    "mas_teh": "",
                    "semnatura": "",
                    "starea": "",
                    "pregatire": "",
                    "admitere": "",
                    "terminare": "",
                    "link": ""
                })

    #  Functie pentru popularea Registrului de autorizatii
    def regPop(self):
        self.alFunc()
        self.abrOficii()
        myDateTime = datetime.datetime.now()

        for i in self.reg_al.find().sort("_id", -1).limit(1):
            self.idAlDs = int(i["id"]) + 1

        self.searchComma = re.search(",", self.ptLine.text())
        if self.searchComma:
            self.ptText = self.ptLine.text()
            myList = self.ptLine.text().split(",")
            self.ptLine.setText(myList[0])

        if self.okControl:
            if self.ptLine.text() == "PT" or self.ptLine.text() == "":
                pt = ""
                localitatea = ""
            else:
                pt = self.ptLine.text()
            if self.ptFidLine.text() == "Fider nr.":
                fid_nr = ""
            else:
                fid_nr = self.ptFidLine.text()

        else:
            deconect = self.decCombo.currentText()
            if deconect == "Programat" or deconect == "Neprogramat":
                mas_teh = self.ptLine.text() + \
                    "\ndec./mont.\n" + self.ptFidLine.text() + "\n" + self.smNr.text()
                pt = self.ptLine.text()
            else:
                mas_teh = ""
                if self.searchComma:
                    pt = self.ptText
                elif self.ptLine.text() == "PT" or self.ptLine.text() == "":
                    pt = ""
                    localitatea = ""
                else: pt = self.ptLine.text()
        if self.angDoc["position"] == "Dispecer":
            starea = "Acceptat:" + "\n" + self.uCombo.currentText() + \
                "\n" + myDateTime.strftime("%d.%m.%y")
        else:
            starea = "Nou inregistrata"

        if self.memEchLine.text() != "":
            mem_ech = "Formatia: " + self.memEchLine.text()
        else:
            mem_ech = ""

            # Calculez localitatile
        for i in range(2, self.wsPt.max_row + 1):
            if self.wsPt.cell(row=i, column=1).value == self.ptLine.text():
                localitatea = self.wsPt.cell(row=i, column=2).value

        # # Deschid fereastra cu autorizatii, creez linkul
        # if self.okControl == False or \
        #          (self.okControl == True and self.formCheck.isChecked() == False):
        #     dlg = QFileDialog()
        #     self.fileName = dlg.getOpenFileName()
        # elif self.okControl == True and self.formCheck.isChecked() == True:
        #     self.dsForm()
        # link = self.fileName[0]

        if self.nameAngajati.find_one({"name": self.nameFMenu})["position"] == "Dispecer":
            emitent = self.emLine.text()
        else:
            emitent = "Confirmat: \n" + self.nameFMenu + " " + "gr. " + \
                        self.angajati.find_one({"name": self.nameFMenu})["gr_ts"] +\
                                            "\n" + myDateTime.strftime("%d.%m.%y")
        try:
            try:
                # Populez registru internet
                if self.okControl:
                    if self.searchComma:
                        pt = self.ptText
                    # for i in self.db.bir_app_al.find({"nr_ds": {"$ne": ""}}).sort("_id", -1).limit(1):
                    #     self.nrAlDs = int(i["nr_ds"]) + 1
                    #Creez linkul dispozitiei
                    if self.formCheck.isChecked() == True:
                        self.dsForm()
                    elif self.formCheck.isChecked() == False:
                        dlg = QFileDialog()
                        self.fileName = dlg.getOpenFileName()
                    link = self.fileName[0]
                    self.reg_al.insert_one({
                        "id": self.idAlDs,
                        "oficiul": self.ofVar,
                        "nr_ds": self.idAlDs,
                        "nr_al": "",
                        "instalatia": self.instLine.currentText(),
                        "pt": pt,
                        "localitatea": localitatea,
                        "fid_nr": fid_nr,
                        "lucrarile": self.lucrLine.currentText(),
                        "sef": self.sfLine.text(),
                        "mem_ech": mem_ech,
                        "emitent": emitent,
                        "cu_dec": "",
                        "mas_teh": "",
                        "semnatura": "Semnatura",
                        "starea": starea,
                        "pregatire": "Pregatire",
                        "admitere": "Admitere",
                        "terminare": "Terminare",
                        "link": link
                    })
                else:
                    # for i in self.db.bir_app_al.find({"nr_al": {"$ne": ""}}).sort("_id", -1).limit(1):
                    #     self.nrAlDs = int(i["nr_al"]) + 1
                    #Creez linkul autorizatiei
                    dlg = QFileDialog()
                    self.fileName = dlg.getOpenFileName()
                    link = self.fileName[0]
                    self.reg_al.insert_one({
                        "id": self.idAlDs,
                        "oficiul": self.ofVar,
                        "nr_ds": "",
                        "nr_al": self.idAlDs,
                        "instalatia": self.instLine.currentText(),
                        "pt": pt,
                        "localitatea": localitatea,
                        "fid_nr": self.ptFidLine.text(),
                        "lucrarile": self.lucrLine.currentText(),
                        "sef": self.sfLine.text(),
                        "mem_ech": mem_ech,
                        "emitent": emitent,
                        "cu_dec": deconect,
                        "mas_teh": mas_teh,
                        "semnatura": "Semnatura",
                        "starea": starea,
                        "pregatire": "Pregatire",
                        "admitere": "Admitere",
                        "terminare": "Terminare",
                        "link": link
                    })

            except UnboundLocalError:
                self.msSecCall('Nu ati introdus corect denumirea de dispecerat "PT"!\n'
                               'Inregistrarea nu se va efectua!')
        except pymongo.errors.DuplicateKeyError:
            self.msSecCall('Cineva concomitent cu dvs. incearca sa inregistreze DS/AL!\n'
                           'Pentru a exclude dublarea nr. DS/AL mai incercati odata!')


        self.centrAlPop()
        self.dialBox.close()

    def corectieFunc(self):
        self.data.at[self.modRow, 0] = self.ofCorectCombo.currentText()
        # if self.data.at[self.modRow, 1] != "":
        #     self.data.at[self.modRow, 1] = self.nrLine.text()
        # elif self.data.at[self.modRow, 2] != "":
        #     self.data.at[self.modRow, 2] = self.nrLine.text()
        self.data.at[self.modRow, 3] = self.instLine.currentText()
        self.data.at[self.modRow, 4] = self.ptLine.text()
        self.data.at[self.modRow, 6] = self.ptFidLine.text()
        self.data.at[self.modRow, 7] = self.lucrLine.currentText()
        self.data.at[self.modRow, 8] = self.sfLine.text()
        if self.memEchLine.text() != "":
            if re.search("Formatia:", self.memEchLine.text()):
                self.data.at[self.modRow, 9] = self.memEchLine.text()
            else:
                self.data.at[self.modRow, 9] = "Formatia: " + self.memEchLine.text()
        else:
            self.data.at[self.modRow, 9] = ""
        self.data.at[self.modRow, 10] = self.emLine.text()
        self.data.at[self.modRow, 11] = self.decCombo.currentText()
        self.data.at[self.modRow, 12] = self.masLine_corect.text()


        self.reg_al.update_one({"_id": self.data.at[self.modRow, 19]},
                                     {"$set": {
                                         "oficiul": self.data.at[self.modRow, 0],
                                         # "nr_ds": self.data.at[self.modRow, 1],
                                         "instalatia": self.data.at[self.modRow, 3],
                                         "pt": self.data.at[self.modRow, 4],
                                         "fid_nr": self.data.at[self.modRow, 6],
                                         "lucrarile": self.data.at[self.modRow, 7],
                                         "sef": self.data.at[self.modRow, 8],
                                         "mem_ech": self.data.at[self.modRow, 9],
                                         "emitent": self.data.at[self.modRow, 10],
                                         "cu_dec": self.data.at[self.modRow, 11],
                                         "mas_teh": self.data.at[self.modRow, 12],
                                     }})
        # if self.data.at[self.modRow, 2] != "":
        #     self.db.bir_app_al.update_one({"_id": self.data.at[self.modRow, 19]},
        #                                  {"$set": {
        #                                      "oficiul": self.data.at[self.modRow, 0],
        #                                      "nr_al": self.data.at[self.modRow, 2],
        #                                      "instalatia": self.data.at[self.modRow, 3],
        #                                      "pt": self.data.at[self.modRow, 4],
        #                                      "fid_nr": self.data.at[self.modRow, 6],
        #                                      "lucrarile": self.data.at[self.modRow, 7],
        #                                      "sef": self.data.at[self.modRow, 8],
        #                                      "mem_ech": self.data.at[self.modRow, 9],
        #                                      "emitent": self.data.at[self.modRow, 10],
        #                                      "cu_dec": self.data.at[self.modRow, 11],
        #                                      "mas_teh": self.data.at[self.modRow, 12],
        #                                  }})

        self.dialBox.close()

    # def dsInform(self):
        # self.progrDial = QProgressDialog("Generare dispozitie.", "Cancel", 0, 100000)
        # self.progrDial.canceled.connect(self.cancel)
        # self.timerPr = QTimer(self)
        # self.timerPr.timeout.connect(self.perform)
        # self.timerPr.start(0)
        # self.steps = 0
    #     pixmap = QPixmap("ataman_splash.png")
    #     label = QLabel(self)
    # #     # # pixmap = QPixmap('cat.jpg')
    #     label.setPixmap(pixmap)
    #     # # self.setCentralWidget(label)
    #     label.resize(pixmap.width(), pixmap.height())
    #     label.show()

    def dsForm(self):
        # self.msSecCall("In citeva momente se genereaza dispozitia. Asteptati!")
        self.dtDispozitie()

        name = self.sfLine.text()
        name = name.split()
        name = name[0] + " " + name[1]
        for i in self.angajati.find({"name": name}):
            sector = i["sector"]
        # print(name)
        #Creez dispozitia in WORD
        document = Document()
        document.add_heading("SA RED-Nord, of. " + self.ofCombo.currentText() + " " + "Subdiviziunea: " + sector, 0)
        document.add_heading("Dispozitia Nr.: " + str(self.idAlDs), 1)

        #Creez paragraf
        document.add_paragraph()
        p = document.add_paragraph('')
        p.add_run('Sef de lucrari: ').bold = True
        run = p.add_run(self.sfLine.text())
        font = run.font
        font.italic = True
        font.underline = True
        p = document.add_paragraph('')
        p.add_run('Membrii formatiei: ').bold = True
        run = p.add_run(self.memEchLine.text())
        font = run.font
        font.italic = True
        font.underline = True
        p = document.add_paragraph('')
        p.add_run('Sa se execute: ').bold = True
        if self.ptLine.text() == "PT":
            self.ptLine.setText("")
        if self.searchComma:
            self.ptLine.setText(self.ptText)
        if self.ptFidLine.text() == "Fider nr.":
            self.ptFidLine.setText("")
        run = p.add_run(self.instLine.currentText() + " / " + self.ptLine.text() + " / " + self.ptFidLine.text() + \
                  " / " + self.lucrLine.currentText())
        font = run.font
        font.italic = True
        font.underline = True
        p = document.add_paragraph('')
        p.add_run('Masurile pentru pregatirea locului de munca: ').bold = True
        run = document.add_paragraph(style= "List Number").add_run(self.masLine_1.text())
        font = run.font
        font.italic = True
        font.underline = True
        run = document.add_paragraph(style="List Number").add_run(self.masLine_2.text())
        font = run.font
        font.italic = True
        font.underline = True
        run = document.add_paragraph(style="List Number").add_run(self.masLine_3.text())
        font = run.font
        font.italic = True
        font.underline = True
        p = document.add_paragraph('')
        p.add_run('Dispozitia a fost emisa pentru data de: ').bold = True
        run = p.add_run(self.dtLine.text())
        font = run.font
        font.italic = True
        font.underline = True
        p = document.add_paragraph('')
        p.add_run('Emitentul dispozitiei: ').bold = True
        # print(self.emLine.text())
        # for i in self.angajati.find({"name": self.emLine.text()}):
        run = p.add_run(self.emLine.text())
        font = run.font
        font.italic = True
        font.underline = True

        name = self.emLine.text()
        name = name.split()
        name = name[0] + " " + name[1]
        # print(self.angajati.find_one({"name": name})["semnatura_el"])
        try:
            if self.angajati.find_one({"name": name})["semnatura_el"] != "":
                document.add_picture(self.angajati.find_one({"name": name})["semnatura_el"])
            else:
                self.msSecCall('Nu exista linkul la semnatura emitentului!\n'
                               'Dispozitia se va genera fara semnatura!')
        except FileNotFoundError:
            self.msSecCall('Linkul semnaturii emitentului e gresit!\n'
                           'Dispozitia se va genera fara semnatura!')

        document.add_picture(os.path.abspath(".") + "/Sources/Dispozitie/semnat_em.png")
        p = document.add_paragraph('Permisiunea la pregatirea locului de munca ')
        p.add_run("(data, ora)").italic = True
        p.add_run(": ____________   Dispecer ")
        p.add_run("(nume)").italic = True
        p.add_run(": _______").italic = True
        p = document.add_paragraph('Permisiunea la admitere la lucru ')
        p.add_run("(data, ora)").italic = True
        p.add_run(": ____________   Dispecer ")
        p.add_run("(nume)").italic = True
        p.add_run(": ____________________").italic = True
        p = document.add_paragraph('Membrii formatiei au fost supusi instructajului'
                                   'la admitere la locul de lucru pregatit de catre: ')
        p = document.add_paragraph('')
        p.add_run('Seful de lucrari: ').bold = True
        p.add_run('(semnatura): ').italic = True
        p.add_run('_________________________________________')
        p = document.add_paragraph('')
        p.add_run('Membrii formatiei: ').bold = True
        p.add_run('(semnatura): ').italic = True
        p.add_run('_________________________________________')
        p = document.add_paragraph('Lucrarile s-au terminat definitiv, formatia de lucru a fost evacuata: ')
        p = document.add_paragraph('Ora: _______,     ')
        p.add_run("Sef de lucrari: ").bold = True
        p.add_run("(semnatura)").italic = True
        p.add_run(": ______________")



        # print(self.angajati.find_one({"name": name})["semnatura_el"])

        # dlg = QFileDialog()
        # self.fileName = dlg.getSaveFile(self, "Salveaza dispozitia", self.dsPath + \
        #                                     "/" + str(self.nrAlDs) + ".docx",\
        #                                 "Text File (*.txt *.docx *.pdf)")
        # document.save(self.fileName[0])
        # print(self.fileName[0])
        document.save(self.dsPath + "/" + str(self.idAlDs) + ".docx")
        convert(self.dsPath + "/" + str(self.idAlDs) + ".docx", \
                self.dsPath + "/" + str(self.idAlDs) + ".pdf")
        self.fileName = []
        self.fileName.append(self.dsPath + "/" + str(self.idAlDs) + ".pdf")
        os.remove(self.dsPath + "/" + str(self.idAlDs) + ".docx")

    def dtDispozitie(self):
        self.destLoad()
        if self.wsDest.cell(row=4, column=2).value == None:
            self.msCall("destinatia DISPOZITIILOR!")
            self.setTrig()
        else:
            destPath = self.wsDest.cell(row=4, column=2).value

            myDate = self.dtLine.text()

            self.myYear = datetime.datetime.now().strftime("%Y")
            self.myMonth = datetime.datetime.now().strftime("%m")
            myDate = myDate.split(".")

            myPath = destPath + "/" + str(self.myYear)
            dirControl = os.path.isdir(myPath)
            if not dirControl:
                os.mkdir(myPath)

            myMonthPath = myPath + "/" + str(self.myMonth)
            dirControl = os.path.isdir(myMonthPath)
            if not dirControl:
                os.mkdir(myMonthPath)

            self.dsPath = myMonthPath + "/" + myDate[0]
            dirControl = os.path.isdir(self.dsPath)
            if not dirControl:
                os.mkdir(self.dsPath)

    def dtContrDecZl(self):
        self.destLoad()
        if self.wsDest.cell(row=3, column=2).value == None:
            self.msCall("destinatia RAPORT PDJT (excel)!")
            self.setTrig()
        else:
            destPath = self.wsDest.cell(row=3, column=2).value
            # Controlez daca exista mapa cu anul, luna, daca nu exista o creez
            self.myYear = datetime.datetime.now().strftime("%Y")
            self.myMonth = datetime.datetime.now().strftime("%m")
            self.myDay = datetime.datetime.now().strftime("%d")
            self.myHour = datetime.datetime.now().strftime("%H")

            myDate = datetime.datetime.now()
            myDelta = datetime.timedelta(days=1)
            myDateNext = myDate + myDelta
            myDeltaMinus = datetime.timedelta(days=-1)
            myDateMinus = myDate + myDeltaMinus
            myDayMinus = myDateMinus.strftime("%d")

            myPath = destPath + "/" + str(self.myYear)
            dirControl = os.path.isdir(myPath)
            if not dirControl:
                os.mkdir(myPath)

            myMonthPath = myPath + "/" + str(self.myMonth)
            dirControl = os.path.isdir(myMonthPath)
            if not dirControl:
                os.mkdir(myMonthPath)

            if int(self.myHour) >= 8:
                self.rapFile = myMonthPath + "/" + str(self.myDay) + " Raport.xlsx"
            else:
                self.rapFile = myMonthPath + "/" + str(myDayMinus) + " Raport.xlsx"
            fileControl = os.path.isfile(self.rapFile)

            if not fileControl:
                myOriginal = os.path.abspath(".") + "/Bundle/Ungheni/Raport/Raport.xlsx"
                shutil.copyfile(myOriginal, self.rapFile)

            self.wbDecZil = load_workbook(self.rapFile)
            self.wsDecPT = self.wbDecZil["Deconectari PT"]
            self.wsNrSol = self.wbDecZil["Numar solicitari"]

            if int(self.myHour) >= 8:
                if self.wsDecPT.cell(row=2, column=8).value == None:
                    self.wsDecPT.cell(row=2, column=8).value = myDate.strftime("%d.%m.%Y")
                    self.wsDecPT.cell(row=2, column=9).value = "- " + myDateNext.strftime("%d.%m.%Y")
                if self.wsNrSol.cell(row=2, column=6).value == None:
                    self.wsNrSol.cell(row=2, column=6).value = myDate.strftime("%d.%m.%Y")
                    self.wsNrSol.cell(row=2, column=7).value = "- " + myDateNext.strftime("%d.%m.%Y")
                    self.wbDecZil.save(self.rapFile)


    def dtContrSaidi(self):
        self.destLoad()
        if self.wsDest.cell(row=2, column=2).value == None:
            self.msCall("destinatia registru SAIDI (excel)!")
            self.setTrig()
        else:
            destPath = self.wsDest.cell(row=2, column=2).value

            self.myYear = datetime.datetime.now().strftime("%Y")
            self.myMonth = datetime.datetime.now().strftime("%m")

            # Introduc datele in SAIDI cu controlul mapelor si file-l
            myPath = destPath + "/" + str(self.myYear)
            dirControl = os.path.isdir(myPath)
            if not dirControl:
                os.mkdir(myPath)

            self.saidiFile = myPath + "/" + str(self.myMonth) + " Deconectari.xlsx"
            fileControl = os.path.isfile(self.saidiFile)
            if not fileControl:
                myOriginal = os.path.abspath(".") + "/Bundle/Ungheni/Deconectari/Deconectari.xlsx"
                shutil.copyfile(myOriginal, self.saidiFile)

            self.wbDec = load_workbook(self.saidiFile)
            self.wsDecProg = self.wbDec["programate-MT+JT"]
            self.wsDecNeProg = self.wbDec["neprogramate-JT"]


    def dtRegAl(self):
        self.destLoad()
        if self.wsDest.cell(row=1, column=2).value == None:
            self.msCall("destinatia registru AUTORIZATII (excel)!")
            self.setTrig()
        else:
            destPath = self.wsDest.cell(row=1, column=2).value

            self.myYear = datetime.datetime.now().strftime("%Y")

            if destPath != "":
                self.regFile = destPath + "/" + "Registru AUTORIZATIILOR UN " + str(self.myYear) + ".xlsx"
                fielControl = os.path.isfile(self.regFile)
                if not fielControl:
                    myOriginal = os.path.abspath(".") + "/Bundle/Ungheni/Registre/Registru AUTORIZATIILOR UN.xlsx"
                    shutil.copyfile(myOriginal, self.regFile)

            self.wbRegAl = load_workbook(self.regFile)
            self.wsRegAl = self.wbRegAl["Registru"]
            self.wsRegAlLink = self.wbRegAl["alLink"]

    def postgresLoad(self):
        # self.destLoad()
        # if self.wsDest.cell(row=5, column=2).value == None:
        #     self.msCall("destinatia ANALIZA ANUALA (excel)!")
        #     self.setTrig()
        # else:
        #     destPath = self.wsDest.cell(row=5, column=2).value
        #     self.myYear = datetime.datetime.now().strftime("%Y")
        #
        #     if destPath != "":
        #         self.fileAnAnual = destPath + "/" + "Analiza_anuala " + str(self.myYear) + ".xlsx"
        #         fielControl = os.path.isfile(self.fileAnAnual)
        #         if not fielControl:
        #             myOriginal = os.path.abspath(".") + "/Bundle/Ungheni/Deconectari/Analiza_anuala.xlsx"
        #             shutil.copyfile(myOriginal, self.fileAnAnual)
        #
        #     self.wbAnAnual = load_workbook(self.fileAnAnual)
        #     self.wsAnAnualP = self.wbAnAnual["Programat"]
        #     self.wsAnAnualN = self.wbAnAnual["Neprogramat"]

        self.conn = psycopg2.connect(
            host = 'red-nord.cwe4mogj2htg.eu-central-1.rds.amazonaws.com',
            database = 'ungheni',
            user = 'postgres',
            password = '123pdj34'
        )

        self.cur = self.conn.cursor()

    def ofChangeAnaliza(self):
        # self.abrOficii()
        if not self.ofChangeContrSec:
            self.ofChangeContr = True
            self.analizaContr = True
            self.ofAfterCh = self.abrOficiiSec(self.ofCombo.currentText())
            self.decAnaliza()
            self.ofChangeContrSec = False

    def decAnalizaLun(self):
        self.dtContrSaidi()
        #Controlez cind schimb oficiul
        self.ofChangeContr = False
        # Controlez cind schimb oficiul II
        self.ofChangeContrSec = False
        # Controlez analiza lunara
        self.analizaContr = True
        # Controlez analiza anuala
        self.analizaAnContr = 0
        self.decAnaliza()

    # def decAnalizaAn(self):
    #     self.postgresLoad()
    #     self.ofChangeContr = False
    #     self.ofChangeContrSec = False
    #     self.analizaContr = True
    #     self.analizaAnContr = 1
    #     self.DecAnFunc()
    #
    # def DecAnFunc(self):
    #     self.cur.execute('SELECT * FROM anlzan21n')
    #     nepl = self.cur.fetchall()
    #     self.cur.execute('SELECT * FROM anlzan21p')
    #     plan = self.cur.fetchall()
    #     self.cur.close()
    #
    #     column_names = ["anlzan21_id", "ociciul", "pt_fider", "nr_consumatori", "ore"]
    #
    #     self.tblAnN = pd.DataFrame(nepl, columns=column_names)
    #     self.tblAnP = pd.DataFrame(plan, columns=column_names)
    #     print(self.tblAnN)

    def AnualProg(self):
        ofList = ["Toate oficiile"]
        ofList = ofList + self.ofList

        self.ofAnProg = QComboBox()
        self.ofAnProg.addItems(ofList)
        self.ofAnProg.setStyleSheet('padding-left:10%; font-size:12px')
        self.ofAnProg.setFixedHeight(25)
        self.ofAnProg.setFixedWidth(100)
        self.ofAnProg.currentTextChanged.connect(self.AnProgFunc)

        self.AnProgFunc()

    def AnProgFunc(self):
        self.postgresLoad()

        if self.ofAnProg.currentText() == "Toate oficiile":
            self.cur.execute("""SELECT * FROM anlzan21p""")
        else:
            self.cur.execute(
                """SELECT * FROM anlzan21p WHERE oficiul='{}'""".format(self.abrOficiiSec(self.ofAnProg.currentText())))
        tuples = self.cur.fetchall()
        column_names = ["anlzan21p_id",
                        "oficiul",
                        "pt_fider",
                        "localitate",
                        "nr_cons",
                        "ore",
                        "nr_dec",
                        "nr_regl",
                        "compens"
                        ]
        data = pd.DataFrame(tuples, columns=column_names)
        data.sort_values(by="pt_fider", inplace=True, ignore_index=True)
        if data.empty != True:
            for i in range(1, len(data)):
                if data.at[i, "pt_fider"] == data.at[i - 1, "pt_fider"]:
                    # Calculez numarul total de deconectari
                    data.at[i, "nr_dec"] = data.at[i, "nr_dec"] + data.at[i - 1, "nr_dec"]

                    # Calculez timpul total de deconectare
                    myTime = datetime.datetime.strptime(data.at[i, "ore"], '%H:%M:%S')
                    myTimeMinus = datetime.datetime.strptime(data.at[i - 1, "ore"], '%H:%M:%S')
                    delta = datetime.timedelta(
                        hours=myTimeMinus.hour,
                        minutes=myTimeMinus.minute,
                    )
                    myTimeDelta = myTime + delta
                    data.at[i, "ore"] = datetime.datetime.strftime(myTimeDelta, '%H:%M:%S')

                    data.drop(i - 1, inplace=True)
            data.sort_values(by="nr_dec", inplace=True, ignore_index=True, ascending=False)
            # Calculez termenul reglementat urban, rural, suma compensatiilor
            for i in range(len(data)):
                if data.at[i, "localitate"] == None:
                    data.at[i, "localitate"] = ""
                if data.at[i, "compens"] == None:
                    data.at[i, "compens"] = 0
                if data.at[i, "localitate"] == "" or bool(re.search("or[.]", data.at[i, "localitate"])):
                    if data.at[i, "nr_dec"] <= 9:
                        data.at[i, "nr_regl"] = "Incadrat"
                    elif data.at[i, "nr_dec"] > 9:
                        data.at[i, "nr_regl"] = "Depasit cu " + \
                                                str(data.at[i, "nr_dec"] - 9) + "dec."
                        data.at[i, "compens"] = round(0.01 * (160 * 12) * 2.04 * \
                                                      (data.at[i, "nr_dec"] - 9), 2)
                    else:
                        data.at[i, "nr_regl"] = "Eroare"
                else:
                    if data.at[i, "nr_dec"] <= 12:
                        data.at[i, "nr_regl"] = "Incadrat"
                    elif data.at[i, "nr_dec"] > 12:
                        data.at[i, "nr_regl"] = "Depasit cu " + \
                                                str(data.at[i, "nr_dec"] - 12) + "dec."
                        data.at[i, "compens"] = round(0.01 * (160 * 12) * 2.04 * \
                                                      (data.at[i, "nr_dec"] - 12), 2)
                    else:
                        data.at[i, "nr_regl"] = "Eroare"

            header = ["anlzan21p_id",
                      "Oficiul",
                      "PT, Fider",
                      "Localitate",
                      "Nr. cons.",
                      "Ore",
                      "Nr. deconect.",
                      "Nr. reglementat",
                      "Compensatie (lei)"
                      ]

            self.tableAnNepr = QTableView()
            model = TableModel(data, header)

            self.tableAnNepr.setModel(model)
            self.tableAnNepr.setStyleSheet('Background-color: rgb(200, 200, 200)')
            self.tableAnNepr.resizeColumnsToContents()
            self.tableAnNepr.verticalHeader().hide()
            self.tableAnNepr.hideColumn(0)
            self.tableAnNepr.setSelectionBehavior(QAbstractItemView.SelectRows)

            title = QLabel()
            title.setText("Analiza anuala a deconectarilor programate, oficiul:")
            title.setStyleSheet('padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)')

            emptyLb = QLabel("")

            # titleFrame = QFrame()
            hbox = QHBoxLayout()
            hbox.addWidget(title)
            hbox.addWidget(self.ofAnProg)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb)

            # Calculez datele pentru SAIDI
            cons_dec_tot = data.sum(axis=0)["nr_cons"]

            data_ore = pd.DataFrame(pd.to_timedelta(data["ore"]))
            time_delta = data_ore.sum(axis=0)["ore"]
            time_sec = time_delta.days * 24 * 60 * 60 + time_delta.seconds
            time_min = time_sec / 60

            nr_dec_tot = data.sum(axis=0)["nr_dec"]

            if self.ofAnProg.currentText() != "Toate oficiile":
                of_ales = self.abrOficiiSec(self.ofAnProg.currentText())
            else:
                of_ales = 'TOTAL'

            self.cur.execute(f"""SELECT cons_tot FROM saidip
                                        WHERE oficiul = '{of_ales}'""")

            for i in self.cur.fetchall():
                cons_tot = i[0]

            self.cur.execute(f"""UPDATE saidip
                                         SET cons_dec = '{cons_dec_tot}',
                                            t_dec = '{time_min}',
                                            nr_dec_tot = '{nr_dec_tot}',
                                            saidi = '{round(cons_dec_tot * time_min / cons_tot, 2)}',
                                            saifi = '{round(cons_dec_tot / cons_tot, 2)}',
                                            caidi = ROUND (saidi / saifi, 1)
                                        WHERE oficiul = '{of_ales}';
                                        """)
            self.conn.commit()

            # Incarc tabelul dec nepr, mongo
            if self.ofAnProg.currentText() != "Toate oficiile":
                tuples = self.db.deconect_app_deconect.find({
                    "oficiul": self.abrOficiiSec(self.ofAnProg.currentText())
                })
            else:
                tuples = self.db.deconect_app_deconect.find()
            column_names = [
                "oficiul",
                "nr_ordine",
                "pt",
                "fid_04kv",
                "data_dec",
                "data_conect",
                "durata",
                "cons_cas",
                "cons_ec",
                "total",
                "localitate",
                "cauza",
                "termen",
                "compens",
                "id",
            ]

            data = pd.DataFrame(tuples, columns=column_names)
            header = [
                "Oficiul",
                "Nr.",
                "PT",
                "Fider",
                "Data si ora\ndeconectarii",
                "Data si ora\nconectarii",
                "Durata\nintreruperii",
                "Consumatori\ncasnici",
                "Consumatori\nnon-casnici",
                "Total",
                "Localitate",
                "Cauza\ndeconectarii",
                "Termen\nreglementat",
                "Compensatie\n(lei)",
                "id",
            ]
            data.sort_values(by="id", ignore_index=True, ascending=False, inplace=True)

            model = TableModel(data, header)

            tableDec = QTableView()
            tableDec.setModel(model)
            tableDec.setStyleSheet('background-color: rgb(200, 200, 200)')
            tableDec.hideColumn(14)
            tableDec.setColumnWidth(0, 40)
            tableDec.setColumnWidth(1, 40)
            tableDec.setColumnWidth(2, 70)
            tableDec.setColumnWidth(3, 40)
            tableDec.setColumnWidth(4, 70)
            tableDec.setColumnWidth(5, 70)
            tableDec.setColumnWidth(6, 70)
            tableDec.setColumnWidth(7, 40)
            tableDec.setColumnWidth(8, 40)
            tableDec.setColumnWidth(9, 40)
            tableDec.resizeRowsToContents()
            tableDec.verticalHeader().hide()

            tablesFrame = QFrame()
            tb_vbox = QVBoxLayout()
            tb_vbox.addWidget(self.tableAnNepr)
            tb_vbox.addWidget(tableDec)
            tablesFrame.setLayout(tb_vbox)

            # Incarc tabelul saidi
            self.cur.execute('SELECT * FROM saidip ORDER BY saidip_id')
            tuples = self.cur.fetchall()
            column_names = [
                "saidip_id",
                "oficiul",
                "cons_cas",
                "cons_ec",
                "cons_tot",
                "cons_dec",
                "t_dec",
                "nr_dec_tot",
                "saidi",
                "saifi",
                "caidi"
            ]

            data = pd.DataFrame(tuples, columns=column_names)
            header = [
                "saidip_id",
                "Oficiul",
                "Cons.casn.",
                "Cons.non-casn.",
                "Cons.total",
                "Cons.total deconect",
                "Timpul total de deconect",
                "Nr.total de deconect.",
                "SAIDI",
                "SAIFI",
                "CAIDI"
            ]

            # Creez tabelul saidin
            self.tableSaidiN = QTableView()
            model = TableModel(data, header)
            self.tableSaidiN.setModel(model)
            self.tableSaidiN.setStyleSheet('Background-color: rgb(200, 200, 200)')
            self.tableSaidiN.resizeColumnsToContents()
            self.tableSaidiN.verticalHeader().hide()
            self.tableSaidiN.hideColumn(0)
            self.tableSaidiN.setSelectionBehavior(QAbstractItemView.SelectRows)

            # Creez graficul saidin
            canvas = FigureCanvas(Figure(figsize=(1, 5), facecolor=(.78, .78, .78)))

            of_labels = []
            for i in self.ofList:
                of_labels.append(i)

            of_saidi = []
            for i in range(len(data) - 1):
                of_saidi.append(int(round(data.at[i, "saidi"])))

            x = np.arange(len(of_labels))

            ax = canvas.figure.subplots()
            rects = ax.bar(x, of_saidi)

            ax.set_facecolor('#999')
            ax.set_ylabel("SAIDI")
            # ax.set_xlabel("Denumire Oficii")
            ax.set_title("SAIDI pe oficii")
            ax.set_xticks(x)
            ax.set_xticklabels(of_labels, rotation=25)
            # ax.xticks(rotation='vertical')

            ax_Frame = QFrame()
            ax_vbox = QVBoxLayout()
            ax_vbox.addWidget(self.tableSaidiN)
            ax_vbox.addWidget(canvas)
            ax_Frame.setLayout(ax_vbox)
            ax.bar_label(rects)

            splitter = QSplitter(Qt.Horizontal)
            splitter.addWidget(tablesFrame)
            splitter.addWidget(ax_Frame)

            totalFrame = QFrame()
            vbox = QVBoxLayout()
            vbox.addLayout(hbox)
            vbox.addWidget(splitter)
            vbox.setStretch(1, 1)
            totalFrame.setLayout(vbox)

            self.anNeprMdi = QMdiSubWindow()
            self.myMidi.addSubWindow(self.anNeprMdi)
            self.anNeprMdi.setWidget(totalFrame)
            self.anNeprMdi.setWindowIcon(QIcon(QPixmap(1, 1)))
            self.anNeprMdi.setGeometry(100, 100, 1000, 600)
            self.anNeprMdi.showMaximized()
            self.anNeprMdi.show()

            self.cur.close()
        else:
            self.msSecCall(f"Pentru oficiul {self.ofAnProg.currentText()} "
                           f"nu exista date!")

    def AnualNepr(self):
        ofList = ["Toate oficiile"]
        ofList = ofList + self.ofList

        self.ofAnNepr = QComboBox()
        self.ofAnNepr.addItems(ofList)
        self.ofAnNepr.setStyleSheet('padding-left:10%; font-size:12px')
        self.ofAnNepr.setFixedHeight(25)
        # self.ofAnNepr.setFixedWidth(100)
        self.ofAnNepr.currentTextChanged.connect(self.AnNeprFunc)

        self.decNeprCombo = QComboBox()
        self.decNeprCombo.addItems(self.mnList)
        self.decNeprCombo.setStyleSheet('padding-left: 10%; font-size: 12px')
        self.decNeprCombo.setFixedHeight(20)
        self.decNeprCombo.setCurrentText(self.NumbToMonth(self.alMonth))
        self.decNeprCombo.currentTextChanged.connect(self.AnNeprFunc)

        self.AnNeprFunc()

    def AnNeprFunc(self):
        self.postgresLoad()

        if self.ofAnNepr.currentText() == "Toate oficiile":
            self.cur.execute("""SELECT * FROM anlzan21n""")
        else:
            self.cur.execute("""SELECT * FROM anlzan21n WHERE oficiul='{}'""".format(self.abrOficiiSec(self.ofAnNepr.currentText())))
        tuples = self.cur.fetchall()
        column_names = ["anlzan21n_id",
                        "oficiul",
                        "pt_fider",
                        "localitate",
                        "nr_cons",
                        "ore",
                        "nr_dec",
                        "nr_regl",
                        "compens"
        ]
        data = pd.DataFrame(tuples, columns=column_names)
        data.sort_values(by="pt_fider", inplace=True, ignore_index=True)
        if data.empty != True:
            for i in range(1, len(data)):
                if data.at[i, "pt_fider"] == data.at[i-1, "pt_fider"]:
                    #Calculez numarul total de deconectari
                    data.at[i, "nr_dec"] = data.at[i, "nr_dec"] + data.at[i-1, "nr_dec"]

                    #Calculez timpul total de deconectare
                    myTime = datetime.datetime.strptime(data.at[i, "ore"], '%H:%M:%S')
                    myTimeMinus = datetime.datetime.strptime(data.at[i-1, "ore"], '%H:%M:%S')
                    delta = datetime.timedelta(
                        hours=myTimeMinus.hour,
                        minutes=myTimeMinus.minute,
                    )
                    myTimeDelta = myTime + delta
                    data.at[i, "ore"] = datetime.datetime.strftime(myTimeDelta, '%H:%M:%S')

                    data.drop(i-1, inplace=True)
            data.sort_values(by="nr_dec", inplace=True, ignore_index=True, ascending=False)
                # Calculez termenul reglementat urban, rural, suma compensatiilor
            for i in range(len(data)):
                if data.at[i, "localitate"] == None:
                    data.at[i, "localitate"] = ""
                if data.at[i, "compens"] == None:
                    data.at[i, "compens"] = 0
                if data.at[i, "localitate"] == "" or bool(re.search("or[.]", data.at[i, "localitate"])):
                    if data.at[i, "nr_dec"] <= 9:
                        data.at[i, "nr_regl"] = "Incadrat"
                    elif data.at[i, "nr_dec"] > 9:
                        data.at[i, "nr_regl"] = "Depasit cu " + \
                            str(data.at[i, "nr_dec"] - 9) + "dec."
                        data.at[i, "compens"] = round(0.01 * (160 * 12) * 2.04 *\
                                                (data.at[i, "nr_dec"] - 9), 2)
                    else:
                        data.at[i, "nr_regl"] = "Eroare"
                else:
                    if data.at[i, "nr_dec"] <= 12:
                        data.at[i, "nr_regl"] = "Incadrat"
                    elif data.at[i, "nr_dec"] > 12:
                        data.at[i, "nr_regl"] = "Depasit cu " + \
                            str(data.at[i, "nr_dec"] - 12) + "dec."
                        data.at[i, "compens"] = round(0.01 * (160 * 12) * 2.04 * \
                                                (data.at[i, "nr_dec"] - 12), 2)
                    else:
                        data.at[i, "nr_regl"] = "Eroare"

            header = ["anlzan21n_id",
                      "Oficiul",
                      "PT, Fider",
                      "Localitate",
                      "Nr. cons.",
                      "Ore",
                      "Nr. deconect.",
                      "Nr. reglementat",
                      "Compensatie (lei)"
                      ]

            self.tableAnNepr = QTableView()
            model = TableModel(data, header)

            self.tableAnNepr.setModel(model)
            self.tableAnNepr.setStyleSheet('Background-color: rgb(200, 200, 200)')
            self.tableAnNepr.resizeColumnsToContents()
            self.tableAnNepr.verticalHeader().hide()
            self.tableAnNepr.hideColumn(0)
            self.tableAnNepr.setSelectionBehavior(QAbstractItemView.SelectRows)

            title = QLabel()
            title.setText("Analiza deconectarilor neprogramate, oficiul:")
            title.setStyleSheet('padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)')

            # emptyLb = QLabel("")

            # titleFrame = QFrame()
            hbox = QHBoxLayout()
            hbox.addWidget(title)
            hbox.addWidget(self.ofAnNepr)
            hbox.setStretch(1,1)
            # hbox.addWidget(emptyLb)
            # hbox.addWidget(emptyLb)

            # Calculez datele pentru SAIDI
            cons_dec_tot = data.sum(axis=0)["nr_cons"]

            data_ore = pd.DataFrame(pd.to_timedelta(data["ore"]))
            time_delta = data_ore.sum(axis=0)["ore"]
            time_sec = time_delta.days * 24 * 60 * 60 + time_delta.seconds
            time_min = time_sec / 60

            nr_dec_tot = data.sum(axis=0)["nr_dec"]

            if self.ofAnNepr.currentText() != "Toate oficiile":
                of_ales = self.abrOficiiSec(self.ofAnNepr.currentText())
            else:
                of_ales = 'TOTAL'

            self.cur.execute(f"""SELECT cons_tot FROM saidin
                                WHERE oficiul = '{of_ales}'""")

            for i in self.cur.fetchall():
                cons_tot = i[0]

            self.cur.execute(f"""UPDATE saidin
                                 SET cons_dec = '{cons_dec_tot}',
                                    t_dec = '{time_min}',
                                    nr_dec_tot = '{nr_dec_tot}',
                                    saidi = '{round(cons_dec_tot * time_min / cons_tot, 2)}',
                                    saifi = '{round(cons_dec_tot / cons_tot, 2)}',
                                    caidi = ROUND (saidi / saifi, 1)
                                WHERE oficiul = '{of_ales}';
                                """)
            self.conn.commit()

            #Incarc tabelul dec nepr, mongo
#             print(f"""decnepr_{self.alYear}_\
# {self.MonthToNumb(self.decNeprCombo.currentText())}""")
            if self.ofAnNepr.currentText() != "Toate oficiile":
                self.cur.execute(f"""SELECT * FROM decnepr_{self.alYear}_\
{self.MonthToNumb(self.decNeprCombo.currentText())} WHERE oficiul=\
'{self.abrOficiiSec(self.ofAnNepr.currentText())}'""")
            else:
                self.cur.execute(f"""SELECT * FROM decnepr_{self.alYear}_\
{self.MonthToNumb(self.decNeprCombo.currentText())}""")
            tuples = self.cur.fetchall()
            column_names = [
                "decnepr_id",
                "oficiul",
                "nr_ordine",
                "pt",
                "fid_04kv",
                "data_dec",
                "data_conect",
                "durata",
                "cons_cas",
                "cons_ec",
                "total",
                "localitate",
                "cauza",
                "termen",
                "compens",
            ]

            data = pd.DataFrame(tuples, columns=column_names)
            header = [
                "decnepr_id",
                "Oficiul",
                "Nr.",
                "PT",
                "Fider",
                "Data si ora\ndeconectarii",
                "Data si ora\nconectarii",
                "Durata\nintreruperii",
                "Cons.\ncas.",
                "Cons.\nn-cas.",
                "Total",
                "Localitate",
                "Cauza\ndeconectarii",
                "Termen\nreglementat",
                "Compensatie\n(lei)",
            ]
            data.sort_values(by="decnepr_id", ignore_index=True, ascending=False, inplace=True)

            model = TableModel(data, header)

            tableDec = QTableView()
            tableDec.setModel(model)
            tableDec.setStyleSheet('background-color: rgb(200, 200, 200)')
            tableDec.hideColumn(0)
            # tableDec.setColumnWidth(0, 40)
            tableDec.setColumnWidth(1, 40)
            tableDec.setColumnWidth(2, 40)
            tableDec.setColumnWidth(3, 70)
            tableDec.setColumnWidth(4, 70)
            tableDec.setColumnWidth(5, 100)
            tableDec.setColumnWidth(6, 100)
            tableDec.setColumnWidth(7, 70)
            tableDec.setColumnWidth(8, 40)
            tableDec.setColumnWidth(9, 40)
            tableDec.setColumnWidth(10, 40)
            # tableDec.setColumnWidth(14, 70)
            tableDec.resizeRowsToContents()
            tableDec.verticalHeader().hide()
            tableDec.setSelectionBehavior(QAbstractItemView.SelectRows)

            yTitle = QLabel("Tabelul anual:")
            yTitle.setStyleSheet('font-weight: bold; padding-left: 10%')

            mTitle = QLabel("Tabelul lunar:")
            mTitle.setStyleSheet('font-weight: bold; padding-left: 10px')

            titleCombo_hbox = QHBoxLayout()
            titleCombo_hbox.addWidget(mTitle)
            titleCombo_hbox.addWidget(self.decNeprCombo)
            titleCombo_hbox.setStretch(1, 1)

            tablesFrame = QFrame()
            tb_vbox = QVBoxLayout()
            tb_vbox.addWidget(yTitle)
            tb_vbox.addWidget(self.tableAnNepr)
            tb_vbox.addLayout(titleCombo_hbox)
            tb_vbox.addWidget(tableDec)
            tablesFrame.setLayout(tb_vbox)

            #Incarc tabelul saidi
            self.cur.execute('SELECT * FROM saidin ORDER BY saidin_id')
            tuples = self.cur.fetchall()
            column_names = [
                "saidi_id",
                "oficiul",
                "cons_cas",
                "cons_ec",
                "cons_tot",
                "cons_dec",
                "t_dec",
                "nr_dec_tot",
                "saidi",
                "saifi",
                "caidi"
            ]

            data = pd.DataFrame(tuples, columns=column_names)
            header = [
                "saidin_id",
                "Oficiul",
                "Cons.casn.",
                "Cons.non-casn.",
                "Cons.total",
                "Cons.total deconect",
                "Timpul total de deconect",
                "Nr.total de deconect.",
                "SAIDI",
                "SAIFI",
                "CAIDI"
            ]

            #Creez tabelul saidin
            self.tableSaidiN = QTableView()
            model = TableModel(data, header)
            self.tableSaidiN.setModel(model)
            self.tableSaidiN.setStyleSheet('Background-color: rgb(200, 200, 200)')
            self.tableSaidiN.resizeColumnsToContents()
            self.tableSaidiN.verticalHeader().hide()
            self.tableSaidiN.hideColumn(0)
            self.tableSaidiN.setSelectionBehavior(QAbstractItemView.SelectRows)

            #Creez graficul saidin
            canvas = FigureCanvas(Figure(figsize=(1, 5), facecolor=(.78, .78, .78)))

            of_labels = []
            for i in self.ofList:
                of_labels.append(i)

            of_saidi = []
            for i in range(len(data)-1):
                of_saidi.append(int(round(data.at[i, "saidi"])))

            x = np.arange(len(of_labels))

            ax = canvas.figure.subplots()
            rects = ax.bar(x, of_saidi)

            ax.set_facecolor('#999')
            ax.set_ylabel("SAIDI")
            # ax.set_xlabel("Denumire Oficii")
            ax.set_title("SAIDI pe oficii")
            ax.set_xticks(x)
            ax.set_xticklabels(of_labels, rotation=25)
            # ax.xticks(rotation='vertical')

            sdTabTitle = QLabel("Tabelul SAIDI (reprezentare grafica):")
            sdTabTitle.setStyleSheet('padding-left: 10%; font-weight: bold')

            grTitle = QLabel("SAIDI pe oficii:")
            grTitle.setStyleSheet('padding-left: 10%; font-weight: bold')

            ax_Frame = QFrame()
            ax_vbox = QVBoxLayout()
            ax_vbox.addWidget(sdTabTitle)
            ax_vbox.addWidget(self.tableSaidiN)
            ax_vbox.addWidget(grTitle)
            ax_vbox.addWidget(canvas)
            ax_Frame.setLayout(ax_vbox)
            ax.bar_label(rects)

            myScr = app.primaryScreen()
            myScrAv = myScr.availableGeometry()
            myWith = myScrAv.width()
            print(myWith)

            splitter = QSplitter(Qt.Horizontal)
            splitter.addWidget(tablesFrame)
            splitter.addWidget(ax_Frame)
            splitter.handle(1).setStyleSheet('Background-color: rgb(191, 60, 60)')
            splitter.setSizes([myWith - myWith/2.2, myWith/2.2])

            totalFrame = QFrame()
            vbox = QVBoxLayout()
            vbox.addLayout(hbox)
            vbox.addWidget(splitter)
            vbox.setStretch(1, 1)
            totalFrame.setLayout(vbox)


            self.anNeprMdi = QMdiSubWindow()
            self.myMidi.addSubWindow(self.anNeprMdi)
            self.anNeprMdi.setWidget(totalFrame)
            self.anNeprMdi.setWindowIcon(QIcon(QPixmap(1, 1)))
            self.anNeprMdi.setGeometry(100, 100, 1000, 600)
            self.anNeprMdi.showMaximized()
            self.anNeprMdi.show()

            self.cur.close()
        else:
            self.msSecCall(f"Pentru oficiul {self.ofAnProg.currentText()} "
                           f"nu exista date!")

    def decAnaliza(self):
        # self.abrOficii()
        self.ofCombo = QComboBox()
        self.ofCombo.addItems(self.ofList)
        self.ofCombo.setStyleSheet("padding-left:10%; font-size:12px")
        self.ofCombo.setFixedHeight(25)
        self.ofCombo.setFixedWidth(100)
        # self.ofCombo.setEditable(True)
        self.ofCombo.currentTextChanged.connect(self.ofChangeAnaliza)

        self.ptFidPl = []
        self.ptFidPlCons = []
        self.ptFidPlOre = []
        self.ptFidPlOf = []
        self.nrDec = []
        if self.analizaContr and self.analizaAnContr == 0:
            for i in range(19, self.wsDecProg.max_row + 1):
                self.ptFidPl.append(self.wsDecProg.cell(row=i, column=5).value + \
                                self.wsDecProg.cell(row=i, column=6).value)
                self.ptFidPlCons.append(self.wsDecProg.cell(row=i, column=15).value)
                self.ptFidPlOre.append(self.wsDecProg.cell(row=i, column=12).value)
                self.ptFidPlOf.append(self.wsDecProg.cell(row=i, column=2).value)
        elif not self.analizaContr and self.analizaAnContr == 0:
            for i in range(21, self.wsDecNeProg.max_row + 1):
                self.ptFidPl.append(self.wsDecNeProg.cell(row=i, column=4).value + \
                                 self.wsDecNeProg.cell(row=i, column=5).value)
                self.ptFidPlCons.append(self.wsDecNeProg.cell(row=i, column=27).value)
                self.ptFidPlOre.append(self.wsDecNeProg.cell(row=i, column=19).value)
                self.ptFidPlOf.append(self.wsDecNeProg.cell(row=i, column=2).value)
        elif self.analizaAnContr == 1:
            self.cur.execute('SELECT * FROM anlzan21p')
            for i in self.cur.fetchall():
                self.ptFidPl.append(i[2])
                self.ptFidPlCons.append(i[3])
                self.ptFidPlOre.append(i[4])
                self.ptFidPlOf.append(i[1])
        elif self.analizaAnContr == 2:
            self.cur.execute('SELECT * FROM anlzan21n')
            for i in self.cur.fetchall():
                self.ptFidPl.append(i[2])
                self.ptFidPlCons.append(i[3])
                self.ptFidPlOre.append(i[4])
                self.ptFidPlOf.append(i[1])
            self.cur.close()

        for i in self.ptFidPl:
            self.nrDec.append(self.ptFidPl.count(i))

        # print(self.ptFidPlOre)
        self.anTabel = pd.DataFrame()
        for i in range(len(self.ptFidPl)):
            self.anTabel.at[i, 0] = self.ptFidPlOf[i]
            self.anTabel.at[i, 1] = self.ptFidPl[i]
            self.anTabel.at[i, 2] = self.nrDec[i]
            self.anTabel.at[i, 3] = self.ptFidPlOre[i]
            self.anTabel.at[i, 4] = self.ptFidPlCons[i]


        try:
            self.anTabel.sort_values(by=1, ignore_index=True, inplace=True)
        except KeyError:
            if self.analizaContr and self.analizaAnContr == 0:
                self.msSecCall("Deconectari lunare programate nu exista!")
            elif not self.analizaContr and self.analizaAnContr == 0:
                self.msSecCall("Deconectari lunare neprogramate nu exista!")
            elif self.analizaAnContr == 1:
                self.msSecCall("Deconectari anuale programate nu exista!")
            elif self.analizaAnContr == 2:
                self.msSecCall("Deconectari anuale neprogramate nu exista!")

        # print(self.anTabel)
        iterIAn = []
        for i in range(1, len(self.ptFidPl)):
            if self.anTabel.at[i, 1] == self.anTabel.at[i - 1, 1]:
                myTime = datetime.datetime.strptime(self.anTabel.at[i, 3], '%H:%M:%S')
                myTimeMinus = datetime.datetime.strptime(self.anTabel.at[i-1, 3], '%H:%M:%S')
                # myHourAdd = myTime.hour + myTimeMinus.hour
                # myMinuteAdd = myTime.minute + myTimeMinus.minute
                myDeltaMinus = myTimeMinus.hour * 3600 + myTimeMinus.minute *60
                myTimeAdd = myTime + datetime.timedelta(seconds=myDeltaMinus)
                self.anTabel.at[i, 3] = datetime.datetime.strftime(myTimeAdd, '%H:%M:%S')
                self.anTabel.drop([i-1], inplace=True)
        try:
            self.anTabel.sort_values(by=2, ignore_index=True, inplace=True, ascending=False)
        except KeyError:
            pass
        #
        # # print(self.anTabel)
        #
        for i in range(len(self.anTabel)):
            if self.ofChangeContr:
                if self.anTabel.at[i, 0] != self.ofAfterCh:
                    iterIAn.append(i)
            else:
                if self.anTabel.at[i, 0] != self.abrOficiiSec(self.ofCombo.currentText()):
                    iterIAn.append(i)

        for i in range(len(iterIAn)):
            self.anTabel.drop([iterIAn[i]], inplace=True)

        # try:
        #     self.anTabel.sort_values(by=2, ignore_index=True, inplace=True, ascending=False)
        # except KeyError:
        #     pass


    # #Pun graficile analizei
    #
    #     categories = []
    #     set0 = QtCharts.QBarSet("Nr. de deconectari")
    #     set1 = QtCharts.QBarSet("Ore deconectate")
    #     set2 = QtCharts.QBarSet("Nr. de consumatori x10")
    #
    #     for i in range(len(self.anTabel)):
    #         set0.append([float(self.anTabel.at[i, 2])])
    #         set1.append([float(self.anTabel.at[i, 3])])
    #         set2.append([float(self.anTabel.at[i, 1]) / 10])
    #         categories.append(self.anTabel.at[i, 0])
    #
    #     series = QtCharts.QBarSeries()
    #     series.append(set0)
    #     series.append(set1)
    #     series.append(set2)
    #
    #     chart = QtCharts.QChart()
    #     chart.addSeries(series)
    #     if self.analizaContr or self.analizaAnContr == 1:
    #         chart.setTitle("Deconectari programate")
    #     elif not self.analizaContr or self.analizaAnContr == 2:
    #         chart.setTitle("Deconectari neprogramate")
    #     chart.setTitleFont(QFont("Calibri", 14))
    #     # chart.setTitleBrush(QColor(191, 60, 60))
    #     chart.setAnimationOptions(QtCharts.QChart.SeriesAnimations)
    #
    #     axis = QtCharts.QBarCategoryAxis()
    #     axis.append(categories)
    #     chart.createDefaultAxes()
    #     chart.setAxisX(axis, series)
    #
    #     chart.legend().setVisible(True)
    #     chart.legend().setAlignment(Qt.AlignBottom)
    #
    #
    #     self.chartView = QtCharts.QChartView(chart)
    #     self.chartView.setRenderHint(QPainter.Antialiasing)
    #     # self.chartView.setRubberBand(QtCharts.QChartView.HorizontalRubberBand)
        self.anFuncWindow()
    #     import matplotlib
    #
        # print(self.anTabel)

    def anFuncWindow(self):
        header = ["Oficiul",
                  "PT, Fider, nr.",
                  "Numar, deconectari",
                  "Ore deconectate",
                  "Numar consumatori"]
        self.tableAnaliza = QTableView()
        self.modelAnaliza = TableModel(self.anTabel, header)
        self.tableAnaliza.setModel(self.modelAnaliza)
        self.tableAnaliza.setWordWrap(True)
        self.tableAnaliza.setTextElideMode(Qt.ElideMiddle)
        self.tableAnaliza.resizeColumnsToContents()
        self.tableAnaliza.setStyleSheet("Background-color: rgb(200, 200, 200)")
        self.tableAnaliza.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tableAnaliza.verticalHeader().hide()

        if self.analizaContr or self.analizaAnContr == 1:
            self.anFrameP = QFrame()
            vbox = QVBoxLayout()
            vbox.addWidget(self.tableAnaliza)
            self.anFrameP.setLayout(vbox)
            self.analizaContr = False
            if self.analizaAnContr == 1:
                self.analizaAnContr = 2
            self.decAnaliza()
        elif not self.analizaContr or self.analizaAnContr == 2:
            if self.analizaAnContr == 0:
                anLunTitle = QLabel("Analiza lunara de deconectari, oficiul:")
            elif self.analizaAnContr == 2:
                anLunTitle = QLabel("Analiza anuala de deconectari, oficiul:")
            anLunTitle.setStyleSheet("padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)")
            self.anFrameN = QFrame()

            emptyLb = QLabel("")
            emptyLb2 = QLabel("")
            emptyLb3 = QLabel("")
            emptyLb4 = QLabel("")
            # emptyLb.setStyleSheet("margin-right:1300%")

            hbox = QHBoxLayout()
            hbox.addWidget(anLunTitle)
            hbox.addWidget(self.ofCombo)
            hbox.addWidget(emptyLb)
            hbox.addWidget(emptyLb2)
            hbox.addWidget(emptyLb3)
            hbox.addWidget(emptyLb4)
            hbox.setStretch(1, 1)

            vbox = QVBoxLayout()
            vbox.addLayout(hbox)
            vbox.addWidget(self.tableAnaliza)
            self.anFrameN.setLayout(vbox)

            myScrRes = app.primaryScreen()
            myScrAvailable = myScrRes.availableGeometry()
            myHeight = myScrAvailable.height()
            mySplitter = QSplitter(Qt.Vertical)
            mySplitter.addWidget(self.anFrameN)
            mySplitter.addWidget(self.anFrameP)
            mySplitter.handle(1).setStyleSheet("Background-color: rgb(191, 60, 60)")
            mySplitter.setSizes([int(myHeight/5), int(myHeight/5)])

            if self.anControl:
                self.anWindow.close()

            self.anWindow = QMdiSubWindow()
            self.myMidi.addSubWindow(self.anWindow)
            self.anWindow.setWindowIcon(QIcon(QPixmap(1, 1)))
            self.anWindow.setWidget(mySplitter)
            self.anWindow.setGeometry(100, 100, 1000, 600)
            self.anControl = True
            if self.ofChangeContr:
                self.ofChangeContrSec = True
                self.ofCombo.setCurrentText(self.ofAfterCh)
            self.anWindow.showMaximized()
            # self.anWindow.show()

    def deranjFunc(self):
        for i in self.db.der_app_deranj.find().sort("_id", -1).limit(1):
            self.nrDeranj = int(i["id"]) + 1
        if self.ptLine.text() == "PT":
            self.ptLine.setText("")
        if self.ptFidLine.text() == "Fider nr.":
            self.ptFidLine.setText("")
        self.abrOficii()
        pyDateTime = datetime.datetime.now()
        self.db.der_app_deranj.insert_one({
            "id": self.nrDeranj,
            "oficiul": self.ofVar,
            "nr_ordine": str(self.nrDeranj),
            "transmis": self.sfLine.text(),
            "sector": self.sectCombo.currentText(),
            "instalatia": self.instalatiaCombo.currentText(),
            "fid_10kv": self.f10kvLine.text(),
            "pt": self.ptLine.text(),
            "fid_04kv": self.ptFidLine.text(),
            "continutul": self.continText.toPlainText() + "!",
            "data": pyDateTime.strftime("%d.%m.%Y %H:%M"),
            "responsabil": "Semnatura",
            "starea": "Neexecutat"
        })
        # self.db.der_app_deranj.find().sort()
        self.dialBox.close()
        self.deranjPop()

    def listAngPop(self):
        if self.listAngControlPop == True:
            self.tabWindowListAng.close()

        self.data_un = pd.DataFrame(self.db.angajati_un.find({}, {"_id": 0})).sort_values("name")
        self.data_fl = pd.DataFrame(self.db.angajati_fl.find({}, {"_id": 0})).sort_values("name")
        self.data_gl = pd.DataFrame(self.db.angajati_gl.find({}, {"_id": 0})).sort_values("name")
        self.data_rs = pd.DataFrame(self.db.angajati_rs.find({}, {"_id": 0})).sort_values("name")
        self.data = self.data_un.\
            append(self.data_fl).\
            append(self.data_gl).\
            append(self.data_rs)
        myList = self.data.pop("oficiul")
        self.data.insert(0, "oficiul", myList)
        myList = self.data.pop("nr_tabel")
        self.data.insert(0, "nr_tabel", myList)
        myList = self.data.pop("gr_ts")
        self.data.insert(3, "gr_ts", myList)
        self.data = self.data.reset_index()
        self.data.pop("index")
        self.data.pop("semnatura_el")
        # print(self.data)
        header = ["Nr. tabel", "Oficiul", "Numele", "Grupa TS", "Functia", "Sector", "Telefon serviciu", \
                  "Telefon personal", "Emitent AL, DS", "Conducator de lucrari", "Admitent",\
                  "Sef de lucrari", "Supraveghetor", "Membrul echipei"]

        self.model = TableModelList(self.data, header)
        self.tableAng.setModel(self.model)
        self.tableAng.setWordWrap(True)
        self.tableAng.setTextElideMode(Qt.ElideMiddle)
        self.tableAng.resizeColumnsToContents()
        self.tableAng.setStyleSheet("Background-color: rgb(200, 200, 200)")
        self.tableAng.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tableAng.verticalHeader().hide()
        self.tableAng.resizeRowsToContents()

        self.tableAng.setColumnWidth(0, 200)
        self.tableAng.setColumnWidth(1, 200)
        self.tableAng.setColumnWidth(2, 200)
        self.tableAng.setColumnWidth(3, 200)
        self.tableAng.setColumnWidth(4, 200)
        self.tableAng.setColumnWidth(5, 200)
        self.tableAng.setColumnWidth(6, 200)
        self.tableAng.setColumnWidth(7, 200)


        regFrame = QFrame()

        regAlTitle = QLabel()
        regAlTitle.setText("Lista angajatilor.:")
        regAlTitle.setStyleSheet("padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)")
        regAlTitle.move(0, 100)

        self.ofAngCombo = QComboBox()
        myList = ["Toate oficiile"]
        myList = myList + self.ofList
        self.ofAngCombo.addItems(myList)
        # self.ofComboReg.setStyleSheet("margin-right:1400%")
        self.ofAngCombo.setStyleSheet("padding-left:10%; font-size:12px")
        self.ofAngCombo.setFixedHeight(25)
        self.ofAngCombo.setFixedWidth(100)
        self.ofAngCombo.currentTextChanged.connect(self.changeOfAng)


        # Combo pentru sectoare
        self.sectorCombo = QComboBox()
        # self.sectorList = []
        # self.sectCombo.addItems(self.sectorList)
        self.sectorCombo.setStyleSheet("padding-left:10%; font-size:12px")
        self.sectorCombo.setFixedHeight(25)
        self.sectorCombo.setFixedWidth(100)
        # self.sectorCombo.setEditable(True)
        self.sectorCombo.textActivated.connect(self.changeSectAng)

       #Pun functia de mai jos pentru ca cind fac refresh dupa modificare stare sa apara toate filialele
        self.changeOfAng()
        emptyLb = QLabel("")

        hbox = QHBoxLayout()
        hbox.addWidget(regAlTitle)
        hbox.addWidget(self.ofAngCombo)
        hbox.addWidget(self.sectorCombo)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        # hbox.setStretch(1, 1)



        vbox = QVBoxLayout()
        vbox.addLayout(hbox)
        vbox.addWidget(self.tableAng)
        regFrame.setLayout(vbox)

        self.tabWindowListAng = QMdiSubWindow()
        self.myMidi.addSubWindow(self.tabWindowListAng)
        self.tabWindowListAng.setWindowIcon(QIcon(QPixmap(1, 1)))
        self.tabWindowListAng.setWidget(regFrame)
        self.tabWindowListAng.setGeometry(100, 100, 1000, 600)
        self.tabWindowListAng.showMaximized()
        self.listAngControlPop = True
        self.tabWindowListAng.show()

        # if self.angDelFunc_control or self.angFunc_control or self.angModifFunc_control:
        #     self.ofAngCombo.setCurrentText(self.ofCombo_constant)
        #     self.loadSectAng()
        #
        #     self.angDelFunc_control = False
        #     self.angFunc_control = False
        #     self.angModifFunc_control = False

    def deranjPop(self):
        if self.deranjControlPop == True:
            self.tabWindowDeranj.close()

        self.data = pd.DataFrame(self.db.der_app_deranj.find({}, {"_id": 0}))
        myColumn = self.data.pop("id")
        self.data.insert(12, "id", myColumn)
        self.data.columns = range(13)
        # print(self.data)
        self.data.sort_index(ascending=False, inplace=True, ignore_index=True)
        header = ["Oficiul", "Nr.", "Transmis", "Sectorul", "Instalatia", "Fider 10kV", "PT", "Fider 0,4kV",
                  "Continutul", "Data, ora", "Semnatura, Responsabil", "Starea", "id"]

        self.tableDeranj = QTableView()
        self.model = TableModelDeranj(self.data, header)
        self.tableDeranj.setModel(self.model)
        self.tableDeranj.setWordWrap(True)
        self.tableDeranj.setTextElideMode(Qt.ElideMiddle)
        self.tableDeranj.resizeColumnsToContents()
        self.tableDeranj.setStyleSheet("Background-color: rgb(200, 200, 200)")
        self.tableDeranj.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tableDeranj.verticalHeader().hide()
        self.tableDeranj.setColumnWidth(0, 100)
        self.tableDeranj.setColumnWidth(1, 100)
        self.tableDeranj.setColumnWidth(2, 200)
        self.tableDeranj.setColumnWidth(3, 100)
        self.tableDeranj.setColumnWidth(4, 150)
        self.tableDeranj.setColumnWidth(5, 150)
        self.tableDeranj.setColumnWidth(6, 150)
        self.tableDeranj.setColumnWidth(7, 100)
        self.tableDeranj.setColumnWidth(8, 250)
        self.tableDeranj.setColumnWidth(11, 250)
        self.tableDeranj.resizeRowsToContents()

        regFrame = QFrame()
        regAlTitle = QLabel()

        regAlTitle.setText("Registru de deranjamente, oficiul:")
        regAlTitle.setStyleSheet("padding-left: 50%; font-size:24px; color:rgb(191, 60, 60)")
        # regAlTitle.move(0, 100)
        # self.loadOficii()
        myList = ['Toate oficiile']
        myList = myList + self.ofList
        self.ofCombo = QComboBox()
        self.ofCombo.addItems(myList)
        # self.ofComboReg.setStyleSheet("margin-right:1400%")
        self.ofCombo.setStyleSheet("padding-left:10%; font-size:12px")
        self.ofCombo.setFixedHeight(25)
        self.ofCombo.setFixedWidth(100)
        self.ofCombo.currentTextChanged.connect(self.ofChangeDeranj)


        #Combo pentru sectoare
        self.sectCombo = QComboBox()
        self.sectorList = []
        self.sectCombo.addItems(self.sectorList)
        self.sectCombo.setStyleSheet("padding-left:10%; font-size:12px")
        self.sectCombo.setFixedHeight(25)
        self.sectCombo.setFixedWidth(100)
        self.sectCombo.textActivated.connect(self.changeSectDeranj)

        self.execCheck = QCheckBox("Neexecutat")
        self.execCheck.setStyleSheet("padding-left:10%; font-size:12px")
        self.execCheck.stateChanged.connect(self.changeExecDeranj)

        # refreshBt = QPushButton("Refresh")
        # refreshBt.setFixedSize(110, 29)
        # refreshBt.clicked.connect(self.refreshDeranj)

        emptyLb = QLabel("")

        hbox = QHBoxLayout()
        hbox.addWidget(regAlTitle)
        hbox.addWidget(self.ofCombo)
        hbox.addWidget(self.sectCombo)
        hbox.addWidget(self.execCheck)
        # hbox.addWidget(refreshBt)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        hbox.addWidget(emptyLb)
        # hbox.addWidget(emptyLb)
        # hbox.addWidget(emptyLb)
        hbox.setStretch(1, 1)

        vbox = QVBoxLayout()
        vbox.addLayout(hbox)
        vbox.addWidget(self.tableDeranj)
        regFrame.setLayout(vbox)

        self.tabWindowDeranj = QMdiSubWindow()
        self.myMidi.addSubWindow(self.tabWindowDeranj)
        self.tabWindowDeranj.setWindowIcon(QIcon(QPixmap(1, 1)))
        self.tabWindowDeranj.setWidget(regFrame)
        self.tabWindowDeranj.setGeometry(100, 100, 1000, 600)
        self.tabWindowDeranj.showMaximized()
        self.deranjControlPop = True
        self.tabWindowDeranj.show()


    def decFunc(self):
        self.dtContrDecZl()
        # self.dtAnAnual()
        self.abrOficii()

        # Calculez diferenta orelor
        valueDate_5 = self.dtLine.text()
        valueDate_6 = datetime.datetime.now().strftime("%d.%m.%y %H:%M")
        strToDate_5 = datetime.datetime.strptime(valueDate_5, "%d.%m.%y %H:%M")
        strToDate_6 = datetime.datetime.strptime(valueDate_6, "%d.%m.%y %H:%M")
        delta_6_5 = strToDate_6 - strToDate_5
        myList = str(delta_6_5).split()
        try:
            if int(myList[0]):
                myDateSplit = myList[2].split(":")
                myDayPlus = 24 * int(myList[0]) + int(myDateSplit[0])
                myDateStr = str(myDayPlus) + ":" + myDateSplit[1] + ":" + myDateSplit[2]
                myDeltaHour = myDateStr
        except ValueError:
            myDeltaHour = str(delta_6_5)

        #Calculez numarul de consumatori casnici si economici pe faza,
        #Determin localitatea
        for i in range(2, self.wsPt.max_row + 1):
            if self.wsPt.cell(row=i, column=1).value == self.ptLine.text():
                self.totNrCas = int(self.wsPt.cell(row=i, column=4).value)
                self.fazaNrCas = round(self.totNrCas / 9)
                if self.fazaNrCas > 25:
                    self.fazaNrCas = random.randrange(20, 30)

                self.totNrEc = self.wsPt.cell(row=i, column=5).value
                if self.totNrEc <= 2 and self.totNrCas == 0:
                    self.fazaNrEc = 1
                elif self.totNrEc > 12:
                    self.fazaNrEc = random.randrange(1, 3)
                else:
                    self.fazaNrEc = round(self.totNrEc / 9)

                myLocalitate = self.wsPt.cell(row=i, column=2).value

        # Determin incadrarea termenului urban, rural
        myList = myDeltaHour.split(":")
        myMinute = int(myList[0]) * 60 + int(myList[1])
        termText = "Incadrat"
        compens = 0
        if myLocalitate == None or bool(re.search("or[.]", myLocalitate)):
            if myMinute > 6 * 60:
                difMinute = myMinute - 6 * 60
                termText = "Depasit cu: " + str(round(difMinute / 60)) + "H " + myList[1] + "min."
                if difMinute / 60 <= 3:
                    k = 1
                elif (difMinute / 60 > 3) and (difMinute / 60 <= 6):
                    k = 4
                elif (difMinute / 60 > 6) and (difMinute / 60 <= 9):
                    k = 7
                elif (difMinute / 60 > 9):
                    k = 10
                compens = round(0.01 * 160 * 2.04 * k, 2)

        else:
            if myMinute > 12 * 60:
                difMinute = myMinute - 12 * 60
                termText = "Depasit cu: " + str(round(difMinute / 60)) + "H " + myList[1] + "min."
                if difMinute / 60 <= 3:
                    k = 1
                elif (difMinute / 60 > 3) and (difMinute / 60 <= 6):
                    k = 4
                elif (difMinute / 60 > 6) and (difMinute / 60 <= 9):
                    k = 7
                elif (difMinute / 60 > 9):
                    k = 10
                compens = round(0.01 * 160 * 2.04 * k, 2)

        #Working with MongoDB
        for i in self.db.deconect_app_deconect.find().sort("_id", -1).limit(1):
            self.nrDec = int(i["id"]) + 1

        self.db.deconect_app_deconect.insert_one({
            "id": self.nrDec,
            "oficiul": self.ofVar,
            "nr_ordine": self.nrDec,
            "pt": self.ptLine.text(),
            "fid_04kv": self.ptFidLine.text(),
            "data_dec": self.dtLine.text(),
            "data_conect": datetime.datetime.now().strftime("%d.%m.%y %H:%M"),
            "durata": myDeltaHour,
            "cons_cas": str(self.fazaNrCas),
            "cons_ec": str(self.fazaNrEc),
            "total": str(self.fazaNrCas + self.fazaNrEc),
            "localitate": myLocalitate,
            "cauza": self.cauzaCombo.currentText(),
            "compens": compens,
            "termen": termText,
        })

        #Intruduc datele in postgres analiza anuala
        self.postgresLoad()
        self.cur.execute(f"""INSERT INTO anlzan21n (
                                oficiul,
                                pt_fider,
                                localitate,
                                nr_cons,
                                ore,
                                nr_dec
                            ) 
                            VALUES (
                                '{self.ofVar}',
                                '{self.ptLine.text() + " " + self.ptFidLine.text()}',
                                '{myLocalitate}',
                                '{self.fazaNrCas + self.fazaNrEc}',
                                '{myDeltaHour}',
                                '{1}'
                            )"""
        )
        self.conn.commit()
        self.cur.close()

        #Working with excel
        myMaxRow = self.wsDecPT.max_row + 1

        if self.wsDecPT.cell(row=2, column=8).value == None:
            self.wsDecPT.cell(row=2, column=8).value = datetime.datetime.now().strftime("%d.%m.%Y")
        self.wsDecPT.cell(row=myMaxRow, column=1).value = \
            str(int(self.wsDecPT.cell(row=myMaxRow - 1, column=1).value) + 1)
        self.wsDecPT.cell(row=myMaxRow, column=1).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=2).value = self.ofVar
        self.wsDecPT.cell(row=myMaxRow, column=2).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=3).value = self.ptLine.text()
        self.wsDecPT.cell(row=myMaxRow, column=4).value = self.ptFidLine.text()
        self.wsDecPT.cell(row=myMaxRow, column=4).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=5).value = self.dtLine.text()
        self.wsDecPT.cell(row=myMaxRow, column=5).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=6).value = datetime.datetime.now().strftime("%d.%m.%y %H:%M")
        self.wsDecPT.cell(row=myMaxRow, column=6).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=7).value = myDeltaHour
        self.wsDecPT.cell(row=myMaxRow, column=7).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=8).value = str(self.fazaNrCas)
        self.wsDecPT.cell(row=myMaxRow, column=8).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=9).value = str(self.fazaNrEc)
        self.wsDecPT.cell(row=myMaxRow, column=9).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=10).value = \
            str(int(self.wsDecPT.cell(row=myMaxRow, column=8).value) + \
            int(self.wsDecPT.cell(row=myMaxRow, column=9).value))
        self.wsDecPT.cell(row=myMaxRow, column=10).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecPT.cell(row=myMaxRow, column=11).value = myLocalitate
        self.wsDecPT.cell(row=myMaxRow, column=12).value = self.cauzaCombo.currentText()
        self.wsDecPT.cell(row=myMaxRow, column=13).value = termText
        self.wsDecPT.cell(row=myMaxRow, column=13).alignment = \
            Alignment(horizontal="center", vertical="center")
        try:
            self.wbDecZil.save(self.rapFile)
        except PermissionError:
            self.msSecCall("Datele nu s-au introdus in fisierul RAPORT PDJT excel \n"
                           "(cineva foloseste aplicatia)!")

        self.dtContrSaidi()
        myMaxRow = self.wsDecNeProg.max_row + 1
        self.wsDecNeProg.cell(row=myMaxRow, column=1).value = \
            self.wsDecNeProg.cell(row=myMaxRow - 1, column=1).value + 1
        self.wsDecNeProg.cell(row=myMaxRow, column=1).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=2).value = \
            self.ofVar
        self.wsDecNeProg.cell(row=myMaxRow, column=2).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=3).value = \
            datetime.date.today().strftime("%d.%m.%y")
        self.wsDecNeProg.cell(row=myMaxRow, column=3).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=4).value = \
            self.ptLine.text()
        self.wsDecNeProg.cell(row=myMaxRow, column=4).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=5).value = \
            self.ptFidLine.text()
        self.wsDecNeProg.cell(row=myMaxRow, column=5).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=6).value = \
            0
        self.wsDecNeProg.cell(row=myMaxRow, column=6).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=7).value = \
            0
        self.wsDecNeProg.cell(row=myMaxRow, column=7).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=8).value = \
            1
        self.wsDecNeProg.cell(row=myMaxRow, column=8).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=9).value = \
            1
        self.wsDecNeProg.cell(row=myMaxRow, column=9).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=10).value = \
            self.dtLine.text()
        self.wsDecNeProg.cell(row=myMaxRow, column=10).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=11).value = \
            datetime.datetime.now().strftime("%d.%m.%y %H:%M")
        self.wsDecNeProg.cell(row=myMaxRow, column=11).alignment = \
            Alignment(horizontal="center", vertical="center")

        self.wsDecNeProg.cell(row=myMaxRow, column=19).value = \
            myDeltaHour
        self.wsDecNeProg.cell(row=myMaxRow, column=19).alignment = \
            Alignment(horizontal="center", vertical="center")

        self.wsDecNeProg.cell(row=myMaxRow, column=25).value = self.fazaNrCas
        self.wsDecNeProg.cell(row=myMaxRow, column=25).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=26).value = self.fazaNrEc
        self.wsDecNeProg.cell(row=myMaxRow, column=26).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=27).value = \
            self.wsDecNeProg.cell(row=myMaxRow, column=25).value + \
            self.wsDecNeProg.cell(row=myMaxRow, column=26).value
        self.wsDecNeProg.cell(row=myMaxRow, column=27).alignment = \
            Alignment(horizontal="center", vertical="center")
        self.wsDecNeProg.cell(row=myMaxRow, column=28).value = \
            myLocalitate

        self.wsDecNeProg.cell(row=myMaxRow, column=32).value = \
            self.cauzaCombo.currentText()

        try:
            self.wbDec.save(self.saidiFile)
        except PermissionError:
            self.msSecCall("Datele nu s-au introdus in fisierul SAIDI excel (cineva foloseste aplicatia)!")

        # Introduc datele in Excel analiza anuala
        # anAnualMaxRow = self.wsAnAnualN.max_row + 1
        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=1).value = \
        #     self.ofVar
        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=2).value = \
        #     self.ptLine.text()+ " " + self.ptFidLine.text()
        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=3).value = \
        #     self.wsDecNeProg.cell(row=self.wsDecNeProg.max_row, column=25).value + \
        #     self.wsDecNeProg.cell(row=self.wsDecNeProg.max_row, column=26).value
        # self.wsAnAnualN.cell(row=anAnualMaxRow, column=4).value = \
        #     myDeltaHour
        # try:
        #     self.wbAnAnual.save(self.fileAnAnual)
        # except PermissionError:
        #     self.msSecCall("Datele nu vor participa la analiza anuala,\n"
        #                    "(undeva este deschisa analiza anuala excel, nu se permite introducerea datelor)!")

        self.dialBox.close()
        self.decTabWindow()

    def okPass(self):
        self.angDoc = self.angajati.find_one({"name": self.uCombo.currentText()})
        if self.tableCheck.isChecked() == True:
            try:
                if self.angDoc["nr_tabel"] == self.psText.text():
                    self.setWindowTitle('Biroul dispecerului    ' + self.nameFMenu.upper())
                    self.dialInt.close()

                    global closeApp
                    closeApp = False
                else:
                    self.msSecCall("Nu ati introdus corect parola sau userul!")
                self.passControl = True
            except TypeError:
                self.msSecCall("Nu ati introdus corect parola sau userul!")
        else:
            try:
                try:
                    result = self.verify.verification_checks.create(to=self.myTel, code=self.authLine.text())
                    if result.status == 'approved':
                        self.setWindowTitle('Biroul dispecerului    ' + self.nameFMenu.upper())
                        self.dialInt.close()

                        closeApp = False
                    else:
                        self.msSecCall("Nu ati introdus corect codul de autentificare!")
                    self.passControl = True
                except TypeError:
                    self.msSecCall("Nu ati introdus corect userul!")
            except AttributeError:
                self.msSecCall("Generati mai intii codul de autentificare!")

    def twilioFunc(self, telArg):
        client = Client('AC1bd4197015848ea29a103e3a14f6a2ff', '662b14ab4bac80c79b53342d15a05746')
        self.verify = client.verify.services('VAe4432b212966dcf9f009e5e77f35c1d0')
        self.verify.verifications.create(to=telArg, channel='sms')

    def authFunc(self):
        try:
            self.angDoc = self.angajati.find_one({"name": self.uCombo.currentText()})
            if self.angDoc['telefon_serv'] != "" or self.angDoc['telefon_pers'] != "":
                if re.match('[+]373 \d\d\d\d\d\d\d\d', self.angDoc['telefon_serv']):
                    self.myTel = self.angDoc['telefon_serv']
                    self.twilioFunc(self.myTel)
                elif re.match('[+]373 \d\d\d\d\d\d\d\d', self.angDoc['telefon_pers']):
                    self.myTel = self.angDoc['telefon_pers']
                    self.twilioFunc(self.myTel)
                else:
                    self.msSecCall("Nici un numar de telefon al userului nu\n"
                                   "corespunde formatului +373 XXXXXXXX!")
            else:
                self.msSecCall("Pentru utilizatorul dat nu exista numar de telefon!")
        except TypeError:
            self.msSecCall("Nu ati introdus corect userul\n"
                           "sau nu exista un asemenea utilizator!")

    def cancelPass(self):
        self.dialInt.close()
        self.close()
        global closeApp
        closeApp = True
        splash.finish(self)

    def logOut(self):
        self.passControl = False
        self.setWindowTitle('Biroul dispecerului')
        # del self.numeMen
        # print(self.numeMen)
        if self.tabWindowControl == True:
            self.tabWindow.close()
        if self.deranjControlPop == True:
            self.tabWindowDeranj.close()
        if self.listAngControlPop == True:
            self.tabWindowListAng.close()
        if self.tabDecControl == True:
            self.tabWindDec.close()
        if self.anControl:
            self.anWindow.close()

        self.intTrig()

    def angFunc(self):
        self.loadAngSec()
        self.abrOficii()
        self.angajati.insert_one({
            "oficiul": self.ofVar,
            "name": self.numeLine.text(),
            "position": self.functiaCombo.currentText(),
            "nr_tabel": self.tabelLine.text(),
            "sector": self.sectCombo.currentText(),
            "gr_ts": self.grtsCombo.currentText(),
            "telefon_serv": self.telefonLine_serv.text(),
            "telefon_pers": self.telefonLine_pers.text(),
            "emitent": "NU!",
            "conducator": "NU!",
            "admitent": "NU!",
            "sef": "NU!",
            "supraveghetor": "NU!",
            "membru": "NU!",
            "semnatura_el": self.semnLine.text(),
            })
        if self.emCheck.isChecked():
            self.angajati.update_one({
                "name": self.numeLine.text()
            }, {
                "$set": {"emitent": "DA!"}
            })
        if self.condCheck.isChecked():
            self.angajati.update_one({
                "name": self.numeLine.text()
            }, {
                "$set": {"conducator": "DA!"}
            })
        if self.admCheck.isChecked():
            self.angajati.update_one({
                "name": self.numeLine.text()
            }, {
                "$set": {"admitent": "DA!"}
            })
        if self.sefCheck.isChecked():
            self.angajati.update_one({
                "name": self.numeLine.text()
            }, {
                "$set": {"sef": "DA!"}
            })
        if self.suprCheck.isChecked():
            self.angajati.update_one({
                "name": self.numeLine.text()
            }, {
                "$set": {"supraveghetor": "DA!"}
            })
        if self.memCheck.isChecked():
            self.angajati.update_one({
                "name": self.numeLine.text()
            }, {
                "$set": {"membru": "DA!"}
            })

        self.listAngPop()
        self.dialBox.close()

    def angModifFunc(self):
        self.loadAngSec()
        self.angajati.update_one({"name": self.uModifCombo.currentText()},
        {"$set": {
            "position": self.functiaCombo.currentText(),
            "nr_tabel": self.tabelLine.text(),
            "sector": self.sectCombo.currentText(),
            "gr_ts": self.grtsCombo.currentText(),
            "telefon_serv": self.telefonLine_serv.text(),
            "telefon_pers": self.telefonLine_pers.text(),
            "emitent": "NU!",
            "conducator": "NU!",
            "admitent": "NU!",
            "sef": "NU!",
            "supraveghetor": "NU!",
            "membru": "NU!",
            "semnatura_el": self.semnLine.text()
        }
        })
        if self.emCheck.isChecked():
            self.angajati.update_one({
                "name": self.uModifCombo.currentText()
            }, {
                "$set": {"emitent": "DA!"}
            })
        if self.condCheck.isChecked():
            self.angajati.update_one({
                "name": self.uModifCombo.currentText()
            }, {
                "$set": {"conducator": "DA!"}
            })
        if self.admCheck.isChecked():
            self.angajati.update_one({
                "name": self.uModifCombo.currentText()
            }, {
                "$set": {"admitent": "DA!"}
            })
        if self.sefCheck.isChecked():
            self.angajati.update_one({
                "name": self.uModifCombo.currentText()
            }, {
                "$set": {"sef": "DA!"}
            })
        if self.suprCheck.isChecked():
            self.angajati.update_one({
                "name": self.uModifCombo.currentText()
            }, {
                "$set": {"supraveghetor": "DA!"}
            })
        if self.memCheck.isChecked():
            self.angajati.update_one({
                "name": self.uModifCombo.currentText()
            }, {
                "$set": {"membru": "DA!"}
            })

        if self.tableAng.isVisible():
            self.data.at[self.modRow, "nr_tabel"] = self.tabelLine.text()
            self.data.at[self.modRow, "gr_ts"] = self.grtsCombo.currentText()
            self.data.at[self.modRow, "position"] = self.functiaCombo.currentText()
            self.data.at[self.modRow, "sector"] = self.sectCombo.currentText()
            self.data.at[self.modRow, "telefon_serv"] = self.telefonLine_serv.text()
            self.data.at[self.modRow, "telefon_pers"] = self.telefonLine_pers.text()
            if self.emCheck.isChecked():
                self.data.at[self.modRow, "emitent"] = "DA!"
            else:
                self.data.at[self.modRow, "emitent"] = "NU!"
            if self.condCheck.isChecked():
                self.data.at[self.modRow, "conducator"] = "DA!"
            else:
                self.data.at[self.modRow, "conducator"] = "NU!"
            if self.admCheck.isChecked():
                self.data.at[self.modRow, "admitent"] = "DA!"
            else:
                self.data.at[self.modRow, "admitent"] = "NU!"
            if self.sefCheck.isChecked():
                self.data.at[self.modRow, "sef"] = "DA!"
            else:
                self.data.at[self.modRow, "sef"] = "NU!"
            if self.suprCheck.isChecked():
                self.data.at[self.modRow, "supraveghetor"] = "DA!"
            else:
                self.data.at[self.modRow, "supraveghetor"] = "NU!"
            if self.memCheck.isChecked():
                self.data.at[self.modRow, "membru"] = "DA!"
            else:
                self.data.at[self.modRow, "membru"] = "NU!"

        self.dialBox.close()

    def angDelFunc(self):
        self.loadAngSec()
        self.angajati.delete_one({"name": self.uModifCombo.currentText()})
        self.angDelFunc_control = True
        self.ofCombo_constant = self.ofCombo.currentText()
        if self.tableAng.isVisible():
            self.tableAng.hideRow(self.modRow)
            self.data.at[self.modRow, "name"] = ""

        # self.listAngPop()
        self.dialBox.close()

    def okTrigger(self):
        self.okControl = False
        if self.sfLine.text() == "":
            self.msCall('"Sef de lucrari:"')
            self.sfLine.setFocus()
        elif self.emLine.text() == "":
            self.msCall('"Emitent:"')
            self.emLine.setFocus()
        elif self.instLine.currentText() == "":
            self.msCall('"Se executa (tipul instalatiei (LEA, ID etc.)):"')
            self.instLine.setFocus()
        elif self.lucrLine.currentText() == "Lucrarile efectuate:"\
                or self.lucrLine.currentText() == "":
            self.msCall('"Lucrarile efectuate:"')
            self.lucrLine.setFocus()
        elif self.decCombo.currentText() == "Programat" or self.decCombo.currentText() == "Neprogramat":
            if self.ptLine.text() == "PT" or self.ptLine.text() == "":
                self.msCall('"PT"')
                self.ptLine.setFocus()
            elif self.ptFidLine.text() == "Fider nr.":
                self.msCall('"Fider nr."')
                self.ptFidLine.setFocus()
            elif self.smNr.text() == "SM nr.":
                self.msCall('"SM nr."')
                self.smNr.setFocus()
            else:
                self.regPop()
        else:
            self.regPop()

    def okAng(self):
        if self.numeLine.text() == "":
            self.msCall('"Nume:"')
            self.numeLine.setFocus()
        elif self.tabelLine.text() == "":
            self.msCall('"Nr. tabel:"')
            self.numeLine.setFocus()
        elif self.functiaCombo.currentText() == "Alege functia:":
            self.msCall('"Functia:"')
            self.functiaCombo.setFocus()
        elif self.sectCombo.currentText() == "Alege sector:":
            self.msCall('"Sector:"')
            self.sectCombo.setFocus()
        else:
            self.angFunc()

    def okAngCaut(self):
        if self.uModifCombo.currentText() == "":
            self.msCall('"Nume:"')
            self.uModifCombo.setFocus()
        else:
            self.dialBox.close()
            self.listAngPop()
            self.ofAngCombo.setCurrentText(self.ofCombo.currentText())
            self.sectorCombo.setCurrentText(self.angajati.find_one({"name": self.uModifCombo.currentText()})["sector"])
            self.changeSectAng()
            self.cautAngControl = False
            self.cautSecAngControl = False

    def okAngModif(self):
        if self.uModifCombo.currentText() == "":
            self.msCall('"Nume:"')
            self.uModifCombo.setFocus()
        else:
            self.angModifFunc()

    def okAngDel(self):
        if self.uModifCombo.currentText() == "":
            self.msCall('"Nume:"')
            self.uModifCombo.setFocus()
        else:
            self.angDelFunc()

    def okDsTrigger(self):
        # self.vboxPr.addWidget(self.progressDs)
        self.okControl = True
        if self.sfLine.text() == "":
            self.msCall('"Sef de lucrari:"')
            self.sfLine.setFocus()
        # elif self.memEchLine.text() == "":
        #     self.msCall('"Membrii echipei:"')
        #     self.memEchLine.setFocus()
        elif self.emLine.text() == "":
            self.msCall('"Emitent:"')
            self.emLine.setFocus()
        # elif self.instLine.text() == "Instalatia:":
        #     self.msCall('"Instalatia:"')
        #     self.instLine.setFocus()
        # elif self.ptLine.text() == "PT" or self.ptLine.text() == "":
        #     self.msCall('"PT:"')
        #     self.ptLine.setFocus()
        elif self.lucrLine.currentText() == "Lucrarile efectuate:" \
                or self.lucrLine.currentText() == "":
            self.msCall('"Lucrarile efectuate:"')
            self.lucrLine.setFocus()
        elif self.formCheck.isChecked() == True and self.dtLine.text() == "":
            self.msCall('"Alege data:"')
            self.dtButton.setFocus()

        else:

            # self.dialBox.close()
            # self.progrDial =  QProgressDialog("Generare dispozitie.", "Cancel", 0, 100000)
            # self.progrDial.canceled.connect(self.cancel)
            # self.timerPr = QTimer(self)
            # self.timerPr.timeout.connect(self.perform)
            # self.timerPr.start(0)
            # self.steps = 0

            self.regPop()

    # def perform(self):
    #     self.progrDial.setValue(self.steps)
    #     self.steps += 1
    #     if self.steps > self.progrDial.maximum():
    #         self.timerPr.stop()
    #         # self.regPop()
    #
    #
    # def cancel(self):
    #     self.timerPr.stop()


    def okDecTrigger(self):
        if self.ptLine.text() == "PT" or self.ptLine.text() == "":
            self.msCall('"PT:"')
            self.ptLine.setFocus()
        elif self.ptFidLine.text() == "Fider nr.":
            self.msCall('"Fider nr.')
            self.ptFidLine.setFocus()
        else:
            self.decFunc()

    def okDeranjTrigger(self):
        if self.sectCombo.currentText() == "Alege:" or self.sectCombo.currentText() == "":
            self.msCall('"Sectorul:"')
            self.sectCombo.setFocus()
        elif self.instalatiaCombo.currentText() == "Alege:" or self.instalatiaCombo.currentText() == "":
            self.msCall('"Instalatia:"')
            self.instalatiaCombo.setFocus()
        elif (self.instalatiaCombo.currentText() == "LEA-0.4kV" or \
                self.instalatiaCombo.currentText() == "ID-10kV" or \
                self.instalatiaCombo.currentText() == "ID-0.4kV" or \
                self.instalatiaCombo.currentText() == "Consumator casnic") and \
                (self.ptLine.text() == "PT" or self.ptLine.text() == ""):
            self.msCall('"PT"')
            self.ptLine.setFocus()
        elif (self.instalatiaCombo.currentText() == "LEA-10kV" or \
                self.instalatiaCombo.currentText() == "LC-10kV") and \
                self.f10kvLine.text() == "":
            self.msCall('"Fid.10kV nr:"')
            self.f10kvLine.setFocus()
        elif self.continText.toPlainText() == "":
            self.msCall('"Continutul:"')
            self.continText.setFocus()
        else:
            self.deranjFunc()

mw = mainWindow()
try:
    if not closeApp:
        mw.show()
        splash.finish(mw)
        app.exec_()
except NameError:
    pass